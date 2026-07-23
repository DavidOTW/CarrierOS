from __future__ import annotations

import asyncio
import csv
import hashlib
import io
import json
import logging
import os
import re
import secrets
import sqlite3
import time
import uuid
from contextlib import asynccontextmanager
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any
from urllib.parse import quote, urlencode, urlsplit

import stripe
from docx import Document
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from itsdangerous import BadSignature, SignatureExpired, URLSafeTimedSerializer
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.middleware.sessions import SessionMiddleware

from .calculations import (
    calculate_quote,
    driver_monthly_fixed,
    fuel_price_for,
    monday_for,
    parse_date,
)
from .audits import AuditFileError, MAX_AUDIT_FILE_BYTES, audit_uploaded_document
from .opportunities import (
    calculate_opportunity,
    compare_drivers,
    offered_revenue,
    opportunity_input_snapshot,
)
from .routing import configured_route_provider
from .dispatch_workflow import HOS_DISCLAIMER, rank_assignments
from .delivery_workflow import parse_delivery_document_kind, validate_driver_transition
from .load_states import LoadState, LoadStateError, normalize_state, transition_load_state
from .money import money_to_cents
from .permissions import has_permission
from .ratecons import (
    MAX_RATECON_BYTES,
    MATERIAL_CLASSIFICATIONS,
    ExtractedField,
    RateConError,
    compare_ratecon_to_booking,
    configured_extraction_provider,
    configured_malware_scan_provider,
    configured_ocr_provider,
    configured_storage_provider,
    default_retention_date,
    document_text_pages,
    suggest_ratecon_matches,
    validate_ratecon_upload,
)
from .release_readiness import evaluate_release_readiness
from .growth import STARTUP_STEPS, equipment_finance_audit, growth_mentor_findings
from .db import (
    as_dict,
    create_database_backup,
    create_password_reset_token,
    db_session,
    execute,
    export_organization_data,
    hash_password,
    init_db,
    new_onboarding_token,
    password_needs_rehash,
    query_all,
    query_one,
    record_audit_event,
    reset_password_with_token,
    token_digest,
    utc_now_iso,
    verify_password,
)
from .emailing import send_password_reset_email, smtp_configured
from .services import (
    LOAD_SORT_OPTIONS,
    filter_and_sort_loads,
    get_bundle,
    get_state,
    loads_with_results,
    selected_month,
    summarize_load_rows,
    summarize_load_rows_by_driver,
    summarize_load_rows_by_month,
)
from .stripe_billing import (
    BillingConfigurationError,
    construct_webhook_event,
    create_checkout_session,
    create_portal_session,
    first_subscription_price_id,
    object_value,
    plan_code_for_price,
    price_id_for_plan,
    stripe_configured,
    stripe_live_configured,
    unix_date,
)

BASE_DIR = Path(__file__).resolve().parent
VERSION = "0.16.0a4"
ENVIRONMENT = os.getenv("CARRIEROS_ENV", "development").strip().lower()
IS_PRODUCTION = ENVIRONMENT == "production"
CANONICAL_BASE_URL = os.getenv(
    "CARRIEROS_CANONICAL_URL", "https://otwcarrieros.com"
).strip().rstrip("/")
SESSION_SECRET = os.getenv("CARRIEROS_SECRET", "")
if IS_PRODUCTION and len(SESSION_SECRET) < 32:
    raise RuntimeError("CARRIEROS_SECRET must be at least 32 characters in production.")
SESSION_SECRET = SESSION_SECRET or "development-only-change-me-v04"
BILLING_MODE = os.getenv("CARRIEROS_BILLING_MODE", "stripe").strip().lower()
if BILLING_MODE not in {"stripe", "beta"}:
    raise RuntimeError("CARRIEROS_BILLING_MODE must be 'stripe' or 'beta'.")
TRIAL_DAYS = 14
TERMS_VERSION = "2026-07-21-audit-v2"
SUPPORT_EMAIL = os.getenv(
    "CARRIEROS_SUPPORT_EMAIL", "david@outsidethewirelogistics.com"
).strip().lower()
FOUNDER_LINKEDIN_URL = "https://www.linkedin.com/in/davidbryant89"
BACKUP_INTERVAL_SECONDS = max(
    3600, int(os.getenv("CARRIEROS_BACKUP_INTERVAL_HOURS", "24")) * 3600
)
logger = logging.getLogger("carrieros")


class SensitiveAccessLogFilter(logging.Filter):
    """Keep reset and onboarding bearer tokens out of platform request logs."""

    def filter(self, record: logging.LogRecord) -> bool:
        if isinstance(record.args, tuple) and len(record.args) >= 3:
            values = list(record.args)
            path = str(values[2])
            if path.startswith("/reset-password?"):
                values[2] = "/reset-password?[redacted]"
            elif path.startswith("/onboard/"):
                values[2] = "/onboard/[redacted]"
            record.args = tuple(values)
        return True


logging.getLogger("uvicorn.access").addFilter(SensitiveAccessLogFilter())


def customer_signups_open() -> bool:
    return not (
        IS_PRODUCTION and BILLING_MODE == "stripe" and not stripe_live_configured()
    )

PLAN_LIMITS = {
    "carrier_startup": {"name": "Carrier Startup", "units": 0, "price": 10},
    "owner_operator": {"name": "Owner-Operator", "units": 2, "price": 25},
    "starter_fleet": {"name": "Starter Fleet", "units": 5, "price": 50},
    "small_fleet": {"name": "Small Fleet", "units": 10, "price": 75},
    "growing_fleet": {"name": "Growing Fleet", "units": 20, "price": 100},
}

QUICK_LINK_CATEGORIES = (
    "Load board",
    "Broker portal",
    "Fuel & routing",
    "Finance",
    "Other",
)
QUICK_LINK_LIMIT = 30

SEO_PAGES = {
    "small-fleet-trucking-software": {
        "title": "Small Fleet Trucking Software | CarrierOS",
        "description": "Small fleet trucking software for owner-operators and carriers with 1–20 trucks. Connect dispatch, driver pay, expenses, and estimated profit per load.",
        "eyebrow": "Small fleet trucking software",
        "heading": "Run a small trucking fleet without running it from five spreadsheets.",
        "lead": "CarrierOS gives owner-operators and small carrier teams one browser-based workspace for loads, driver pay, operating costs, receivables, and the profit each load actually keeps.",
        "audience": "Built for owner-operators growing beyond one truck, fleet owners managing 2–20 power units, and the dispatch or office people keeping those operations moving.",
        "problem_title": "A practical operating system for the part of the market enterprise TMS platforms overlook.",
        "problem_copy": "Small fleets need more than a load list, but they should not need enterprise software, dedicated IT staff, or a long implementation. CarrierOS keeps the daily operating picture focused on the decisions that affect cash and margin.",
        "benefits": [
            ("Dispatch and load control", "Keep lanes, dates, drivers, units, revenue, miles, and load status connected."),
            ("Driver pay flexibility", "Use profit split, contractor gross split, owner-operator split, flat per load, loaded mile, total mile, or day rate by driver."),
            ("Profit per load", "Compare revenue with fuel, driver pay, maintenance, direct costs, fixed costs, and overhead."),
            ("Carrier administration", "Track receivables, detention, compliance dates, onboarding records, and operating documents."),
        ],
        "workflow_title": "From booked load to a clearer carrier result",
        "workflow": [
            ("Enter the load", "Record the lane, dates, miles, revenue, assigned driver, and power unit."),
            ("Apply your real costs", "CarrierOS uses your fuel, maintenance, fixed-cost, fee, and driver-pay assumptions."),
            ("Review the result", "See estimated carrier profit, margin, driver obligation, and the items that need attention."),
        ],
        "faqs": [
            ("Who is CarrierOS built for?", "Owner-operators and small U.S. motor carriers that want a clearer view of dispatch, driver pay, and profitability without enterprise complexity."),
            ("Does CarrierOS replace my ELD or accounting system?", "No. CarrierOS is an operations and profitability workspace. It does not replace an ELD, payroll provider, tax professional, or regulated accounting system."),
            ("Can office users and drivers be added?", "Plans are based on active power units and include unlimited driver records and office users."),
        ],
    },
    "driver-settlement-software": {
        "title": "Driver Pay Tracking for Small Fleets | CarrierOS",
        "description": "Driver pay tracking for small trucking fleets. Estimate seven driver pay structures, record payments, and keep load pay and carrier profit connected.",
        "eyebrow": "Driver pay tracking",
        "heading": "Driver pay that follows the agreement—not a one-size-fits-all formula.",
        "lead": "CarrierOS helps small carriers calculate, review, and track driver pay across seven compensation structures while keeping the load economics visible.",
        "audience": "For carrier owners, dispatchers, and back-office teams paying company drivers, contractors, and owner-operators under different agreements.",
        "problem_title": "Make every driver-pay calculation easier to explain and repeat.",
        "problem_copy": "Mixed pay arrangements become fragile when the math lives in memory or separate spreadsheets. CarrierOS keeps each driver's method attached to the operating record so pay and carrier margin can be reviewed together.",
        "benefits": [
            ("Profit split", "Calculate the agreed driver share after the allowed load costs in your operating model."),
            ("Loaded and total miles", "Apply a driver-specific rate to loaded miles or total paid miles."),
            ("Contractor and owner-operator splits", "Apply an agreed share to gross revenue or the configured profit base."),
            ("Flat and day rate", "Use a fixed amount per load or per qualifying trip day."),
        ],
        "workflow_title": "A consistent driver-pay workflow",
        "workflow": [
            ("Choose the driver's model", "Set the compensation structure and rate that match the operating agreement."),
            ("Connect pay to the load", "Use the load's revenue, miles, dates, and allowed costs to calculate the obligation."),
            ("Track what was paid", "Record payments against cumulative driver or contractor balances for a clearer audit trail."),
        ],
        "faqs": [
            ("Which driver pay methods are supported?", "Profit split, contractor gross split, owner-operator split, flat per load, loaded mile, total mile, and day rate."),
            ("Can different drivers use different pay models?", "Yes. CarrierOS keeps compensation settings driver-specific, so a fleet can use mixed models."),
            ("Is CarrierOS a payroll processor?", "No. CarrierOS calculates and tracks operating pay obligations; it does not file payroll taxes or replace professional payroll, tax, legal, or accounting advice."),
        ],
    },
    "load-profitability-calculator": {
        "title": "Truck Load Profitability Calculator | CarrierOS",
        "description": "Calculate truck load profitability using revenue, miles, fuel, driver pay, maintenance, fixed costs, fees, and overhead. Try the free CarrierOS live demo.",
        "eyebrow": "Truck load profitability calculator",
        "heading": "Know what a load can keep—not only what it pays per mile.",
        "lead": "CarrierOS turns a load's rate, miles, fuel, driver pay, direct costs, fixed costs, and company assumptions into a clearer estimate of carrier profit and margin.",
        "audience": "For owner-operators, dispatchers, and small fleet owners comparing freight opportunities or reviewing completed load performance.",
        "problem_title": "Rate per mile is useful. It is not the whole profit story.",
        "problem_copy": "Two loads with the same rate per mile can produce different results after deadhead, fuel, driver compensation, maintenance, fixed-cost days, fees, and overhead. CarrierOS keeps those inputs together so the tradeoff is visible before or after booking.",
        "benefits": [
            ("Cost-based lane quotes", "Estimate break-even, target, and opening quote from your operating assumptions."),
            ("Fuel and maintenance", "Apply fuel price, MPG, maintenance reserve, and total-mile assumptions."),
            ("Driver pay", "Compare the effect of the driver's actual compensation model on the load result."),
            ("Carrier margin", "See estimated profit dollars and margin instead of relying on gross revenue alone."),
        ],
        "workflow_title": "Turn a broker offer into a profit decision",
        "workflow": [
            ("Enter the freight economics", "Add rate, loaded and empty miles, dates, and direct load expenses."),
            ("Apply company assumptions", "Use the unit, driver, fuel, maintenance, fees, fixed costs, and overhead that fit your operation."),
            ("Compare the outcome", "Review break-even, target quote, driver pay, carrier profit, and margin before relying on the number."),
        ],
        "faqs": [
            ("What costs should be included in load profitability?", "CarrierOS can model fuel, driver or contractor pay, maintenance, direct load costs, fixed costs, company fees, and overhead using the assumptions you enter."),
            ("Can I use the calculator before booking a load?", "Yes. The lane quote tool is designed to compare break-even, target, and opening quote values before a load is accepted."),
            ("Are the calculated results financial advice?", "No. Results are operational estimates based on user-entered assumptions and should be reviewed before business, tax, accounting, payroll, or compliance decisions."),
        ],
    },
}

INDEXABLE_PATHS = {"/", "/demo", *(f"/{slug}" for slug in SEO_PAGES)}
PUBLIC_CRAWL_FILES = {"/robots.txt", "/sitemap.xml"}
LOGIN_WINDOW_SECONDS = 15 * 60
LOGIN_MAX_ATTEMPTS = 10
SIGNUP_WINDOW_SECONDS = 60 * 60
SIGNUP_MAX_ATTEMPTS = 5
RESET_WINDOW_SECONDS = 60 * 60
RESET_MAX_ATTEMPTS = 5
login_attempts: dict[str, list[float]] = {}
signup_attempts: dict[str, list[float]] = {}
reset_attempts: dict[str, list[float]] = {}


async def database_backup_worker() -> None:
    while True:
        await asyncio.sleep(BACKUP_INTERVAL_SECONDS)
        try:
            await asyncio.to_thread(create_database_backup)
        except Exception:
            logger.exception("Scheduled CarrierOS database backup failed")

@asynccontextmanager
async def lifespan(_: FastAPI):
    init_db()
    backup_task: asyncio.Task[None] | None = None
    if IS_PRODUCTION:
        try:
            await asyncio.to_thread(create_database_backup)
        except Exception:
            logger.exception("CarrierOS startup database backup failed")
        backup_task = asyncio.create_task(database_backup_worker())
    try:
        yield
    finally:
        if backup_task:
            backup_task.cancel()
            try:
                await backup_task
            except asyncio.CancelledError:
                pass


app = FastAPI(title="CarrierOS", version=VERSION, lifespan=lifespan)
app.add_middleware(
    SessionMiddleware,
    secret_key=SESSION_SECRET,
    session_cookie="carrieros_session",
    same_site="lax",
    https_only=IS_PRODUCTION,
    max_age=60 * 60 * 12,
)


class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        path = request.url.path
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
        response.headers["Cross-Origin-Opener-Policy"] = "same-origin"
        response.headers["Cross-Origin-Resource-Policy"] = "same-origin"
        response.headers["X-Permitted-Cross-Domain-Policies"] = "none"
        if request.method == "GET" and (path in INDEXABLE_PATHS or path in PUBLIC_CRAWL_FILES):
            response.headers["Cache-Control"] = "public, max-age=300, stale-while-revalidate=3600"
        elif not path.startswith("/static/"):
            response.headers["Cache-Control"] = "no-store"
        if path not in INDEXABLE_PATHS and path not in PUBLIC_CRAWL_FILES and not path.startswith("/static/"):
            response.headers["X-Robots-Tag"] = "noindex, nofollow"
        production_upgrade = "; upgrade-insecure-requests" if IS_PRODUCTION else ""
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; img-src 'self' data:; style-src 'self' 'unsafe-inline'; "
            "script-src 'self' 'unsafe-inline'; font-src 'self'; frame-ancestors 'none'; "
            f"base-uri 'self'; form-action 'self'{production_upgrade}"
        )
        if IS_PRODUCTION:
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response


app.add_middleware(SecurityHeadersMiddleware)
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))


def money(value: Any) -> str:
    try:
        return f"${float(value):,.2f}"
    except (TypeError, ValueError):
        return "$0.00"


def number(value: Any, default: float = 0.0) -> float:
    try:
        if value in (None, ""):
            return default
        return float(value)
    except (TypeError, ValueError):
        return default


def optional_number(value: Any) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def integer(value: Any, default: int = 0) -> int:
    try:
        if value in (None, ""):
            return default
        return int(value)
    except (TypeError, ValueError):
        return default


def yes(value: Any) -> bool:
    return str(value or "").strip().lower() in {"1", "yes", "true", "on"}


def normalized_external_url(value: Any) -> str | None:
    """Return a safe browser link without ever requesting the destination."""
    raw = str(value or "").strip()
    if not raw or len(raw) > 2048 or any(char.isspace() for char in raw):
        return None
    if "://" not in raw and ":" in raw.split("/", 1)[0]:
        return None
    candidate = raw if "://" in raw else f"https://{raw}"
    try:
        parsed = urlsplit(candidate)
        parsed.port
    except ValueError:
        return None
    if parsed.scheme.lower() not in {"http", "https"} or not parsed.hostname:
        return None
    if parsed.username or parsed.password:
        return None
    return parsed.geturl()


def callable_phone(value: Any) -> str | None:
    raw = str(value or "").strip()
    digits = "".join(character for character in raw if "0" <= character <= "9")
    if not 7 <= len(digits) <= 15:
        return None
    return f"+{digits}" if raw.startswith("+") else digits


def _dispatch_datetime(value: Any) -> str:
    raw = str(value or "").strip()
    if not raw:
        return ""
    try:
        parsed = datetime.fromisoformat(raw)
    except ValueError:
        return raw.replace("T", " ")
    return parsed.strftime("%a %b %d, %Y %-I:%M %p") if os.name != "nt" else parsed.strftime("%a %b %d, %Y %#I:%M %p")


def _stop_navigation_links(address: Any) -> dict[str, str]:
    destination = str(address or "").strip()
    if not destination:
        return {"apple": "", "google": ""}
    return {
        "apple": f"https://maps.apple.com/?{urlencode({'daddr': destination, 'dirflg': 'd'})}",
        "google": "https://www.google.com/maps/dir/?"
        + urlencode({"api": "1", "destination": destination, "travelmode": "driving"}),
    }


def _stop_dispatch_lines(load: dict[str, Any], prefix: str, label: str) -> tuple[list[str], dict[str, str]]:
    address = str(load.get(f"{prefix}_address") or "").strip()
    start = _dispatch_datetime(load.get(f"{prefix}_window_start"))
    end = _dispatch_datetime(load.get(f"{prefix}_window_end"))
    contact_name = str(load.get(f"{prefix}_contact_name") or "").strip()
    contact_phone = str(load.get(f"{prefix}_contact_phone") or "").strip()
    instructions = str(load.get(f"{prefix}_instructions") or "").strip()
    timezone_name = str(load.get(f"{prefix}_timezone") or "").strip()
    navigation = _stop_navigation_links(address)
    lines = [f"{label}: {start} - {end} ({timezone_name})", address]
    if contact_name or contact_phone:
        lines.append("Contact: " + " | ".join(value for value in (contact_name, contact_phone) if value))
    if instructions:
        lines.append(f"Instructions: {instructions}")
    lines.extend((f"Apple Maps: {navigation['apple']}", f"Google Maps: {navigation['google']}"))
    return lines, navigation


def driver_dispatch_package(load: dict[str, Any]) -> dict[str, Any]:
    driver = load.get("driver") or {}
    driver_phone = callable_phone(driver.get("phone"))
    blockers: list[str] = []
    if not load.get("ratecon_received_at"):
        blockers.append("Mark the RateCon as received and verified.")
    if not driver:
        blockers.append("Assign a driver.")
    elif not driver_phone:
        blockers.append("Add a valid mobile number to the assigned driver's profile.")
    for prefix, label in (("pickup", "pickup"), ("delivery", "delivery")):
        if not str(load.get(f"{prefix}_address") or "").strip():
            blockers.append(f"Add the full {label} address.")
        if not load.get(f"{prefix}_window_start") or not load.get(f"{prefix}_window_end"):
            blockers.append(f"Add the complete {label} appointment window.")
        if not str(load.get(f"{prefix}_timezone") or "").strip():
            blockers.append(f"Add the IANA time zone for the {label} facility.")
    if blockers:
        return {"ready": False, "blockers": blockers}

    pickup_lines, pickup_navigation = _stop_dispatch_lines(load, "pickup", "PICKUP")
    delivery_lines, delivery_navigation = _stop_dispatch_lines(load, "delivery", "DELIVERY")
    header = [f"CarrierOS dispatch | Load {load.get('load_number') or ''}"]
    if load.get("broker"):
        header.append(f"Broker/customer: {load['broker']}")
    if load.get("ratecon_reference"):
        header.append(f"RateCon ref: {load['ratecon_reference']}")
    message = "\n".join(
        header
        + ["", *pickup_lines, "", *delivery_lines, "", "Please confirm receipt and report any address or appointment discrepancy before departure."]
    )
    encoded_message = quote(message, safe="")
    return {
        "ready": True,
        "driver_name": driver.get("name") or "Driver",
        "phone": driver_phone,
        "call_url": f"tel:{driver_phone}",
        "sms_iphone_url": f"sms:{driver_phone}&body={encoded_message}",
        "sms_android_url": f"sms:{driver_phone}?body={encoded_message}",
        "message": message,
        "pickup_navigation": pickup_navigation,
        "delivery_navigation": delivery_navigation,
    }


def _ratecon_form_fields(form: Any, existing: Any = None) -> dict[str, Any]:
    received = yes(form.get("ratecon_received"))
    existing_received_at = str(existing["ratecon_received_at"] or "") if existing is not None else ""
    fields = {
        "pickup_address": str(form.get("pickup_address", "")).strip(),
        "pickup_window_start": str(form.get("pickup_window_start", "")).strip() or None,
        "pickup_window_end": str(form.get("pickup_window_end", "")).strip() or None,
        "pickup_contact_name": str(form.get("pickup_contact_name", "")).strip(),
        "pickup_contact_phone": str(form.get("pickup_contact_phone", "")).strip(),
        "pickup_instructions": str(form.get("pickup_instructions", "")).strip(),
        "pickup_timezone": str(form.get("pickup_timezone", "")).strip(),
        "delivery_address": str(form.get("delivery_address", "")).strip(),
        "delivery_window_start": str(form.get("delivery_window_start", "")).strip() or None,
        "delivery_window_end": str(form.get("delivery_window_end", "")).strip() or None,
        "delivery_contact_name": str(form.get("delivery_contact_name", "")).strip(),
        "delivery_contact_phone": str(form.get("delivery_contact_phone", "")).strip(),
        "delivery_instructions": str(form.get("delivery_instructions", "")).strip(),
        "delivery_timezone": str(form.get("delivery_timezone", "")).strip(),
        "ratecon_reference": str(form.get("ratecon_reference", "")).strip(),
        "ratecon_received_at": (existing_received_at or utc_now_iso()) if received else None,
    }
    limits = {
        "pickup_address": 500, "delivery_address": 500,
        "pickup_contact_name": 200, "delivery_contact_name": 200,
        "pickup_contact_phone": 100, "delivery_contact_phone": 100,
        "pickup_instructions": 2000, "delivery_instructions": 2000,
        "ratecon_reference": 255,
        "pickup_timezone": 100, "delivery_timezone": 100,
    }
    for name, limit in limits.items():
        if len(str(fields[name] or "")) > limit:
            raise HTTPException(400, f"{name.replace('_', ' ').title()} is too long")
    for prefix, label in (("pickup", "Pickup"), ("delivery", "Delivery")):
        start = fields[f"{prefix}_window_start"]
        end = fields[f"{prefix}_window_end"]
        try:
            parsed_start = datetime.fromisoformat(start) if start else None
            parsed_end = datetime.fromisoformat(end) if end else None
        except ValueError as exc:
            raise HTTPException(400, f"Enter a valid {label.lower()} appointment time") from exc
        if parsed_start and parsed_end and parsed_end < parsed_start:
            raise HTTPException(400, f"{label} window end must be after its start")
        timezone_name = str(fields[f"{prefix}_timezone"] or "")
        if timezone_name and not re.fullmatch(r"[A-Za-z_+-]+(?:/[A-Za-z_+-]+)+", timezone_name):
            raise HTTPException(400, f"Enter a valid IANA time zone for {label.lower()} (example: America/Chicago)")
    return fields


def driver_availability(organization_id: int) -> list[dict[str, Any]]:
    rows = query_all(
        """SELECT d.id AS driver_id, d.name AS driver_name, d.phone,
        d.vehicle_id AS assigned_vehicle_id, assigned_vehicle.name AS assigned_vehicle_name,
        scheduled_load.id AS load_id, scheduled_load.load_number,
        scheduled_load.delivery_date, scheduled_load.destination, scheduled_load.status,
        scheduled_load.vehicle_id AS load_vehicle_id, load_vehicle.name AS load_vehicle_name
        FROM drivers d
        LEFT JOIN vehicles assigned_vehicle
          ON assigned_vehicle.id=d.vehicle_id AND assigned_vehicle.organization_id=d.organization_id
        LEFT JOIN loads scheduled_load ON scheduled_load.id=(
          SELECT candidate.id FROM loads candidate
          WHERE candidate.organization_id=d.organization_id
            AND candidate.driver_id=d.id
            AND lower(trim(candidate.status)) NOT IN ('cancelled', 'canceled', 'quote')
          ORDER BY
            CASE WHEN candidate.delivery_date IS NULL OR candidate.delivery_date='' THEN 1 ELSE 0 END,
            candidate.delivery_date DESC,
            candidate.id DESC
          LIMIT 1
        )
        LEFT JOIN vehicles load_vehicle
          ON load_vehicle.id=scheduled_load.vehicle_id
          AND load_vehicle.organization_id=d.organization_id
        WHERE d.organization_id=? AND d.active=1
        ORDER BY lower(d.name)""",
        (organization_id,),
    )
    today = date.today()
    availability = []
    for row in rows:
        item = as_dict(row) or {}
        delivery = parse_date(item.get("delivery_date"))
        status = str(item.get("status") or "").strip()
        completed = status.lower() == "delivered"
        destination = str(item.get("destination") or "Destination not set")
        item["vehicle_name"] = item.get("load_vehicle_name") or item.get("assigned_vehicle_name") or "No unit assigned"
        item["phone_href"] = callable_phone(item.get("phone"))
        item["delivery_display"] = (
            delivery.strftime("%b %d, %Y").replace(" 0", " ") if delivery else None
        )
        item["next_empty_date"] = item.get("delivery_date") or None
        item["next_empty_date_display"] = item["delivery_display"]
        item["next_empty_location"] = destination if item.get("load_id") else None
        if not item.get("load_id"):
            item.update(
                availability_label="Available now",
                availability_detail="No scheduled load",
                availability_tone="good",
            )
        elif not delivery:
            item.update(
                availability_label="Delivery date needed",
                availability_detail=f"{item['load_number']} to {destination}",
                availability_tone="warn",
            )
        elif completed or delivery < today:
            item.update(
                availability_label="Available now",
                availability_detail=f"Last delivery: {destination} · {item['delivery_display']}",
                availability_tone="good",
            )
        else:
            item.update(
                availability_label=f"Available {item['delivery_display']}",
                availability_detail=f"Delivering {item['load_number']} to {destination}",
                availability_tone="info",
            )
        availability.append(item)
    return availability


def percent(value: Any, digits: int = 1) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return "0.0%"


def current_user(request: Request) -> dict[str, Any] | None:
    user_id = request.session.get("user_id")
    if not user_id:
        return None
    row = query_one(
        """SELECT u.*, o.name AS organization_name, o.owner_name, o.owner_email,
        o.source_filename, o.source_sync_date, o.plan_code, o.active_unit_limit,
        o.subscription_status, o.trial_ends_at, o.billing_customer_reference,
        o.billing_subscription_reference, o.billing_price_reference,
        o.subscription_current_period_end, o.subscription_cancel_at_period_end
        FROM users u JOIN organizations o ON o.id=u.organization_id WHERE u.id=?""",
        (user_id,),
    )
    return as_dict(row)


def subscription_allows_access(user: dict[str, Any]) -> bool:
    status = str(user.get("subscription_status") or "").lower()
    if status == "active":
        return True
    if status != "trialing":
        return False
    trial_end = parse_date(user.get("trial_ends_at"))
    return bool(trial_end and trial_end >= date.today())


def require_user(request: Request) -> dict[str, Any]:
    user = current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Login required")
    if not subscription_allows_access(user):
        raise HTTPException(status_code=402, detail="Subscription required")
    return user


def require_permission(request: Request, permission: str) -> dict[str, Any]:
    user = require_user(request)
    if not has_permission(str(user.get("role") or "read_only"), permission):
        raise HTTPException(status_code=403, detail="You do not have permission for this action")
    return user


WORKFLOW_LABELS = {
    LoadState.BOOKED_AWAITING_RATECON: "Booked — Awaiting RateCon",
    LoadState.RATECON_REVIEW: "RateCon Review",
    LoadState.NEEDS_ASSIGNMENT: "Needs Assignment",
    LoadState.DISPATCH_AWAITING_APPROVAL: "Dispatch Awaiting Approval",
    LoadState.DISPATCHED_AWAITING_ACK: "Dispatched — Awaiting Driver Acknowledgment",
    LoadState.DISPATCH_ACKNOWLEDGED: "Dispatch Acknowledged",
}


def _transition_workflow(
    conn: Any,
    *,
    organization_id: int,
    load_id: int,
    target: LoadState,
    actor_user_id: int | None,
    idempotency_key: str,
    reason: str,
) -> Any:
    try:
        history = transition_load_state(
            conn,
            organization_id=organization_id,
            load_id=load_id,
            target=target,
            actor_user_id=actor_user_id,
            idempotency_key=idempotency_key,
            reason=reason,
        )
    except LoadStateError as exc:
        raise HTTPException(409, str(exc)) from exc
    conn.execute(
        "UPDATE loads SET status=? WHERE id=? AND organization_id=?",
        (WORKFLOW_LABELS.get(target, target.value.replace("_", " ").title()), load_id, organization_id),
    )
    return history


def _load_by_public_uuid(organization_id: int, public_uuid: str) -> dict[str, Any]:
    row = query_one(
        "SELECT * FROM loads WHERE public_uuid=? AND organization_id=?",
        (public_uuid, organization_id),
    )
    if not row:
        raise HTTPException(404, "Load not found")
    return dict(row)


def _ratecon_document(organization_id: int, public_uuid: str) -> dict[str, Any]:
    row = query_one(
        """SELECT d.*,l.public_uuid AS load_public_uuid,l.load_number
        FROM operational_documents d
        LEFT JOIN loads l ON l.id=d.load_id AND l.organization_id=d.organization_id
        WHERE d.public_uuid=? AND d.organization_id=? AND d.deleted_at IS NULL""",
        (public_uuid, organization_id),
    )
    if not row:
        raise HTTPException(404, "RateCon not found")
    return dict(row)


def _effective_extracted_fields(
    organization_id: int, extraction_id: int, *, conn: Any | None = None
) -> list[ExtractedField]:
    sql = """SELECT * FROM ratecon_extracted_fields
    WHERE organization_id=? AND extraction_id=? ORDER BY id"""
    params = (organization_id, extraction_id)
    rows = conn.execute(sql, params).fetchall() if conn is not None else query_all(sql, params)
    return [
        ExtractedField(
            str(row["field_name"]),
            str(row["reviewed_value"] if row["reviewed_value"] is not None else row["extracted_value"]),
            int(row["confidence_millis"]) / 1000,
            int(row["document_page"]) if row["document_page"] is not None else None,
            str(row["evidence_text"]),
            str(row["bounding_reference"]) if row["bounding_reference"] else None,
        )
        for row in rows
        if str(row["human_review_status"]) != "REJECTED"
    ]


def _replace_ratecon_differences(
    conn: Any,
    *,
    organization_id: int,
    document_id: int,
    load: dict[str, Any],
    fields: list[ExtractedField],
) -> list[dict[str, Any]]:
    approved = conn.execute(
        """SELECT 1 FROM ratecon_differences
        WHERE organization_id=? AND document_id=? AND approval_status='APPROVED' LIMIT 1""",
        (organization_id, document_id),
    ).fetchone()
    if approved:
        raise HTTPException(409, "Approved RateCon differences are locked")
    conn.execute(
        "DELETE FROM ratecon_differences WHERE organization_id=? AND document_id=?",
        (organization_id, document_id),
    )
    differences = compare_ratecon_to_booking(load, fields)
    for difference in differences:
        conn.execute(
            """INSERT INTO ratecon_differences
            (organization_id,document_id,load_id,field_name,booked_value,ratecon_value,
             classification,financial_impact_cents,operational_impact,confidence_millis,
             evidence_text,approval_status)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                organization_id,
                document_id,
                load["id"],
                difference.field_name,
                difference.booked_value,
                difference.ratecon_value,
                difference.classification,
                difference.financial_impact_cents,
                difference.operational_impact,
                round(difference.confidence * 1000),
                difference.evidence,
                "PENDING" if difference.material else "NOT_REQUIRED",
            ),
        )
    return [difference.to_dict() for difference in differences]


def _insert_load_snapshot(
    conn: Any,
    *,
    organization_id: int,
    load_id: int,
    stage: str,
    inputs: dict[str, Any],
    result: dict[str, Any],
    user_id: int,
) -> int:
    revision = int(
        conn.execute(
            "SELECT COALESCE(MAX(revision),0)+1 AS revision FROM load_financial_snapshots WHERE load_id=? AND stage=?",
            (load_id, stage),
        ).fetchone()["revision"]
    )
    input_json = json.dumps(inputs, sort_keys=True, separators=(",", ":"), default=str)
    result_json = json.dumps(result, sort_keys=True, separators=(",", ":"), default=str)
    checksum = hashlib.sha256(f"{input_json}|{result_json}".encode()).hexdigest()
    return int(
        conn.execute(
            """INSERT INTO load_financial_snapshots
            (public_uuid,organization_id,load_id,stage,revision,calculation_version,
             input_json,result_json,checksum_sha256,created_by)
            VALUES (?,?,?,?,?,'0.16-phase2',?,?,?,?)""",
            (
                str(uuid.uuid4()), organization_id, load_id, stage, revision,
                input_json, result_json, checksum, user_id,
            ),
        ).lastrowid
    )


def _dispatch_token(approval_public_uuid: str) -> str:
    serializer = URLSafeTimedSerializer(SESSION_SECRET, salt="carrieros-driver-dispatch-v1")
    return serializer.dumps({"scope": "driver_dispatch_ack", "approval": approval_public_uuid})


def _read_dispatch_token(token: str) -> str:
    serializer = URLSafeTimedSerializer(SESSION_SECRET, salt="carrieros-driver-dispatch-v1")
    try:
        payload = serializer.loads(token, max_age=48 * 3600)
    except SignatureExpired as exc:
        raise HTTPException(410, "This dispatch link expired. Ask dispatch for a new link.") from exc
    except BadSignature as exc:
        raise HTTPException(404, "Dispatch link not found") from exc
    if payload.get("scope") != "driver_dispatch_ack" or not payload.get("approval"):
        raise HTTPException(404, "Dispatch link not found")
    return str(payload["approval"])


def render(request: Request, name: str, context: dict[str, Any] | None = None, status_code: int = 200):
    context = context or {}
    csrf_token = request.session.get("csrf_token")
    if not csrf_token:
        csrf_token = secrets.token_urlsafe(32)
        request.session["csrf_token"] = csrf_token
    context.update({
        "request": request,
        "user": current_user(request),
        "money": money,
        "percent": percent,
        "today": date.today(),
        "version": VERSION,
        "support_email": SUPPORT_EMAIL,
        "billing_mode": BILLING_MODE,
        "signups_open": customer_signups_open(),
        "signup_href": (
            "/signup"
            if customer_signups_open()
            else f"mailto:{SUPPORT_EMAIL}?subject=CarrierOS%20launch%20access"
        ),
        "signup_cta": (
            "Start free trial" if customer_signups_open() else "Get launch access"
        ),
        "email_delivery_ready": smtp_configured(),
        "csrf_token": csrf_token,
        "flash": request.session.pop("flash", None) if hasattr(request, "session") else None,
    })
    return templates.TemplateResponse(request=request, name=name, context=context, status_code=status_code)


def seo_context(
    path: str,
    title: str,
    description: str,
    *,
    page_type: str = "WebPage",
    faqs: list[tuple[str, str]] | None = None,
) -> dict[str, Any]:
    canonical_url = f"{CANONICAL_BASE_URL}{path}"
    graph: list[dict[str, Any]] = [
        {
            "@type": "Organization",
            "@id": f"{CANONICAL_BASE_URL}/#organization",
            "name": "Outside The Wire Logistics LLC",
            "url": "https://www.outsidethewirelogistics.com/",
            "email": SUPPORT_EMAIL,
            "brand": {"@type": "Brand", "name": "CarrierOS"},
            "founder": {"@id": f"{CANONICAL_BASE_URL}/#founder"},
        },
        {
            "@type": "Person",
            "@id": f"{CANONICAL_BASE_URL}/#founder",
            "name": "David Bryant",
            "jobTitle": "Founder of CarrierOS and Outside The Wire Logistics LLC",
            "description": "United States Marine Corps combat veteran, Purple Heart recipient, and logistics leader with 20 years of experience.",
            "award": "Purple Heart",
            "sameAs": [FOUNDER_LINKEDIN_URL],
        },
        {
            "@type": page_type,
            "@id": f"{canonical_url}#webpage",
            "url": canonical_url,
            "name": title,
            "description": description,
            "inLanguage": "en-US",
            "isPartOf": {"@id": f"{CANONICAL_BASE_URL}/#website"},
            "publisher": {"@id": f"{CANONICAL_BASE_URL}/#organization"},
        },
        {
            "@type": "WebSite",
            "@id": f"{CANONICAL_BASE_URL}/#website",
            "url": f"{CANONICAL_BASE_URL}/",
            "name": "CarrierOS",
            "description": "Fleet operations and profitability software for owner-operators and small motor carriers.",
            "publisher": {"@id": f"{CANONICAL_BASE_URL}/#organization"},
        },
    ]
    if path == "/":
        graph.append(
            {
                "@type": "WebApplication",
                "@id": f"{CANONICAL_BASE_URL}/#software",
                "name": "CarrierOS",
                "url": f"{CANONICAL_BASE_URL}/",
                "applicationCategory": "BusinessApplication",
                "applicationSubCategory": "Trucking management software",
                "operatingSystem": "Any operating system with a modern web browser",
                "browserRequirements": "Requires JavaScript and a modern web browser",
                "description": description,
                "featureList": [
                    "Dispatch and load tracking",
                    "Seven driver compensation structures",
                    "Driver settlement tracking",
                    "Load profitability calculations",
                    "Pre-book rate offers, profit checks, and negotiation tracking",
                    "Immutable quote-to-load booking snapshots",
                    "Receivables and detention tracking",
                ],
                "audience": {
                    "@type": "BusinessAudience",
                    "audienceType": "Owner-operators and small motor carriers with 1 to 20 power units",
                },
                "provider": {"@id": f"{CANONICAL_BASE_URL}/#organization"},
                "offers": [
                    {
                        "@type": "Offer",
                        "name": plan["name"],
                        "price": plan["price"],
                        "priceCurrency": "USD",
                        "description": f"Up to {plan['units']} active power units; unlimited driver records and office users.",
                        "url": f"{CANONICAL_BASE_URL}/signup?plan={code}",
                    }
                    for code, plan in PLAN_LIMITS.items()
                ],
            }
        )
    if faqs:
        graph.append(
            {
                "@type": "FAQPage",
                "mainEntity": [
                    {
                        "@type": "Question",
                        "name": question,
                        "acceptedAnswer": {"@type": "Answer", "text": answer},
                    }
                    for question, answer in faqs
                ],
            }
        )
    return {
        "seo_title": title,
        "seo_description": description,
        "canonical_url": canonical_url,
        "robots_content": "index, follow, max-image-preview:large, max-snippet:-1, max-video-preview:-1",
        "social_image_url": f"{CANONICAL_BASE_URL}/static/carrieros-launch-og.png",
        "founder_linkedin_url": FOUNDER_LINKEDIN_URL,
        "structured_data": {"@context": "https://schema.org", "@graph": graph},
    }


async def verified_form(request: Request):
    form = await request.form()
    if not IS_PRODUCTION:
        return form
    expected = str(request.session.get("csrf_token") or "")
    supplied = str(form.get("_csrf") or "")
    if not expected or not supplied or not secrets.compare_digest(expected, supplied):
        raise HTTPException(status_code=403, detail="Invalid form token")
    return form


def valid_password(password: str) -> bool:
    return (
        len(password) >= 12
        and any(char.islower() for char in password)
        and any(char.isupper() for char in password)
        and any(char.isdigit() for char in password)
        and any(not char.isalnum() for char in password)
    )


def client_key(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for", "").split(",", 1)[0].strip()
    return forwarded or (request.client.host if request.client else "unknown")


def login_rate_limited(request: Request) -> bool:
    key = client_key(request)
    cutoff = time.time() - LOGIN_WINDOW_SECONDS
    recent = [stamp for stamp in login_attempts.get(key, []) if stamp >= cutoff]
    login_attempts[key] = recent
    return len(recent) >= LOGIN_MAX_ATTEMPTS


def signup_rate_limited(request: Request) -> bool:
    key = client_key(request)
    cutoff = time.time() - SIGNUP_WINDOW_SECONDS
    recent = [stamp for stamp in signup_attempts.get(key, []) if stamp >= cutoff]
    signup_attempts[key] = recent
    return len(recent) >= SIGNUP_MAX_ATTEMPTS


def reset_rate_limited(request: Request) -> bool:
    key = client_key(request)
    cutoff = time.time() - RESET_WINDOW_SECONDS
    recent = [stamp for stamp in reset_attempts.get(key, []) if stamp >= cutoff]
    reset_attempts[key] = recent
    return len(recent) >= RESET_MAX_ATTEMPTS


def redirect(url: str) -> RedirectResponse:
    return RedirectResponse(url=url, status_code=303)


def public_url(request: Request) -> str:
    configured = os.getenv("CARRIEROS_PUBLIC_URL", "").strip().rstrip("/")
    if configured:
        return configured
    render_hostname = os.getenv("RENDER_EXTERNAL_HOSTNAME", "").strip().rstrip("/")
    if render_hostname:
        return f"https://{render_hostname}"
    if IS_PRODUCTION:
        raise BillingConfigurationError("CARRIEROS_PUBLIC_URL is not configured.")
    return str(request.base_url).rstrip("/")


def set_flash(request: Request, text: str) -> None:
    request.session["flash"] = text


def compliance_status(expiration_date: str | None) -> tuple[str, int | None]:
    exp = parse_date(expiration_date)
    if not expiration_date:
        return "Missing date", None
    if not exp:
        return "Invalid", None
    days = (exp - date.today()).days
    if days < 0:
        return "Expired", days
    if days <= 30:
        return "Due soon", days
    return "Current", days


def invoice_aging(row: Any) -> dict[str, Any]:
    data = dict(row)
    if row["paid_date"] or str(row["status"]).lower() == "paid":
        data["age_days"] = 0
        data["aging_label"] = "Paid"
        return data
    due = parse_date(row["due_date"])
    age = (date.today() - due).days if due else 0
    data["age_days"] = max(0, age)
    if age <= 0:
        data["aging_label"] = "Not due"
    elif age <= 30:
        data["aging_label"] = "1-30 days"
    elif age <= 60:
        data["aging_label"] = "31-60 days"
    elif age <= 90:
        data["aging_label"] = "61-90 days"
    else:
        data["aging_label"] = "90+ days"
    return data


@app.exception_handler(401)
async def unauthorized_handler(request: Request, exc: HTTPException):
    return redirect("/login")


@app.exception_handler(402)
async def subscription_required_handler(request: Request, exc: HTTPException):
    return redirect("/billing")


@app.get("/health")
def health() -> dict[str, str]:
    query_one("SELECT 1 AS healthy")
    billing_state = "live" if stripe_live_configured() else "prelaunch"
    return {
        "status": "ok",
        "version": VERSION,
        "database": "ok",
        "billing": billing_state,
    }


@app.get("/health/ready")
def health_ready() -> JSONResponse:
    """Expose the release gate for monitoring without revealing configuration values."""

    report = evaluate_release_readiness()
    payload = {
        "status": "ok" if report["ready"] else "blocked",
        "ready": bool(report["ready"]),
        "environment": report["environment"],
        "checks": report["checks"],
        "blockers": report["blockers"],
        "warnings": report["warnings"],
    }
    return JSONResponse(payload, status_code=200 if report["ready"] else 503, headers={"Cache-Control": "no-store"})


@app.get("/manifest.webmanifest")
def manifest() -> Response:
    return Response((BASE_DIR / "static" / "manifest.webmanifest").read_text(encoding="utf-8"), media_type="application/manifest+json")


@app.get("/service-worker.js")
def service_worker() -> Response:
    return Response((BASE_DIR / "static" / "service-worker.js").read_text(encoding="utf-8"), media_type="application/javascript")


@app.get("/robots.txt")
def robots() -> Response:
    blocked_paths = (
        "/billing",
        "/audits",
        "/compliance",
        "/dashboard",
        "/detention/",
        "/documents",
        "/drivers",
        "/financials",
        "/fuel",
        "/growth",
        "/health",
        "/health/ready",
        "/idle",
        "/loads",
        "/onboard/",
        "/onboarding",
        "/payments",
        "/quotes",
        "/rate-quotes",
        "/receivables",
        "/settings",
        "/stripe/",
        "/startup",
        "/vehicles",
    )
    body = "User-agent: *\n" + "".join(f"Disallow: {path}\n" for path in blocked_paths)
    body += f"\nSitemap: {CANONICAL_BASE_URL}/sitemap.xml\n"
    return Response(body, media_type="text/plain")


@app.get("/sitemap.xml")
def sitemap() -> Response:
    updated = "2026-07-21"
    paths = ["/", "/demo", *(f"/{slug}" for slug in SEO_PAGES)]
    urls = "".join(
        f"<url><loc>{CANONICAL_BASE_URL}{path}</loc><lastmod>{updated}</lastmod></url>"
        for path in paths
    )
    xml = f'<?xml version="1.0" encoding="UTF-8"?><urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">{urls}</urlset>'
    return Response(xml, media_type="application/xml")


@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    description = "Small fleet trucking software for owner-operators and carriers with 1–20 trucks. Manage dispatch, seven driver pay models, payments, and estimated profit per load."
    return render(
        request,
        "marketing.html",
        {
            "public_page": True,
            "plans": PLAN_LIMITS,
            **seo_context(
                "/",
                "Small Fleet Trucking Software | CarrierOS",
                description,
                page_type="WebPage",
            ),
        },
    )


@app.get("/demo", response_class=HTMLResponse)
def demo(request: Request):
    return render(
        request,
        "demo.html",
        {
            "public_page": True,
            "plans": PLAN_LIMITS,
            **seo_context(
                "/demo",
                "CarrierOS Live Demo | Small Fleet Trucking Software",
                "Try the CarrierOS live demo with fictional fleet data. Explore dispatch, seven supported driver pay structures, pricing, and estimated profit per load.",
            ),
        },
    )


@app.get("/small-fleet-trucking-software", response_class=HTMLResponse)
@app.get("/driver-settlement-software", response_class=HTMLResponse)
@app.get("/load-profitability-calculator", response_class=HTMLResponse)
def seo_landing_page(request: Request):
    seo_slug = request.url.path.strip("/")
    page = SEO_PAGES.get(seo_slug)
    if not page:
        raise HTTPException(status_code=404, detail="Not found")
    return render(
        request,
        "seo_page.html",
        {
            "public_page": True,
            "page": page,
            "seo_slug": seo_slug,
            **seo_context(
                f"/{seo_slug}",
                page["title"],
                page["description"],
                faqs=page["faqs"],
            ),
        },
    )


@app.get("/privacy", response_class=HTMLResponse)
def privacy(request: Request):
    return render(request, "privacy.html", {"public_page": True})


@app.get("/terms", response_class=HTMLResponse)
def terms(request: Request):
    return render(
        request,
        "terms.html",
        {"public_page": True, "terms_version": TERMS_VERSION},
    )


@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request):
    if current_user(request):
        return redirect("/dashboard")
    return render(request, "login.html", {"public_page": True})


@app.post("/login")
async def login(request: Request):
    if login_rate_limited(request):
        return render(request, "login.html", {"error": "Too many sign-in attempts. Try again in 15 minutes.", "public_page": True}, 429)
    form = await verified_form(request)
    email = str(form.get("email", "")).strip().lower()
    password = str(form.get("password", ""))
    row = query_one("SELECT * FROM users WHERE lower(email)=?", (email,))
    if not row or not verify_password(password, row["password_hash"]):
        login_attempts.setdefault(client_key(request), []).append(time.time())
        return render(request, "login.html", {
            "error": "Invalid email or password.",
            "public_page": True,
        }, 400)
    if password_needs_rehash(str(row["password_hash"])):
        execute(
            "UPDATE users SET password_hash=? WHERE id=?",
            (hash_password(password), row["id"]),
        )
    request.session.clear()
    request.session["user_id"] = row["id"]
    request.session["csrf_token"] = secrets.token_urlsafe(32)
    login_attempts.pop(client_key(request), None)
    record_audit_event(
        "user.login",
        organization_id=int(row["organization_id"]),
        user_id=int(row["id"]),
    )
    return redirect("/dashboard")


@app.get("/signup", response_class=HTMLResponse)
def signup_page(request: Request, plan: str = "owner_operator"):
    if current_user(request):
        return redirect("/dashboard")
    if not customer_signups_open():
        return render(request, "launch_pending.html", {"public_page": True})
    selected_plan = plan if plan in PLAN_LIMITS else "owner_operator"
    return render(request, "signup.html", {"plans": PLAN_LIMITS, "selected_plan": selected_plan, "public_page": True})


@app.post("/signup", response_class=HTMLResponse)
async def signup(request: Request):
    if not customer_signups_open():
        return render(request, "launch_pending.html", {"public_page": True}, 503)
    if signup_rate_limited(request):
        return render(request, "signup.html", {
            "error": "Too many accounts were created from this connection. Try again in one hour.",
            "plans": PLAN_LIMITS,
            "selected_plan": "owner_operator",
            "public_page": True,
        }, 429)
    form = await verified_form(request)
    full_name = str(form.get("full_name", "")).strip()
    company_name = str(form.get("company_name", "")).strip()
    email = str(form.get("email", "")).strip().lower()
    password = str(form.get("password", ""))
    accepted_terms = yes(form.get("accepted_terms"))
    plan_code = str(form.get("plan", "owner_operator"))
    if plan_code not in PLAN_LIMITS:
        plan_code = "owner_operator"
    plan = PLAN_LIMITS[plan_code]
    error = None
    if not full_name or not company_name or "@" not in email:
        error = "Enter your name, company, and a valid email address."
    elif not valid_password(password):
        error = "Use at least 12 characters with uppercase, lowercase, a number, and a symbol."
    elif not accepted_terms:
        error = "Accept the Terms of Service and Privacy Policy to create an account."
    elif query_one("SELECT id FROM users WHERE lower(email)=?", (email,)):
        error = "An account already exists for that email address."
    if error:
        return render(request, "signup.html", {
            "error": error,
            "plans": PLAN_LIMITS,
            "selected_plan": plan_code,
            "values": {"full_name": full_name, "company_name": company_name, "email": email},
            "public_page": True,
        }, 400)

    trial_ends_at = (date.today() + timedelta(days=TRIAL_DAYS)).isoformat() if BILLING_MODE == "beta" else None
    subscription_status = "trialing" if BILLING_MODE == "beta" else "incomplete"
    accepted_at = utc_now_iso()
    with db_session() as conn:
        org_id = conn.execute(
            """INSERT INTO organizations
            (name, owner_name, owner_email, plan_code, active_unit_limit,
             subscription_status, trial_ends_at, reporting_start_month, default_report_month,
             terms_accepted_at, terms_version)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                company_name, full_name, email, plan_code, int(plan["units"]),
                subscription_status, trial_ends_at,
                date.today().replace(day=1).isoformat(), date.today().replace(day=1).isoformat(),
                accepted_at, TERMS_VERSION,
            ),
        ).lastrowid
        user_id = conn.execute(
            """INSERT INTO users
            (organization_id, full_name, email, password_hash, is_admin, role)
            VALUES (?, ?, ?, ?, 1, 'owner')""",
            (org_id, full_name, email, hash_password(password)),
        ).lastrowid
        conn.execute(
            "INSERT INTO overhead_items (organization_id, name, monthly_cost, sort_order) VALUES (?, 'Other overhead', 0, 1)",
            (org_id,),
        )
    signup_attempts.setdefault(client_key(request), []).append(time.time())
    record_audit_event(
        "organization.created",
        organization_id=int(org_id),
        user_id=int(user_id),
        details={"billing_mode": BILLING_MODE, "plan_code": plan_code, "terms_version": TERMS_VERSION},
    )
    request.session.clear()
    request.session["user_id"] = user_id
    request.session["csrf_token"] = secrets.token_urlsafe(32)
    return redirect("/dashboard" if BILLING_MODE == "beta" else "/billing?new=1")


@app.get("/forgot-password", response_class=HTMLResponse)
def forgot_password_page(request: Request):
    if current_user(request):
        return redirect("/dashboard")
    return render(request, "forgot_password.html", {"public_page": True})


@app.post("/forgot-password", response_class=HTMLResponse)
async def forgot_password(request: Request):
    if reset_rate_limited(request):
        return render(
            request,
            "forgot_password.html",
            {
                "error": "Too many reset requests. Try again in one hour or contact support.",
                "public_page": True,
            },
            429,
        )
    form = await verified_form(request)
    email = str(form.get("email", "")).strip().lower()
    row = query_one("SELECT * FROM users WHERE lower(email)=?", (email,))
    reset_attempts.setdefault(client_key(request), []).append(time.time())
    if row and smtp_configured():
        raw_token = create_password_reset_token(int(row["id"]))
        try:
            await asyncio.to_thread(
                send_password_reset_email,
                recipient=str(row["email"]),
                full_name=str(row["full_name"]),
                reset_url=f"{public_url(request)}/reset-password?token={raw_token}",
            )
            record_audit_event(
                "password.reset_requested",
                organization_id=int(row["organization_id"]),
                user_id=int(row["id"]),
                details={"delivery": "accepted"},
            )
        except Exception:
            logger.exception("CarrierOS password-reset delivery failed")
            record_audit_event(
                "password.reset_requested",
                organization_id=int(row["organization_id"]),
                user_id=int(row["id"]),
                details={"delivery": "failed"},
            )
    return render(
        request,
        "forgot_password.html",
        {"sent": True, "public_page": True},
    )


@app.get("/reset-password", response_class=HTMLResponse)
def reset_password_page(request: Request, token: str = ""):
    return render(
        request,
        "reset_password.html",
        {"token": token, "public_page": True},
    )


@app.post("/reset-password", response_class=HTMLResponse)
async def reset_password(request: Request):
    if reset_rate_limited(request):
        return render(
            request,
            "reset_password.html",
            {
                "token": "",
                "error": "Too many reset attempts. Try again in one hour or contact support.",
                "public_page": True,
            },
            429,
        )
    form = await verified_form(request)
    reset_attempts.setdefault(client_key(request), []).append(time.time())
    token = str(form.get("token", ""))
    password = str(form.get("password", ""))
    confirm_password = str(form.get("confirm_password", ""))
    error = None
    if not token:
        error = "This reset link is invalid or has expired."
    elif password != confirm_password:
        error = "The passwords do not match."
    elif not valid_password(password):
        error = "Use at least 12 characters with uppercase, lowercase, a number, and a symbol."
    row = None if error else reset_password_with_token(token, hash_password(password))
    if not error and not row:
        error = "This reset link is invalid or has expired."
    if error:
        return render(
            request,
            "reset_password.html",
            {"token": token, "error": error, "public_page": True},
            400,
        )
    record_audit_event(
        "password.reset_completed",
        organization_id=int(row["organization_id"]),
        user_id=int(row["user_id"]),
    )
    request.session.clear()
    return render(
        request,
        "login.html",
        {"success": "Your password has been changed. Sign in with the new password.", "public_page": True},
    )


@app.get("/billing", response_class=HTMLResponse)
def billing(request: Request, checkout: str | None = None, new: int = 0):
    user = current_user(request)
    if not user:
        return redirect("/login")
    plan = PLAN_LIMITS.get(str(user.get("plan_code")), PLAN_LIMITS["owner_operator"])
    return render(request, "billing.html", {
        "plan": plan,
        "plans": PLAN_LIMITS,
        "stripe_ready": stripe_configured() and (not IS_PRODUCTION or stripe_live_configured()),
        "billing_mode": BILLING_MODE,
        "has_access": subscription_allows_access(user),
        "checkout_result": checkout,
        "new_account": bool(new),
    })


@app.post("/billing/checkout")
async def billing_checkout(request: Request):
    user = current_user(request)
    if not user:
        return redirect("/login")
    if BILLING_MODE == "beta":
        set_flash(request, "No payment is required during the founding beta.")
        return redirect("/billing")
    if IS_PRODUCTION and not stripe_live_configured():
        set_flash(
            request,
            "Live subscription billing is awaiting final activation. Please contact support for launch access.",
        )
        return redirect("/billing")
    form = await verified_form(request)
    plan_code = str(form.get("plan", user.get("plan_code") or "owner_operator"))
    if plan_code not in PLAN_LIMITS:
        raise HTTPException(status_code=400, detail="Unknown plan")
    if user.get("billing_subscription_reference") and str(user.get("subscription_status")) not in {
        "canceled", "incomplete_expired"
    }:
        set_flash(request, "Use the secure billing portal to change an existing subscription.")
        return redirect("/billing")
    try:
        base_url = public_url(request)
        session = await asyncio.to_thread(
            create_checkout_session,
            organization_id=int(user["organization_id"]),
            owner_email=str(user["owner_email"] or user["email"]),
            plan_code=plan_code,
            expected_monthly_price=int(PLAN_LIMITS[plan_code]["price"]),
            existing_customer_id=str(user["billing_customer_reference"] or "") or None,
            success_url=f"{base_url}/billing?checkout=success&session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{base_url}/billing?checkout=cancelled",
        )
    except BillingConfigurationError:
        set_flash(request, "Secure billing is not configured yet. Please contact support.")
        return redirect("/billing")
    except stripe.StripeError:
        set_flash(request, "Stripe could not start checkout. Please try again or contact support.")
        return redirect("/billing")
    checkout_url = str(object_value(session, "url") or "")
    if not checkout_url.startswith("https://"):
        set_flash(request, "Stripe did not return a secure checkout link. Please contact support.")
        return redirect("/billing")
    return redirect(checkout_url)


@app.post("/billing/portal")
async def billing_portal(request: Request):
    user = current_user(request)
    if not user:
        return redirect("/login")
    await verified_form(request)
    customer_id = str(user.get("billing_customer_reference") or "")
    if not customer_id:
        set_flash(request, "Start a subscription before opening the billing portal.")
        return redirect("/billing")
    try:
        session = await asyncio.to_thread(
            create_portal_session,
            customer_id=customer_id,
            return_url=f"{public_url(request)}/billing",
        )
    except (BillingConfigurationError, stripe.StripeError):
        set_flash(request, "Stripe could not open the billing portal. Please try again or contact support.")
        return redirect("/billing")
    portal_url = str(object_value(session, "url") or "")
    if not portal_url.startswith("https://"):
        set_flash(request, "Stripe did not return a secure billing link. Please contact support.")
        return redirect("/billing")
    return redirect(portal_url)


def _stripe_metadata(obj: Any) -> dict[str, Any]:
    metadata = object_value(obj, "metadata") or {}
    return dict(metadata) if hasattr(metadata, "items") else {}


def _organization_for_stripe_object(conn: Any, obj: Any) -> Any:
    metadata = _stripe_metadata(obj)
    raw_org_id = metadata.get("carrieros_org_id") or object_value(obj, "client_reference_id")
    if str(raw_org_id or "").isdigit():
        row = conn.execute("SELECT * FROM organizations WHERE id=?", (int(raw_org_id),)).fetchone()
        if row:
            return row
    subscription_id = object_value(obj, "subscription") or object_value(obj, "id")
    if isinstance(subscription_id, str) and subscription_id.startswith("sub_"):
        row = conn.execute(
            "SELECT * FROM organizations WHERE billing_subscription_reference=?", (subscription_id,)
        ).fetchone()
        if row:
            return row
    customer_id = object_value(obj, "customer")
    if isinstance(customer_id, str):
        return conn.execute(
            "SELECT * FROM organizations WHERE billing_customer_reference=?", (customer_id,)
        ).fetchone()
    return None


def _apply_checkout_completed(conn: Any, checkout_session: Any) -> None:
    if str(object_value(checkout_session, "mode") or "") != "subscription":
        return
    organization = _organization_for_stripe_object(conn, checkout_session)
    metadata = _stripe_metadata(checkout_session)
    plan_code = str(metadata.get("carrieros_plan_code") or "")
    if not organization or plan_code not in PLAN_LIMITS:
        return
    customer_id = object_value(checkout_session, "customer")
    subscription_id = object_value(checkout_session, "subscription")
    plan = PLAN_LIMITS[plan_code]
    conn.execute(
        """UPDATE organizations SET plan_code=?, active_unit_limit=?, subscription_status='trialing',
        trial_ends_at=?, billing_customer_reference=COALESCE(?, billing_customer_reference),
        billing_subscription_reference=COALESCE(?, billing_subscription_reference),
        billing_price_reference=? WHERE id=?""",
        (
            plan_code,
            int(plan["units"]),
            (date.today() + timedelta(days=14)).isoformat(),
            customer_id,
            subscription_id,
            price_id_for_plan(plan_code),
            organization["id"],
        ),
    )


def _subscription_period_end(subscription: Any) -> str | None:
    direct = unix_date(object_value(subscription, "current_period_end"))
    if direct:
        return direct
    items = object_value(subscription, "items") or {}
    data = object_value(items, "data") or []
    return unix_date(object_value(data[0], "current_period_end")) if data else None


def _apply_subscription_event(conn: Any, subscription: Any, event_type: str) -> None:
    organization = _organization_for_stripe_object(conn, subscription)
    if not organization:
        return
    price_id = first_subscription_price_id(subscription)
    metadata = _stripe_metadata(subscription)
    plan_code = plan_code_for_price(price_id) or str(metadata.get("carrieros_plan_code") or "")
    if plan_code not in PLAN_LIMITS:
        plan_code = str(organization["plan_code"])
    plan = PLAN_LIMITS.get(plan_code, PLAN_LIMITS["owner_operator"])
    status = "canceled" if event_type == "customer.subscription.deleted" else str(
        object_value(subscription, "status") or organization["subscription_status"]
    )
    conn.execute(
        """UPDATE organizations SET plan_code=?, active_unit_limit=?, subscription_status=?,
        trial_ends_at=?, billing_customer_reference=COALESCE(?, billing_customer_reference),
        billing_subscription_reference=COALESCE(?, billing_subscription_reference),
        billing_price_reference=COALESCE(?, billing_price_reference),
        subscription_current_period_end=?, subscription_cancel_at_period_end=? WHERE id=?""",
        (
            plan_code,
            int(plan["units"]),
            status,
            unix_date(object_value(subscription, "trial_end")),
            object_value(subscription, "customer"),
            object_value(subscription, "id"),
            price_id,
            _subscription_period_end(subscription),
            1 if object_value(subscription, "cancel_at_period_end") else 0,
            organization["id"],
        ),
    )


def _apply_invoice_event(conn: Any, invoice: Any, event_type: str) -> None:
    organization = _organization_for_stripe_object(conn, invoice)
    if not organization:
        return
    subscription_id = object_value(invoice, "subscription")
    if not subscription_id:
        parent = object_value(invoice, "parent") or {}
        details = object_value(parent, "subscription_details") or {}
        subscription_id = object_value(details, "subscription")
    if not subscription_id or subscription_id != organization["billing_subscription_reference"]:
        return
    status = "active" if event_type == "invoice.paid" else "past_due"
    conn.execute(
        "UPDATE organizations SET subscription_status=? WHERE id=?",
        (status, organization["id"]),
    )


@app.post("/stripe/webhook")
async def stripe_webhook(request: Request):
    payload = await request.body()
    try:
        event = construct_webhook_event(payload, request.headers.get("stripe-signature"))
    except BillingConfigurationError:
        return JSONResponse({"detail": "Webhook is not configured"}, status_code=503)
    except Exception:
        return JSONResponse({"detail": "Invalid webhook"}, status_code=400)

    event_id = str(object_value(event, "id") or "")
    event_type = str(object_value(event, "type") or "")
    data = object_value(event, "data") or {}
    stripe_object = object_value(data, "object") or {}
    if not event_id or not event_type:
        return JSONResponse({"detail": "Invalid event"}, status_code=400)

    with db_session() as conn:
        if conn.execute(
            "SELECT 1 FROM processed_stripe_events WHERE event_id=?", (event_id,)
        ).fetchone():
            return {"received": True, "duplicate": True}
        if event_type == "checkout.session.completed":
            _apply_checkout_completed(conn, stripe_object)
        elif event_type in {
            "customer.subscription.created",
            "customer.subscription.updated",
            "customer.subscription.deleted",
            "customer.subscription.trial_will_end",
        }:
            _apply_subscription_event(conn, stripe_object, event_type)
        elif event_type in {"invoice.paid", "invoice.payment_failed"}:
            _apply_invoice_event(conn, stripe_object, event_type)
        conn.execute(
            "INSERT INTO processed_stripe_events (event_id, event_type) VALUES (?, ?)",
            (event_id, event_type),
        )
    return {"received": True}


@app.get("/logout")
def logout(request: Request):
    user = current_user(request)
    if user:
        record_audit_event(
            "user.logout",
            organization_id=int(user["organization_id"]),
            user_id=int(user["id"]),
        )
    request.session.clear()
    return redirect("/login")


@app.get("/dashboard", response_class=HTMLResponse)
def dashboard(request: Request, month: str | None = None):
    user = require_user(request)
    settings = query_one("SELECT * FROM organizations WHERE id=?", (user["organization_id"],))
    report_month = selected_month(month, settings["default_report_month"] if settings else None)
    bundle, state = get_state(user["organization_id"], report_month)
    all_loads = loads_with_results(bundle, state)

    compliance_rows = query_all("SELECT * FROM compliance_items WHERE organization_id=? ORDER BY expiration_date", (user["organization_id"],))
    compliance = []
    due_count = 0
    for row in compliance_rows:
        item = dict(row)
        item["status"], item["days"] = compliance_status(row["expiration_date"])
        compliance.append(item)
        if item["status"] != "Current":
            due_count += 1
    invoice_rows = query_all("SELECT * FROM invoices WHERE organization_id=? ORDER BY due_date", (user["organization_id"],))
    invoices = [invoice_aging(row) for row in invoice_rows]
    unpaid_total = sum(number(row["amount"]) for row in invoice_rows if str(row["status"]).lower() != "paid")
    overdue_total = sum(number(row["amount"]) for row in invoices if str(row["status"]).lower() != "paid" and row["age_days"] > 0)
    quick_link_rows = query_all(
        """SELECT * FROM quick_links WHERE organization_id=?
        ORDER BY sort_order, lower(label) LIMIT 8""",
        (user["organization_id"],),
    )
    warnings = state["warnings"]
    open_exceptions = (
        warnings["loss_making"]
        + warnings["below_target"]
        + warnings["high_deadhead"]
        + warnings["low_company_profit"]
        + warnings["fixed_cost_reviews"]
        + warnings["idle_input_errors"]
        + due_count
    )

    return render(request, "dashboard.html", {
        "bundle": bundle,
        "state": state,
        "stats": state["summary"],
        "warnings": warnings,
        "balances": state["driver_balances"],
        "loads": all_loads[:10],
        "month": report_month,
        "monthly": state["selected_month_financial"],
        "bridges": state["selected_month_bridges"],
        "compliance_due": due_count,
        "unpaid_total": unpaid_total,
        "overdue_total": overdue_total,
        "compliance": compliance[:5],
        "invoices": invoices[:5],
        "quick_links": [as_dict(row) for row in quick_link_rows],
        "open_exceptions": open_exceptions,
        "driver_availability": driver_availability(int(user["organization_id"])),
    })


@app.get("/dispatch", response_class=HTMLResponse)
def dispatch_page(request: Request):
    """Provide a dispatch-first view of the next open window for each driver.

    Availability is intentionally derived from the latest non-cancelled,
    non-quote operational load assigned to each driver.  CarrierOS does not
    infer live GPS, HOS, or appointment completion from this view.
    """
    user = require_user(request)
    organization_id = int(user["organization_id"])
    availability = driver_availability(organization_id)
    operational_loads = query_all(
        """SELECT id, load_number, status, driver_id, delivery_date, destination
        FROM loads
        WHERE organization_id=?
          AND lower(trim(status)) NOT IN ('cancelled', 'canceled', 'quote')
        ORDER BY delivery_date, id""",
        (organization_id,),
    )
    quick_link_rows = query_all(
        """SELECT * FROM quick_links WHERE organization_id=?
        ORDER BY sort_order, lower(label) LIMIT 8""",
        (organization_id,),
    )
    return render(request, "dispatch.html", {
        "driver_availability": availability,
        "operational_loads": [as_dict(row) for row in operational_loads],
        "quick_links": [as_dict(row) for row in quick_link_rows],
        "dispatch_summary": {
            "drivers": len(availability),
            "scheduled": sum(1 for item in availability if item.get("load_id")),
            "open_now": sum(1 for item in availability if item.get("availability_tone") == "good"),
            "needs_contact": sum(1 for item in availability if not item.get("phone_href")),
        },
    })


def quick_links_context(
    organization_id: int,
    *,
    values: dict[str, str] | None = None,
    error: str | None = None,
) -> dict[str, Any]:
    rows = query_all(
        """SELECT * FROM quick_links WHERE organization_id=?
        ORDER BY sort_order, lower(label)""",
        (organization_id,),
    )
    return {
        "links": [as_dict(row) for row in rows],
        "quick_link_categories": QUICK_LINK_CATEGORIES,
        "quick_link_limit": QUICK_LINK_LIMIT,
        "values": values or {},
        "error": error,
    }


@app.get("/links", response_class=HTMLResponse)
def quick_links_page(request: Request):
    user = require_user(request)
    return render(
        request,
        "quick_links.html",
        quick_links_context(int(user["organization_id"])),
    )


@app.post("/links")
async def create_quick_link(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    organization_id = int(user["organization_id"])
    label = str(form.get("label") or "").strip()
    submitted_url = str(form.get("url") or "").strip()
    url = normalized_external_url(submitted_url)
    category = str(form.get("category") or "Other").strip()
    values = {"label": label, "url": submitted_url, "category": category}

    error = None
    if not 2 <= len(label) <= 60:
        error = "Enter a link name between 2 and 60 characters."
    elif not url:
        error = "Enter a valid http or https website address."
    elif category not in QUICK_LINK_CATEGORIES:
        error = "Choose a valid link category."
    elif query_one(
        "SELECT 1 FROM quick_links WHERE organization_id=? AND lower(label)=lower(?)",
        (organization_id, label),
    ):
        error = "A business link with that name already exists."
    elif int(
        query_one(
            "SELECT COUNT(*) AS total FROM quick_links WHERE organization_id=?",
            (organization_id,),
        )["total"]
    ) >= QUICK_LINK_LIMIT:
        error = f"Each workspace can save up to {QUICK_LINK_LIMIT} business links."

    if error:
        return render(
            request,
            "quick_links.html",
            quick_links_context(organization_id, values=values, error=error),
            status_code=400,
        )

    sort_row = query_one(
        "SELECT COALESCE(MAX(sort_order), 0) AS current_sort FROM quick_links WHERE organization_id=?",
        (organization_id,),
    )
    execute(
        """INSERT INTO quick_links
        (organization_id, label, url, category, sort_order)
        VALUES (?, ?, ?, ?, ?)""",
        (organization_id, label, url, category, int(sort_row["current_sort"]) + 10),
    )
    record_audit_event(
        "quick_link.created",
        organization_id=organization_id,
        user_id=int(user["id"]),
        details={"label": label, "category": category},
    )
    request.session["flash"] = "Business link added."
    return redirect("/links")


@app.post("/links/{link_id}/delete")
async def delete_quick_link(request: Request, link_id: int):
    user = require_user(request)
    await verified_form(request)
    organization_id = int(user["organization_id"])
    link = query_one(
        "SELECT * FROM quick_links WHERE id=? AND organization_id=?",
        (link_id, organization_id),
    )
    if not link:
        raise HTTPException(status_code=404, detail="Business link not found")
    execute(
        "DELETE FROM quick_links WHERE id=? AND organization_id=?",
        (link_id, organization_id),
    )
    record_audit_event(
        "quick_link.deleted",
        organization_id=organization_id,
        user_id=int(user["id"]),
        details={"label": str(link["label"]), "category": str(link["category"])},
    )
    request.session["flash"] = "Business link removed."
    return redirect("/links")


def _selected_query_ids(request: Request, name: str, allowed: set[int]) -> set[int]:
    return {
        value
        for raw in request.query_params.getlist(name)
        if (value := integer(raw)) in allowed
    }


def load_report_context(
    request: Request, bundle: dict[str, Any], state: dict[str, Any]
) -> dict[str, Any]:
    all_rows = loads_with_results(bundle, state)
    allowed_driver_ids = {int(driver["id"]) for driver in bundle["drivers"]}
    allowed_load_ids = {int(load["id"]) for load in bundle["loads"]}
    selected_driver_ids = _selected_query_ids(request, "driver_id", allowed_driver_ids)
    selected_load_ids = _selected_query_ids(request, "load_id", allowed_load_ids)
    date_from = parse_date(request.query_params.get("date_from"))
    date_to = parse_date(request.query_params.get("date_to"))
    date_swapped = bool(date_from and date_to and date_from > date_to)
    if date_swapped:
        date_from, date_to = date_to, date_from
    date_field = str(request.query_params.get("date_field") or "delivery_date")
    if date_field not in {"pickup_date", "delivery_date"}:
        date_field = "delivery_date"
    allowed_sorts = {value for value, _ in LOAD_SORT_OPTIONS}
    sort = str(request.query_params.get("sort") or "delivery_desc")
    if sort not in allowed_sorts:
        sort = "delivery_desc"
    statuses = sorted({str(load.get("status") or "") for load in bundle["loads"] if load.get("status")})
    status = str(request.query_params.get("status") or "")
    if status not in statuses:
        status = ""
    filters = {
        "date_from": date_from,
        "date_to": date_to,
        "date_field": date_field,
        "driver_ids": selected_driver_ids,
        "load_ids": selected_load_ids,
        "status": status,
        "sort": sort,
        "date_swapped": date_swapped,
    }
    rows = filter_and_sort_loads(
        all_rows,
        date_from=date_from,
        date_to=date_to,
        date_field=date_field,
        driver_ids=selected_driver_ids,
        load_ids=selected_load_ids,
        status=status,
        sort=sort,
    )
    query_values: list[tuple[str, str]] = []
    if date_from:
        query_values.append(("date_from", date_from.isoformat()))
    if date_to:
        query_values.append(("date_to", date_to.isoformat()))
    if date_field != "delivery_date":
        query_values.append(("date_field", date_field))
    query_values.extend(("driver_id", str(value)) for value in sorted(selected_driver_ids))
    query_values.extend(("load_id", str(value)) for value in sorted(selected_load_ids))
    if status:
        query_values.append(("status", status))
    if sort != "delivery_desc":
        query_values.append(("sort", sort))
    query_string = urlencode(query_values)
    return {
        "all_loads": all_rows,
        "loads": rows,
        "drivers": bundle["drivers"],
        "statuses": statuses,
        "filters": filters,
        "sort_options": LOAD_SORT_OPTIONS,
        "filter_query": query_string,
        "selection": summarize_load_rows(rows),
        "selected_by_driver": summarize_load_rows_by_driver(rows, bundle["drivers"]),
        "selected_by_month": summarize_load_rows_by_month(rows),
    }


@app.get("/loads", response_class=HTMLResponse)
def loads_page(request: Request):
    user = require_user(request)
    bundle, state = get_state(user["organization_id"])
    context = load_report_context(request, bundle, state)
    context.update({"filter_action": "/loads", "filter_page": "loads"})
    return render(request, "loads.html", context)


def csv_safe(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    if value.lstrip().startswith(("=", "+", "-", "@")):
        return f"'{value}"
    return value


@app.get("/loads/export.csv")
def export_filtered_loads(request: Request):
    user = require_user(request)
    bundle, state = get_state(user["organization_id"])
    context = load_report_context(request, bundle, state)
    output = io.StringIO(newline="")
    writer = csv.writer(output)
    writer.writerow([
        "Load number", "Pickup date", "Delivery date", "Status", "Driver", "Unit", "Broker",
        "Origin", "Destination", "Revenue", "Loaded miles", "Deadhead miles", "Total miles",
        "Revenue per mile", "Fuel cost", "Allocated fixed cost", "Maintenance reserve", "Company fees",
        "Tolls / misc", "Other direct costs", "Operating expense", "Driver / contractor pay",
        "Owner-operator load pay", "Company profit before owner distribution", "Owner profit distribution",
        "Retained company profit", "Company margin", "Decision", "Included", "Exclusion reason", "Notes",
    ])
    for load in context["loads"]:
        result = load["result"]
        writer.writerow([
            csv_safe(str(load.get("load_number") or "")), load.get("pickup_date") or "",
            load.get("delivery_date") or "", csv_safe(str(load.get("status") or "")),
            csv_safe(str((load.get("driver") or {}).get("name") or "Unassigned")),
            csv_safe(str((load.get("vehicle") or {}).get("name") or "")),
            csv_safe(str(load.get("broker") or "")), csv_safe(str(load.get("origin") or "")),
            csv_safe(str(load.get("destination") or "")), float(load.get("revenue") or 0),
            float(load.get("loaded_miles") or 0), float(load.get("deadhead_miles") or 0),
            result.total_miles, result.all_in_revenue_per_mile, result.fuel_cost,
            result.allocated_fixed_cost, result.maintenance_reserve, result.company_fees,
            result.tolls_misc, result.other_direct_costs, result.total_operating_expense,
            result.driver_contractor_earned, result.owner_operator_load_pay,
            result.company_profit_before_owner_distribution, result.owner_profit_distribution,
            result.retained_company_profit, result.company_margin_pct, result.decision,
            "Yes" if result.included else "No", csv_safe(result.exclusion_reason),
            csv_safe(str(load.get("notes") or "")),
        ])
    filename = f"carrieros-loads-{date.today().isoformat()}.csv"
    return Response(
        content="\ufeff" + output.getvalue(),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def load_form_context(org_id: int, values: dict[str, Any] | None = None, editing: bool = False):
    bundle = get_bundle(org_id)
    current_fuel = fuel_price_for(date.today(), bundle["weekly_fuel"], number(bundle["settings"]["fallback_diesel_price"]))
    return {
        "drivers": bundle["drivers"],
        "vehicles": bundle["vehicles"],
        "values": values or {
            "pickup_date": date.today().isoformat(),
            "delivery_date": date.today().isoformat(),
            "status": "Booked",
            "include_in_model": 1,
        },
        "current_fuel": current_fuel,
        "editing": editing,
    }


@app.get("/loads/new", response_class=HTMLResponse)
def new_load_page(request: Request):
    user = require_user(request)
    return render(request, "load_form.html", load_form_context(user["organization_id"]))


@app.post("/loads/new")
async def create_load(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    driver_id = integer(form.get("driver_id"))
    driver = query_one("SELECT * FROM drivers WHERE id=? AND organization_id=?", (driver_id, user["organization_id"]))
    if not driver:
        raise HTTPException(400, "Select a valid driver")
    vehicle_id = integer(form.get("vehicle_id")) or integer(driver["vehicle_id"])
    if vehicle_id and not query_one(
        "SELECT 1 FROM vehicles WHERE id=? AND organization_id=?",
        (vehicle_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid unit")
    ratecon = _ratecon_form_fields(form)
    status = str(form.get("status", "Booked"))
    if ratecon["ratecon_received_at"] and "Awaiting RateCon" in status:
        status = "RateCon Received"
    elif not ratecon["ratecon_received_at"] and status == "RateCon Received":
        status = "Booked — Awaiting RateCon"
    status_state = (
        LoadState.RATECON_REVIEW if ratecon["ratecon_received_at"] else normalize_state(status)
    )
    load_public_uuid = str(uuid.uuid4())
    created_at = utc_now_iso()
    revenue = optional_number(form.get("revenue"))
    load_id = execute(
        """INSERT INTO loads
        (organization_id,public_uuid,status_code,updated_at,load_number,pickup_date,delivery_date,driver_id,vehicle_id,
         broker, origin, destination, status, revenue, loaded_miles, deadhead_miles,
         fuel_override, tolls_misc, other_direct_costs, notes, include_in_model,
         pickup_address,pickup_window_start,pickup_window_end,pickup_contact_name,
         pickup_contact_phone,pickup_instructions,pickup_timezone,delivery_address,delivery_window_start,
         delivery_window_end,delivery_contact_name,delivery_contact_phone,
         delivery_instructions,delivery_timezone,ratecon_reference,ratecon_received_at)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            user["organization_id"], load_public_uuid, status_state.value, created_at,
            str(form.get("load_number", "")).strip() or f"LOAD-{datetime.now():%Y%m%d%H%M}",
            str(form.get("pickup_date", "")) or None, str(form.get("delivery_date", "")) or None,
            driver_id, vehicle_id or None, str(form.get("broker", "")).strip(), str(form.get("origin", "")).strip(),
            str(form.get("destination", "")).strip(), status, revenue,
            number(form.get("loaded_miles")), number(form.get("deadhead_miles")), optional_number(form.get("fuel_override")),
            number(form.get("tolls_misc")), number(form.get("other_direct_costs")), str(form.get("notes", "")).strip(),
            1 if yes(form.get("include_in_model")) else 0,
            ratecon["pickup_address"], ratecon["pickup_window_start"], ratecon["pickup_window_end"],
            ratecon["pickup_contact_name"], ratecon["pickup_contact_phone"], ratecon["pickup_instructions"], ratecon["pickup_timezone"],
            ratecon["delivery_address"], ratecon["delivery_window_start"], ratecon["delivery_window_end"],
            ratecon["delivery_contact_name"], ratecon["delivery_contact_phone"], ratecon["delivery_instructions"],
            ratecon["delivery_timezone"], ratecon["ratecon_reference"], ratecon["ratecon_received_at"],
        ),
    )
    with db_session() as conn:
        conn.execute(
            """INSERT INTO load_status_history
            (organization_id,load_id,prior_status,new_status,changed_by,idempotency_key,reason)
            VALUES (?,?,NULL,?,?,?,?)""",
            (
                user["organization_id"], load_id, status_state.value, user["id"],
                f"manual-load:{load_public_uuid}:initial", "Manual load created",
            ),
        )
        for sequence, stop_type, address, appointment_start, appointment_end, timezone_name, contact_name, contact_phone, instructions in (
            (1, "PICKUP", ratecon["pickup_address"], ratecon["pickup_window_start"], ratecon["pickup_window_end"], ratecon["pickup_timezone"], ratecon["pickup_contact_name"], ratecon["pickup_contact_phone"], ratecon["pickup_instructions"]),
            (2, "DELIVERY", ratecon["delivery_address"], ratecon["delivery_window_start"], ratecon["delivery_window_end"], ratecon["delivery_timezone"], ratecon["delivery_contact_name"], ratecon["delivery_contact_phone"], ratecon["delivery_instructions"]),
        ):
            conn.execute(
                """INSERT INTO load_stops
                (public_uuid,organization_id,load_id,sequence_number,stop_type,address_line1,
                 iana_timezone,appointment_local_start,appointment_local_end,contact_name,
                 contact_phone,instructions,detention_eligible)
                VALUES (?,?,?,?,?,?,?,?,?,?,?,?,1)""",
                (
                    str(uuid.uuid4()), user["organization_id"], load_id, sequence, stop_type,
                    address, timezone_name or None, appointment_start, appointment_end,
                    contact_name, contact_phone, instructions,
                ),
            )
        if revenue is not None:
            conn.execute(
                """INSERT INTO load_revenue_items
                (public_uuid,organization_id,load_id,category,description,amount_cents,stage,source)
                VALUES (?,?,?,'LINEHAUL','Manual load revenue',?,'BOOKED','manual_load')""",
                (str(uuid.uuid4()), user["organization_id"], load_id, money_to_cents(revenue, field="Load revenue")),
            )
        power_unit = conn.execute(
            "SELECT id FROM power_units WHERE organization_id=? AND legacy_vehicle_id=?",
            (user["organization_id"], vehicle_id),
        ).fetchone()
        if power_unit:
            conn.execute(
                """INSERT INTO load_assignments
                (public_uuid,organization_id,load_id,driver_id,power_unit_id,assignment_stage,
                 provisional,deadhead_miles_micros,route_source)
                VALUES (?,?,?, ?,?,'MANUAL_ENTRY',1,?,'manual')""",
                (
                    str(uuid.uuid4()), user["organization_id"], load_id, driver_id, power_unit["id"],
                    round(number(form.get("deadhead_miles")) * 1_000_000),
                ),
            )
    record_audit_event(
        "load.created",
        organization_id=int(user["organization_id"]),
        user_id=int(user["id"]),
        details={"load_id": load_id, "load_number": str(form.get("load_number", "")).strip()},
    )
    set_flash(request, "Load saved and the full model was recalculated.")
    return redirect(f"/loads/{load_id}")


@app.get("/loads/{load_id}/edit", response_class=HTMLResponse)
def edit_load_page(request: Request, load_id: int):
    user = require_user(request)
    row = query_one("SELECT * FROM loads WHERE id=? AND organization_id=?", (load_id, user["organization_id"]))
    if not row:
        raise HTTPException(404, "Load not found")
    return render(request, "load_form.html", load_form_context(user["organization_id"], dict(row), True))


@app.post("/loads/{load_id}/edit")
async def update_load(request: Request, load_id: int):
    user = require_user(request)
    form = await verified_form(request)
    existing = query_one(
        "SELECT * FROM loads WHERE id=? AND organization_id=?",
        (load_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Load not found")
    driver_id = integer(form.get("driver_id"))
    driver = query_one("SELECT * FROM drivers WHERE id=? AND organization_id=?", (driver_id, user["organization_id"]))
    if not driver:
        raise HTTPException(400, "Select a valid driver")
    vehicle_id = integer(form.get("vehicle_id")) or integer(driver["vehicle_id"])
    if vehicle_id and not query_one(
        "SELECT 1 FROM vehicles WHERE id=? AND organization_id=?",
        (vehicle_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid unit")
    ratecon = _ratecon_form_fields(form, existing)
    status = str(form.get("status", "Booked"))
    if ratecon["ratecon_received_at"] and "Awaiting RateCon" in status:
        status = "RateCon Received"
    elif not ratecon["ratecon_received_at"] and status == "RateCon Received":
        status = "Booked — Awaiting RateCon"
    with db_session() as conn:
        conn.execute(
            """UPDATE loads SET load_number=?,pickup_date=?,delivery_date=?,driver_id=?,vehicle_id=?,
            broker=?,origin=?,destination=?,status=?,revenue=?,loaded_miles=?,deadhead_miles=?,
            fuel_override=?,tolls_misc=?,other_direct_costs=?,notes=?,include_in_model=?,
            pickup_address=?,pickup_window_start=?,pickup_window_end=?,pickup_contact_name=?,
            pickup_contact_phone=?,pickup_instructions=?,pickup_timezone=?,delivery_address=?,delivery_window_start=?,
            delivery_window_end=?,delivery_contact_name=?,delivery_contact_phone=?,
            delivery_instructions=?,delivery_timezone=?,ratecon_reference=?,ratecon_received_at=?
            WHERE id=? AND organization_id=?""",
            (
                str(form.get("load_number", "")).strip() or f"LOAD-{load_id}", str(form.get("pickup_date", "")) or None,
                str(form.get("delivery_date", "")) or None, driver_id, vehicle_id or None, str(form.get("broker", "")).strip(),
                str(form.get("origin", "")).strip(), str(form.get("destination", "")).strip(), status,
                optional_number(form.get("revenue")), number(form.get("loaded_miles")), number(form.get("deadhead_miles")),
                optional_number(form.get("fuel_override")), number(form.get("tolls_misc")), number(form.get("other_direct_costs")),
                str(form.get("notes", "")).strip(), 1 if yes(form.get("include_in_model")) else 0,
                ratecon["pickup_address"], ratecon["pickup_window_start"], ratecon["pickup_window_end"],
                ratecon["pickup_contact_name"], ratecon["pickup_contact_phone"], ratecon["pickup_instructions"], ratecon["pickup_timezone"],
                ratecon["delivery_address"], ratecon["delivery_window_start"], ratecon["delivery_window_end"],
                ratecon["delivery_contact_name"], ratecon["delivery_contact_phone"], ratecon["delivery_instructions"],
                ratecon["delivery_timezone"], ratecon["ratecon_reference"], ratecon["ratecon_received_at"],
                load_id, user["organization_id"],
            ),
        )
        for stop_type, prefix in (("PICKUP", "pickup"), ("DELIVERY", "delivery")):
            conn.execute(
                """UPDATE load_stops SET address_line1=?,iana_timezone=?,appointment_local_start=?,
                appointment_local_end=?,contact_name=?,contact_phone=?,instructions=?,updated_at=?
                WHERE organization_id=? AND load_id=? AND stop_type=?""",
                (
                    ratecon[f"{prefix}_address"], ratecon[f"{prefix}_timezone"],
                    ratecon[f"{prefix}_window_start"], ratecon[f"{prefix}_window_end"],
                    ratecon[f"{prefix}_contact_name"], ratecon[f"{prefix}_contact_phone"],
                    ratecon[f"{prefix}_instructions"], utc_now_iso(),
                    user["organization_id"], load_id, stop_type,
                ),
            )
        revised_revenue = optional_number(form.get("revenue"))
        if revised_revenue is not None:
            conn.execute(
                """UPDATE load_revenue_items SET amount_cents=?
                WHERE organization_id=? AND load_id=? AND source='manual_load' AND stage='BOOKED'""",
                (money_to_cents(revised_revenue, field="Load revenue"), user["organization_id"], load_id),
            )
    record_audit_event(
        "load.updated",
        organization_id=int(user["organization_id"]),
        user_id=int(user["id"]),
        details={"load_id": load_id, "load_number": str(existing["load_number"])},
    )
    set_flash(request, "Load updated. Overlapping truck-day costs were recalculated across every included load.")
    return redirect(f"/loads/{load_id}")


@app.post("/loads/{load_id}/cancel")
async def cancel_load(request: Request, load_id: int):
    user = require_user(request)
    await verified_form(request)
    existing = query_one(
        "SELECT * FROM loads WHERE id=? AND organization_id=?",
        (load_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Load not found")
    execute(
        "UPDATE loads SET status='Cancelled',include_in_model=0 WHERE id=? AND organization_id=?",
        (load_id, user["organization_id"]),
    )
    record_audit_event(
        "load.cancelled",
        organization_id=int(user["organization_id"]),
        user_id=int(user["id"]),
        details={"load_id": load_id, "load_number": str(existing["load_number"])},
    )
    set_flash(request, "Load cancelled and removed from pay and profitability calculations. Its history was retained.")
    return redirect(f"/loads/{load_id}")


@app.get("/loads/{load_id}", response_class=HTMLResponse)
def load_detail(request: Request, load_id: int):
    user = require_user(request)
    bundle, state = get_state(user["organization_id"])
    item = next((row for row in loads_with_results(bundle, state) if int(row["id"]) == load_id), None)
    if not item:
        raise HTTPException(404, "Load not found")
    payments = query_all(
        """SELECT p.*,d.name AS driver_name FROM payments p LEFT JOIN drivers d ON d.id=p.driver_id
        WHERE p.organization_id=? AND p.load_id=? ORDER BY p.paid_at DESC,p.id DESC""",
        (user["organization_id"], load_id),
    )
    delivery_documents = _delivery_documents_for_load(int(user["organization_id"]), load_id)
    status_history = [dict(row) for row in query_all(
        """SELECT h.*,u.full_name AS changed_by_name FROM load_status_history h
        LEFT JOIN users u ON u.id=h.changed_by AND u.organization_id=h.organization_id
        WHERE h.organization_id=? AND h.load_id=? ORDER BY h.id DESC""",
        (user["organization_id"], load_id),
    )]
    return render(request, "load_detail.html", {
        "load": item,
        "payments": [dict(p) for p in payments],
        "delivery_documents": delivery_documents,
        "status_history": status_history,
        "delivery_status_options": [
            LoadState.AT_PICKUP.value,
            LoadState.IN_TRANSIT.value,
            LoadState.AT_DELIVERY.value,
            LoadState.DELIVERED_DOCUMENTS_PENDING.value,
        ],
        "dispatch": driver_dispatch_package(item),
    })


def _ratecon_page_context(organization_id: int, document: dict[str, Any]) -> dict[str, Any]:
    extraction = query_one(
        """SELECT * FROM ratecon_extractions
        WHERE organization_id=? AND document_id=? ORDER BY id DESC LIMIT 1""",
        (organization_id, document["id"]),
    )
    fields = []
    if extraction:
        fields = [dict(row) for row in query_all(
            """SELECT * FROM ratecon_extracted_fields
            WHERE organization_id=? AND extraction_id=? ORDER BY field_name""",
            (organization_id, extraction["id"]),
        )]
    candidates = [dict(row) for row in query_all(
        """SELECT c.*,l.public_uuid AS load_public_uuid,l.load_number,l.broker,l.origin,l.destination
        FROM ratecon_match_candidates c JOIN loads l ON l.id=c.load_id
        WHERE c.organization_id=? AND c.document_id=? ORDER BY c.score DESC,l.load_number""",
        (organization_id, document["id"]),
    )]
    differences = [dict(row) for row in query_all(
        """SELECT * FROM ratecon_differences
        WHERE organization_id=? AND document_id=? ORDER BY id""",
        (organization_id, document["id"]),
    )]
    material_pending = any(
        row["classification"] in MATERIAL_CLASSIFICATIONS and row["approval_status"] == "PENDING"
        for row in differences
    )
    delivery_documents = (
        _delivery_documents_for_load(organization_id, int(document["load_id"]))
        if document.get("load_id")
        else []
    )
    return {
        "document": document,
        "extraction": dict(extraction) if extraction else None,
        "fields": fields,
        "field_values": {
            row["field_name"]: (
                row["reviewed_value"] if row["reviewed_value"] is not None else row["extracted_value"]
            )
            for row in fields
            if row["human_review_status"] != "REJECTED"
        },
        "candidates": candidates,
        "differences": differences,
        "material_pending": material_pending,
        "max_ratecon_mb": 12,
        "proof_of_delivery_documents": [
            row for row in delivery_documents if row.get("document_kind") == "POD"
        ],
    }


@app.get("/ratecons", response_class=HTMLResponse)
def ratecon_inbox(request: Request):
    user = require_permission(request, "documents.view_operational")
    documents = [dict(row) for row in query_all(
        """SELECT d.*,l.public_uuid AS load_public_uuid,l.load_number
        FROM operational_documents d
        LEFT JOIN loads l ON l.id=d.load_id AND l.organization_id=d.organization_id
        WHERE d.organization_id=? AND d.document_type='RATECON' AND d.deleted_at IS NULL
        ORDER BY d.created_at DESC,d.id DESC""",
        (user["organization_id"],),
    )]
    loads = [dict(row) for row in query_all(
        """SELECT id,public_uuid,load_number,broker,origin,destination,status,status_code
        FROM loads WHERE organization_id=? AND include_in_model=1
        AND status_code IN ('BOOKED_AWAITING_RATECON','RATECON_REVIEW')
        ORDER BY pickup_date,id""",
        (user["organization_id"],),
    )]
    storage = configured_storage_provider()
    scanner = configured_malware_scan_provider()
    return render(request, "ratecons.html", {
        "documents": documents,
        "loads": loads,
        "storage_ready": (not IS_PRODUCTION) or storage.secure_at_rest,
        "scanner_ready": scanner.name != "manual",
        "max_ratecon_mb": 12,
    })


@app.post("/ratecons/upload")
async def upload_ratecon(request: Request):
    user = require_permission(request, "documents.manage_operational")
    form = await verified_form(request)
    upload = form.get("ratecon")
    if upload is None or not hasattr(upload, "read"):
        raise HTTPException(400, "Choose a RateCon PDF or phone image")
    payload = await upload.read(12 * 1024 * 1024 + 1)
    filename = Path(str(getattr(upload, "filename", "ratecon"))).name.replace("\x00", "")[:255]
    try:
        validated = validate_ratecon_upload(
            payload,
            filename=filename,
            claimed_content_type=str(getattr(upload, "content_type", "")),
        )
    except RateConError as exc:
        raise HTTPException(400, str(exc)) from exc
    storage = configured_storage_provider()
    if IS_PRODUCTION and not storage.secure_at_rest:
        raise HTTPException(503, "Private encrypted document storage is not configured")
    scan = configured_malware_scan_provider().scan(payload, filename=filename)
    if scan.status == "REJECTED":
        raise HTTPException(400, "The file failed malware screening and was not stored")
    document_uuid = str(uuid.uuid4())
    storage_key = f"organizations/{user['organization_id']}/ratecons/{document_uuid}.{validated.extension}"
    storage.put(storage_key, payload, content_type=validated.media_type)
    extraction_result = None
    if scan.status == "CLEAN":
        pages = document_text_pages(payload, validated.media_type, configured_ocr_provider())
        extraction_result = configured_extraction_provider().extract(pages)
    selected_load_uuid = str(form.get("load_public_uuid") or "").strip()
    load = _load_by_public_uuid(user["organization_id"], selected_load_uuid) if selected_load_uuid else None
    try:
        with db_session() as conn:
            document_id = int(conn.execute(
                """INSERT INTO operational_documents
                (public_uuid,organization_id,load_id,document_type,storage_key,storage_provider,
                 original_filename,content_type,size_bytes,page_count,sha256,malware_status,
                 processing_status,retention_date,created_by)
                VALUES (?,?,?,'RATECON',?,?,?,?,?,?,?,?,?,?,?)""",
                (
                    document_uuid, user["organization_id"], load["id"] if load else None,
                    storage_key, storage.name, filename, validated.media_type,
                    validated.size_bytes, validated.page_count, validated.sha256, scan.status,
                    extraction_result.status if extraction_result else "MALWARE_REVIEW",
                    default_retention_date(), user["id"],
                ),
            ).lastrowid)
            extraction_id = int(conn.execute(
                """INSERT INTO ratecon_extractions
                (public_uuid,organization_id,document_id,extraction_provider,provider_version,
                 status,detail,extracted_at)
                VALUES (?,?,?,?,?,?,?,?)""",
                (
                    str(uuid.uuid4()), user["organization_id"], document_id,
                    extraction_result.provider if extraction_result else "manual_fallback",
                    extraction_result.provider_version if extraction_result else "1",
                    extraction_result.status if extraction_result else "MALWARE_REVIEW",
                    extraction_result.detail if extraction_result else scan.detail,
                    utc_now_iso(),
                ),
            ).lastrowid)
            for field in extraction_result.fields if extraction_result else ():
                conn.execute(
                    """INSERT INTO ratecon_extracted_fields
                    (organization_id,extraction_id,field_name,extracted_value,confidence_millis,
                     document_page,evidence_text,bounding_reference)
                    VALUES (?,?,?,?,?,?,?,?)""",
                    (
                        user["organization_id"], extraction_id, field.name, field.value,
                        round(field.confidence * 1000), field.page, field.evidence,
                        field.bounding_reference,
                    ),
                )
            fields = list(extraction_result.fields) if extraction_result else []
            candidate_loads = [dict(row) for row in conn.execute(
                """SELECT * FROM loads WHERE organization_id=? AND include_in_model=1
                AND status_code IN ('BOOKED_AWAITING_RATECON','RATECON_REVIEW')""",
                (user["organization_id"],),
            ).fetchall()]
            for candidate in suggest_ratecon_matches(candidate_loads, {item.name: item.value for item in fields}):
                conn.execute(
                    """INSERT INTO ratecon_match_candidates
                    (organization_id,document_id,load_id,score,reasons_json)
                    VALUES (?,?,?,?,?)""",
                    (user["organization_id"], document_id, candidate.load_id, candidate.score, json.dumps(candidate.reasons)),
                )
            if load and scan.status == "CLEAN":
                _transition_workflow(
                    conn,
                    organization_id=user["organization_id"], load_id=load["id"],
                    target=LoadState.RATECON_REVIEW, actor_user_id=user["id"],
                    idempotency_key=f"ratecon:{document_uuid}:review",
                    reason="RateCon uploaded and malware screening passed",
                )
                conn.execute(
                    """UPDATE loads SET ratecon_reference=?,ratecon_received_at=?
                    WHERE id=? AND organization_id=?""",
                    (filename, utc_now_iso(), load["id"], user["organization_id"]),
                )
                _replace_ratecon_differences(
                    conn, organization_id=user["organization_id"], document_id=document_id,
                    load=load, fields=fields,
                )
    except sqlite3.IntegrityError as exc:
        storage.delete(storage_key)
        if "sha256" in str(exc).lower() or "unique" in str(exc).lower():
            raise HTTPException(409, "This RateCon was already uploaded for the selected load") from exc
        raise
    record_audit_event(
        "ratecon.uploaded", int(user["organization_id"]), int(user["id"]),
        {"document_uuid": document_uuid, "load_id": load["id"] if load else None, "sha256": validated.sha256},
    )
    set_flash(request, "RateCon stored privately. Review the extracted facts and every difference before dispatch.")
    return redirect(f"/ratecons/{document_uuid}")


@app.get("/ratecons/{document_uuid}", response_class=HTMLResponse)
def ratecon_detail_page(request: Request, document_uuid: str):
    user = require_permission(request, "documents.view_operational")
    document = _ratecon_document(user["organization_id"], document_uuid)
    return render(request, "ratecon_detail.html", _ratecon_page_context(user["organization_id"], document))


@app.post("/ratecons/{document_uuid}/pod", response_class=HTMLResponse)
async def upload_ratecon_pod(request: Request, document_uuid: str):
    """Store an office-uploaded POD alongside the load's RateCon."""
    user = require_permission(request, "documents.manage_operational")
    form = await verified_form(request)
    document = _ratecon_document(user["organization_id"], document_uuid)
    if document.get("document_type") != "RATECON":
        raise HTTPException(404, "RateCon not found")
    if not document.get("load_id"):
        raise HTTPException(409, "Attach the RateCon to a booked load before uploading a POD")
    upload = form.get("pod")
    if upload is None or not hasattr(upload, "read"):
        raise HTTPException(400, "Choose a signed proof of delivery file")
    payload = await upload.read(MAX_RATECON_BYTES + 1)
    filename = Path(str(getattr(upload, "filename", "proof-of-delivery"))).name.replace("\x00", "")[:255]
    try:
        validated = validate_ratecon_upload(
            payload,
            filename=filename,
            claimed_content_type=str(getattr(upload, "content_type", "")),
        )
    except RateConError as exc:
        raise HTTPException(400, str(exc)) from exc
    storage = configured_storage_provider()
    if not storage.secure_at_rest:
        raise HTTPException(503, "Private encrypted document storage is not configured")
    scan = configured_malware_scan_provider().scan(payload, filename=filename)
    if scan.status != "CLEAN":
        raise HTTPException(
            400,
            f"The proof of delivery cannot be accepted until malware screening is CLEAN ({scan.status})",
        )
    pod_uuid = str(uuid.uuid4())
    storage_key = (
        f"organizations/{user['organization_id']}/loads/{document['load_id']}"
        f"/delivery/{pod_uuid}.{validated.extension}"
    )
    storage.put(storage_key, payload, content_type=validated.media_type)
    try:
        with db_session() as conn:
            document_id = int(
                conn.execute(
                    """INSERT INTO operational_documents
                    (public_uuid,organization_id,load_id,document_type,storage_key,storage_provider,
                     original_filename,content_type,size_bytes,page_count,sha256,malware_status,
                     processing_status,retention_date,created_by)
                    VALUES (?,?,?,'POD',?,?,?,?,?,?,?,?,?,?,?)""",
                    (
                        pod_uuid,
                        user["organization_id"],
                        document["load_id"],
                        storage_key,
                        storage.name,
                        filename,
                        validated.media_type,
                        validated.size_bytes,
                        validated.page_count,
                        validated.sha256,
                        scan.status,
                        "RECEIVED",
                        default_retention_date(),
                        user["id"],
                    ),
                ).lastrowid
            )
            conn.execute(
                """INSERT INTO delivery_document_links
                (public_uuid,organization_id,document_id,load_id,document_kind,source,captured_by_user,notes)
                VALUES (?,?,?,?,?,'office',?,?)""",
                (
                    str(uuid.uuid4()),
                    user["organization_id"],
                    document_id,
                    document["load_id"],
                    "POD",
                    user["id"],
                    str(form.get("notes") or "").strip()[:500],
                ),
            )
    except sqlite3.IntegrityError as exc:
        storage.delete(storage_key)
        raise HTTPException(409, "This proof of delivery was already uploaded for the load") from exc
    record_audit_event(
        "delivery_document.uploaded",
        int(user["organization_id"]),
        int(user["id"]),
        {
            "load_id": document["load_id"],
            "document_uuid": pod_uuid,
            "document_kind": "POD",
            "source": "office",
            "sha256": validated.sha256,
        },
    )
    set_flash(request, "Proof of delivery stored privately and attached to the load.")
    return redirect(f"/ratecons/{document['public_uuid']}")


@app.post("/ratecons/{document_uuid}/attach/{load_uuid}")
async def attach_ratecon(request: Request, document_uuid: str, load_uuid: str):
    user = require_permission(request, "documents.manage_operational")
    await verified_form(request)
    document = _ratecon_document(user["organization_id"], document_uuid)
    if document["load_id"]:
        raise HTTPException(409, "This RateCon is already attached to a load")
    if document["malware_status"] != "CLEAN":
        raise HTTPException(409, "Malware screening must pass before attaching the RateCon")
    load = _load_by_public_uuid(user["organization_id"], load_uuid)
    extraction = query_one(
        "SELECT id FROM ratecon_extractions WHERE organization_id=? AND document_id=? ORDER BY id DESC LIMIT 1",
        (user["organization_id"], document["id"]),
    )
    fields = _effective_extracted_fields(user["organization_id"], int(extraction["id"])) if extraction else []
    with db_session() as conn:
        conn.execute(
            "UPDATE operational_documents SET load_id=? WHERE id=? AND organization_id=?",
            (load["id"], document["id"], user["organization_id"]),
        )
        conn.execute(
            """UPDATE ratecon_match_candidates SET selected_at=?,selected_by=?
            WHERE organization_id=? AND document_id=? AND load_id=?""",
            (utc_now_iso(), user["id"], user["organization_id"], document["id"], load["id"]),
        )
        _transition_workflow(
            conn, organization_id=user["organization_id"], load_id=load["id"],
            target=LoadState.RATECON_REVIEW, actor_user_id=user["id"],
            idempotency_key=f"ratecon:{document_uuid}:attach",
            reason="User selected the RateCon match",
        )
        conn.execute(
            "UPDATE loads SET ratecon_reference=?,ratecon_received_at=? WHERE id=? AND organization_id=?",
            (document["original_filename"], utc_now_iso(), load["id"], user["organization_id"]),
        )
        _replace_ratecon_differences(
            conn, organization_id=user["organization_id"], document_id=document["id"], load=load, fields=fields,
        )
    return redirect(f"/ratecons/{document_uuid}")


RATECON_REVIEW_FIELDS = (
    "broker_customer", "load_number", "ratecon_number", "total_rate",
    "pickup_date", "delivery_date", "added_stop", "tracking_penalty",
    "driver_assist", "factoring_restriction", "pickup_address",
    "pickup_window_start", "pickup_window_end", "pickup_timezone",
    "pickup_contact_name", "pickup_contact_phone", "pickup_instructions",
    "delivery_address", "delivery_window_start", "delivery_window_end",
    "delivery_timezone", "delivery_contact_name", "delivery_contact_phone",
    "delivery_instructions",
)


@app.post("/ratecons/{document_uuid}/review")
async def review_ratecon_fields(request: Request, document_uuid: str):
    user = require_permission(request, "documents.manage_operational")
    form = await verified_form(request)
    document = _ratecon_document(user["organization_id"], document_uuid)
    if not document["load_id"]:
        raise HTTPException(409, "Attach the RateCon to a load before reviewing its differences")
    extraction = query_one(
        "SELECT id FROM ratecon_extractions WHERE organization_id=? AND document_id=? ORDER BY id DESC LIMIT 1",
        (user["organization_id"], document["id"]),
    )
    if not extraction:
        raise HTTPException(409, "RateCon extraction record not found")
    with db_session() as conn:
        for name in RATECON_REVIEW_FIELDS:
            value = str(form.get(name) or "").strip()
            if not value:
                continue
            existing = conn.execute(
                "SELECT id FROM ratecon_extracted_fields WHERE extraction_id=? AND field_name=?",
                (extraction["id"], name),
            ).fetchone()
            if existing:
                conn.execute(
                    """UPDATE ratecon_extracted_fields SET reviewed_value=?,human_review_status='CORRECTED',
                    reviewed_by=?,reviewed_at=? WHERE id=? AND organization_id=?""",
                    (value, user["id"], utc_now_iso(), existing["id"], user["organization_id"]),
                )
            else:
                conn.execute(
                    """INSERT INTO ratecon_extracted_fields
                    (organization_id,extraction_id,field_name,extracted_value,confidence_millis,
                     evidence_text,human_review_status,reviewed_value,reviewed_by,reviewed_at)
                    VALUES (?,?,?,?,1000,'Human-entered review value','APPROVED',?,?,?)""",
                    (user["organization_id"], extraction["id"], name, value, value, user["id"], utc_now_iso()),
                )
        fields = _effective_extracted_fields(
            user["organization_id"], int(extraction["id"]), conn=conn
        )
        load = dict(conn.execute(
            "SELECT * FROM loads WHERE id=? AND organization_id=?",
            (document["load_id"], user["organization_id"]),
        ).fetchone())
        _replace_ratecon_differences(
            conn, organization_id=user["organization_id"], document_id=document["id"], load=load, fields=fields,
        )
    record_audit_event(
        "ratecon.fields_reviewed", int(user["organization_id"]), int(user["id"]),
        {"document_uuid": document_uuid},
    )
    set_flash(request, "Reviewed values saved. CarrierOS recalculated the comparison without overwriting the booking snapshot.")
    return redirect(f"/ratecons/{document_uuid}")


@app.post("/ratecons/{document_uuid}/approve")
async def approve_ratecon_differences(request: Request, document_uuid: str):
    user = require_permission(request, "loads.manage")
    form = await verified_form(request)
    document = _ratecon_document(user["organization_id"], document_uuid)
    if not document["load_id"] or document["malware_status"] != "CLEAN":
        raise HTTPException(409, "A clean, attached RateCon is required")
    extraction = query_one(
        "SELECT id FROM ratecon_extractions WHERE organization_id=? AND document_id=? ORDER BY id DESC LIMIT 1",
        (user["organization_id"], document["id"]),
    )
    fields = _effective_extracted_fields(user["organization_id"], int(extraction["id"])) if extraction else []
    values = {field.name: field.value for field in fields}
    required = (
        "total_rate", "pickup_address", "pickup_window_start", "pickup_window_end",
        "pickup_timezone", "delivery_address", "delivery_window_start",
        "delivery_window_end", "delivery_timezone",
    )
    missing = [name.replace("_", " ") for name in required if not values.get(name)]
    if missing:
        raise HTTPException(400, "Review and confirm: " + ", ".join(missing))
    for timezone_name in (values["pickup_timezone"], values["delivery_timezone"]):
        if not re.fullmatch(r"[A-Za-z_+-]+(?:/[A-Za-z_+-]+)+", timezone_name):
            raise HTTPException(400, "Use IANA facility time zones such as America/Chicago")
    differences = [dict(row) for row in query_all(
        "SELECT * FROM ratecon_differences WHERE organization_id=? AND document_id=?",
        (user["organization_id"], document["id"]),
    )]
    material = [row for row in differences if row["classification"] in MATERIAL_CLASSIFICATIONS]
    if material and not yes(form.get("approve_material")):
        raise HTTPException(400, "Explicitly approve the material RateCon differences")
    with db_session() as conn:
        now = utc_now_iso()
        conn.execute(
            """UPDATE ratecon_differences SET approval_status='APPROVED',approved_by=?,approved_at=?
            WHERE organization_id=? AND document_id=? AND approval_status='PENDING'""",
            (user["id"], now, user["organization_id"], document["id"]),
        )
        load = dict(conn.execute(
            "SELECT * FROM loads WHERE id=? AND organization_id=?",
            (document["load_id"], user["organization_id"]),
        ).fetchone())
        conn.execute(
            """UPDATE loads SET pickup_address=?,pickup_window_start=?,pickup_window_end=?,
            pickup_timezone=?,pickup_contact_name=?,pickup_contact_phone=?,pickup_instructions=?,
            delivery_address=?,delivery_window_start=?,delivery_window_end=?,delivery_timezone=?,
            delivery_contact_name=?,delivery_contact_phone=?,delivery_instructions=?
            WHERE id=? AND organization_id=?""",
            (
                values["pickup_address"], values["pickup_window_start"], values["pickup_window_end"],
                values["pickup_timezone"], values.get("pickup_contact_name", ""),
                values.get("pickup_contact_phone", ""), values.get("pickup_instructions", ""),
                values["delivery_address"], values["delivery_window_start"], values["delivery_window_end"],
                values["delivery_timezone"], values.get("delivery_contact_name", ""),
                values.get("delivery_contact_phone", ""), values.get("delivery_instructions", ""),
                load["id"], user["organization_id"],
            ),
        )
        for stop_type, prefix in (("PICKUP", "pickup"), ("DELIVERY", "delivery")):
            conn.execute(
                """UPDATE load_stops SET address_line1=?,iana_timezone=?,appointment_local_start=?,
                appointment_local_end=?,contact_name=?,contact_phone=?,instructions=?,updated_at=?
                WHERE organization_id=? AND load_id=? AND stop_type=?""",
                (
                    values[f"{prefix}_address"], values[f"{prefix}_timezone"],
                    values[f"{prefix}_window_start"], values[f"{prefix}_window_end"],
                    values.get(f"{prefix}_contact_name", ""), values.get(f"{prefix}_contact_phone", ""),
                    values.get(f"{prefix}_instructions", ""), utc_now_iso(),
                    user["organization_id"], load["id"], stop_type,
                ),
            )
        _insert_load_snapshot(
            conn,
            organization_id=user["organization_id"], load_id=load["id"], stage="RATECON_CONFIRMED",
            inputs={"booking_snapshot_id": load.get("booking_snapshot_id"), "reviewed_fields": values},
            result={"confirmed_rate_cents": money_to_cents(values["total_rate"], field="RateCon total"), "differences": differences},
            user_id=user["id"],
        )
        _transition_workflow(
            conn, organization_id=user["organization_id"], load_id=load["id"],
            target=LoadState.NEEDS_ASSIGNMENT, actor_user_id=user["id"],
            idempotency_key=f"ratecon:{document_uuid}:approved",
            reason="Human approved RateCon facts and material differences",
        )
        conn.execute(
            "UPDATE operational_documents SET processing_status='REVIEWED' WHERE id=? AND organization_id=?",
            (document["id"], user["organization_id"]),
        )
    record_audit_event(
        "ratecon.approved", int(user["organization_id"]), int(user["id"]),
        {"document_uuid": document_uuid, "load_id": document["load_id"], "material_differences": len(material)},
    )
    set_flash(request, "RateCon review approved. The load is ready for driver, truck, and trailer assignment.")
    return redirect(f"/loads/{document['load_public_uuid']}/dispatch")


@app.get("/ratecons/{document_uuid}/download")
@app.get("/documents/{document_uuid}/download")
def issue_ratecon_download(request: Request, document_uuid: str):
    user = require_permission(request, "documents.view_operational")
    _ratecon_document(user["organization_id"], document_uuid)
    serializer = URLSafeTimedSerializer(SESSION_SECRET, salt="carrieros-private-document-v1")
    token = serializer.dumps({"document": document_uuid, "organization": user["organization_id"]})
    record_audit_event(
        "ratecon.download_link_issued", int(user["organization_id"]), int(user["id"]),
        {"document_uuid": document_uuid},
    )
    return redirect(f"/private-documents/{document_uuid}?token={quote(token)}")


@app.get("/private-documents/{document_uuid}")
def download_private_document(request: Request, document_uuid: str, token: str):
    user = require_permission(request, "documents.view_operational")
    serializer = URLSafeTimedSerializer(SESSION_SECRET, salt="carrieros-private-document-v1")
    try:
        payload = serializer.loads(token, max_age=300)
    except SignatureExpired as exc:
        raise HTTPException(410, "This private download link expired") from exc
    except BadSignature as exc:
        raise HTTPException(404, "Private document not found") from exc
    if payload.get("document") != document_uuid or int(payload.get("organization") or 0) != int(user["organization_id"]):
        raise HTTPException(404, "Private document not found")
    document = _ratecon_document(user["organization_id"], document_uuid)
    try:
        body = configured_storage_provider().get(document["storage_key"])
    except (FileNotFoundError, KeyError) as exc:
        raise HTTPException(404, "Private document object not found") from exc
    record_audit_event(
        "ratecon.downloaded", int(user["organization_id"]), int(user["id"]),
        {"document_uuid": document_uuid},
    )
    safe_name = quote(document["original_filename"])
    return Response(
        body,
        media_type=document["content_type"],
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}", "Cache-Control": "no-store"},
    )


def _dispatch_page_context(organization_id: int, load_uuid: str) -> dict[str, Any]:
    load = _load_by_public_uuid(organization_id, load_uuid)
    settings = dict(query_one("SELECT * FROM organizations WHERE id=?", (organization_id,)))
    opportunity = query_one(
        "SELECT equipment_type FROM load_opportunities WHERE id=? AND organization_id=?",
        (load.get("opportunity_id") or 0, organization_id),
    )
    stops = [dict(row) for row in query_all(
        """SELECT * FROM load_stops WHERE organization_id=? AND load_id=?
        ORDER BY sequence_number""",
        (organization_id, load["id"]),
    )]
    pickup_stop = next((item for item in stops if item["stop_type"] == "PICKUP"), {})
    delivery_stop = next((item for item in reversed(stops) if item["stop_type"] == "DELIVERY"), {})
    load.update(
        origin_city=pickup_stop.get("city") or "",
        origin_state=pickup_stop.get("state") or "",
        origin_postal_code=pickup_stop.get("postal_code") or "",
        destination_city=delivery_stop.get("city") or "",
        destination_state=delivery_stop.get("state") or "",
        destination_postal_code=delivery_stop.get("postal_code") or "",
        pickup_at=pickup_stop.get("appointment_local_start") or load.get("pickup_window_start") or load.get("pickup_date"),
        delivery_at=delivery_stop.get("appointment_local_start") or load.get("delivery_window_start") or load.get("delivery_date"),
        equipment_requirement=str(opportunity["equipment_type"] if opportunity else ""),
    )
    drivers = [dict(row) for row in query_all(
        "SELECT * FROM drivers WHERE organization_id=? AND active=1 ORDER BY lower(name)",
        (organization_id,),
    )]
    power_units_by_driver: dict[int, dict[str, Any] | None] = {}
    for driver in drivers:
        assignment = query_one(
            """SELECT p.* FROM equipment_assignments e JOIN power_units p ON p.id=e.power_unit_id
            WHERE e.organization_id=? AND e.driver_id=? AND e.end_at IS NULL AND p.active=1
            ORDER BY e.start_at DESC,e.id DESC LIMIT 1""",
            (organization_id, driver["id"]),
        )
        if not assignment and driver.get("vehicle_id"):
            assignment = query_one(
                """SELECT * FROM power_units WHERE organization_id=? AND legacy_vehicle_id=? AND active=1""",
                (organization_id, driver["vehicle_id"]),
            )
        power_units_by_driver[int(driver["id"])] = dict(assignment) if assignment else None
    trailers = [dict(row) for row in query_all(
        "SELECT * FROM trailers WHERE organization_id=? AND active=1 ORDER BY unit_number",
        (organization_id,),
    )]
    schedule_blockers: dict[int, list[str]] = {}
    compliance_blockers: dict[int, list[str]] = {}
    for driver in drivers:
        conflict = query_one(
            """SELECT load_number,pickup_date,delivery_date FROM loads
            WHERE organization_id=? AND driver_id=? AND id<>?
              AND include_in_model=1 AND status_code NOT IN ('CANCELLED','CLOSED')
              AND COALESCE(pickup_date,'9999-12-31')<=COALESCE(?,'9999-12-31')
              AND COALESCE(delivery_date,'0001-01-01')>=COALESCE(?,'0001-01-01')
            ORDER BY pickup_date LIMIT 1""",
            (organization_id, driver["id"], load["id"], load.get("delivery_date"), load.get("pickup_date")),
        )
        if conflict:
            schedule_blockers[int(driver["id"])] = [
                f"Schedule conflicts with {conflict['load_number']} ({conflict['pickup_date']} to {conflict['delivery_date']})"
            ]
        expired = query_all(
            """SELECT document_type,expiration_date FROM compliance_items
            WHERE organization_id=? AND lower(subject_type)='driver' AND subject_id=?
              AND expiration_date IS NOT NULL AND expiration_date<COALESCE(?,'9999-12-31')""",
            (organization_id, driver["id"], load.get("delivery_date")),
        )
        if expired:
            compliance_blockers[int(driver["id"])] = [
                f"{row['document_type']} expires {row['expiration_date']}" for row in expired
            ]
    bundle = get_bundle(organization_id)
    pickup_date = parse_date(load.get("pickup_date")) or date.today()
    fuel_price = fuel_price_for(pickup_date, bundle["weekly_fuel"], number(settings["fallback_diesel_price"]))
    candidates = rank_assignments(
        settings, load, drivers, power_units_by_driver, trailers,
        latest_driver_locations(organization_id), configured_route_provider(),
        schedule_blockers=schedule_blockers, compliance_blockers=compliance_blockers,
        fuel_price=fuel_price,
    )
    current_assignment = query_one(
        """SELECT a.*,d.name AS driver_name,p.company_unit_number,t.unit_number AS trailer_number
        FROM load_assignments a
        LEFT JOIN drivers d ON d.id=a.driver_id
        LEFT JOIN power_units p ON p.id=a.power_unit_id
        LEFT JOIN trailers t ON t.id=a.trailer_id
        WHERE a.organization_id=? AND a.load_id=? AND a.ended_at IS NULL
        ORDER BY a.created_at DESC,a.id DESC LIMIT 1""",
        (organization_id, load["id"]),
    )
    approval = query_one(
        """SELECT * FROM dispatch_approvals WHERE organization_id=? AND load_id=?
        ORDER BY id DESC LIMIT 1""",
        (organization_id, load["id"]),
    )
    ack_url = ""
    if approval and approval["status"] == "AWAITING_DRIVER_ACK":
        ack_url = f"{CANONICAL_BASE_URL}/driver/dispatch/{_dispatch_token(approval['public_uuid'])}"
    return {
        "load": load,
        "stops": stops,
        "candidates": candidates,
        "current_assignment": dict(current_assignment) if current_assignment else None,
        "approval": dict(approval) if approval else None,
        "ack_url": ack_url,
        "hos_disclaimer": HOS_DISCLAIMER,
    }


@app.get("/loads/{load_uuid}/dispatch", response_class=HTMLResponse)
def dispatch_assignment_page(request: Request, load_uuid: str):
    user = require_permission(request, "loads.view")
    return render(request, "dispatch_assignment.html", _dispatch_page_context(user["organization_id"], load_uuid))


@app.post("/loads/{load_uuid}/dispatch/assign")
async def approve_load_assignment(request: Request, load_uuid: str):
    user = require_permission(request, "loads.manage")
    form = await verified_form(request)
    context = _dispatch_page_context(user["organization_id"], load_uuid)
    load = context["load"]
    if normalize_state(load["status_code"]) != LoadState.NEEDS_ASSIGNMENT:
        raise HTTPException(409, "Complete and approve the RateCon review before assignment")
    driver_id = integer(form.get("driver_id"))
    power_unit_id = integer(form.get("power_unit_id"))
    trailer_id = integer(form.get("trailer_id")) or None
    candidate = next(
        (
            item for item in context["candidates"]
            if int(item.driver["id"]) == driver_id
            and item.power_unit and int(item.power_unit["id"]) == power_unit_id
            and int((item.trailer or {}).get("id") or 0) == int(trailer_id or 0)
        ),
        None,
    )
    if not candidate or not candidate.eligible:
        raise HTTPException(400, "Choose an eligible driver, power unit, and trailer combination")
    now = utc_now_iso()
    with db_session() as conn:
        conn.execute(
            "UPDATE load_assignments SET ended_at=? WHERE organization_id=? AND load_id=? AND ended_at IS NULL",
            (now, user["organization_id"], load["id"]),
        )
        assignment_id = int(conn.execute(
            """INSERT INTO load_assignments
            (public_uuid,organization_id,load_id,driver_id,power_unit_id,trailer_id,
             assignment_stage,provisional,deadhead_origin,deadhead_miles_micros,
             route_source,approved_by,approved_at)
            VALUES (?,?,?,?,?,?,'RATECON_APPROVED',0,?,?,?,?,?)""",
            (
                str(uuid.uuid4()), user["organization_id"], load["id"], driver_id,
                power_unit_id, trailer_id,
                str((candidate.location or {}).get("city") or ""),
                round(candidate.result.deadhead_miles * 1_000_000), candidate.deadhead_source,
                user["id"], now,
            ),
        ).lastrowid)
        conn.execute(
            """UPDATE loads SET driver_id=?,vehicle_id=?,deadhead_miles=?
            WHERE id=? AND organization_id=?""",
            (
                driver_id, candidate.power_unit.get("legacy_vehicle_id"),
                candidate.result.deadhead_miles, load["id"], user["organization_id"],
            ),
        )
        _insert_load_snapshot(
            conn, organization_id=user["organization_id"], load_id=load["id"], stage="RATECON_CONFIRMED",
            inputs={
                "assignment": {
                    "driver_id": driver_id, "power_unit_id": power_unit_id, "trailer_id": trailer_id,
                    "deadhead_source": candidate.deadhead_source,
                },
                "prior_ratecon_snapshot": "preserved",
            },
            result=candidate.result.to_dict(), user_id=user["id"],
        )
        _transition_workflow(
            conn, organization_id=user["organization_id"], load_id=load["id"],
            target=LoadState.DISPATCH_AWAITING_APPROVAL, actor_user_id=user["id"],
            idempotency_key=f"assignment:{assignment_id}:approved",
            reason="Authorized user approved driver and equipment assignment",
        )
        conn.execute(
            """INSERT INTO dispatch_approvals
            (public_uuid,organization_id,load_id,load_assignment_id,status)
            VALUES (?,?,?,?,'AWAITING_APPROVAL')""",
            (str(uuid.uuid4()), user["organization_id"], load["id"], assignment_id),
        )
    record_audit_event(
        "dispatch.assignment_approved", int(user["organization_id"]), int(user["id"]),
        {"load_id": load["id"], "driver_id": driver_id, "power_unit_id": power_unit_id, "trailer_id": trailer_id},
    )
    set_flash(request, "Assignment approved and profitability recalculated. Dispatch still requires final approval.")
    return redirect(f"/loads/{load_uuid}/dispatch")


@app.post("/loads/{load_uuid}/dispatch/approve")
async def approve_dispatch(request: Request, load_uuid: str):
    user = require_permission(request, "dispatch.approve")
    await verified_form(request)
    load = _load_by_public_uuid(user["organization_id"], load_uuid)
    approval = query_one(
        """SELECT * FROM dispatch_approvals WHERE organization_id=? AND load_id=?
        AND status='AWAITING_APPROVAL' ORDER BY id DESC LIMIT 1""",
        (user["organization_id"], load["id"]),
    )
    if not approval:
        raise HTTPException(409, "An approved assignment is required before dispatch approval")
    with db_session() as conn:
        now = utc_now_iso()
        conn.execute(
            """UPDATE dispatch_approvals SET status='AWAITING_DRIVER_ACK',approved_by=?,approved_at=?
            WHERE id=? AND organization_id=? AND status='AWAITING_APPROVAL'""",
            (user["id"], now, approval["id"], user["organization_id"]),
        )
        _transition_workflow(
            conn, organization_id=user["organization_id"], load_id=load["id"],
            target=LoadState.DISPATCHED_AWAITING_ACK, actor_user_id=user["id"],
            idempotency_key=f"dispatch:{approval['public_uuid']}:approved",
            reason="Authorized dispatcher approved dispatch",
        )
    record_audit_event(
        "dispatch.approved", int(user["organization_id"]), int(user["id"]),
        {"load_id": load["id"], "approval_uuid": approval["public_uuid"]},
    )
    set_flash(request, "Dispatch approved. Send the secure acknowledgment link to the assigned driver.")
    return redirect(f"/loads/{load_uuid}/dispatch")


def _driver_dispatch_record(token: str) -> dict[str, Any]:
    approval_uuid = _read_dispatch_token(token)
    row = query_one(
        """SELECT a.*,l.public_uuid AS load_public_uuid,l.load_number,l.broker,l.status_code,
        l.ratecon_reference,l.ratecon_received_at,
        l.pickup_address,l.pickup_window_start,l.pickup_window_end,l.pickup_contact_name,
        l.pickup_contact_phone,l.pickup_instructions,l.pickup_timezone,l.delivery_address,
        l.delivery_window_start,l.delivery_window_end,l.delivery_contact_name,
        l.delivery_contact_phone,l.delivery_instructions,l.delivery_timezone,
        d.name AS driver_name,d.phone AS driver_phone
        FROM dispatch_approvals a
        JOIN loads l ON l.id=a.load_id AND l.organization_id=a.organization_id
        JOIN load_assignments la ON la.id=a.load_assignment_id AND la.organization_id=a.organization_id
        LEFT JOIN drivers d ON d.id=la.driver_id AND d.organization_id=a.organization_id
        WHERE a.public_uuid=?""",
        (approval_uuid,),
    )
    if not row:
        raise HTTPException(404, "Dispatch link not found")
    return dict(row)


@app.get("/driver/dispatch/{token}", response_class=HTMLResponse)
def driver_dispatch_ack_page(request: Request, token: str):
    return _render_driver_dispatch_page(request, token, _driver_dispatch_record(token))


@app.post("/driver/dispatch/{token}/ack", response_class=HTMLResponse)
async def acknowledge_driver_dispatch(request: Request, token: str):
    record = _driver_dispatch_record(token)
    form = await request.form()
    note = str(form.get("note") or "").strip()[:500]
    if record["status"] == "AWAITING_DRIVER_ACK":
        with db_session() as conn:
            now = utc_now_iso()
            conn.execute(
                """UPDATE dispatch_approvals SET status='ACKNOWLEDGED',acknowledged_at=?,acknowledgement_note=?
                WHERE id=? AND organization_id=? AND status='AWAITING_DRIVER_ACK'""",
                (now, note, record["id"], record["organization_id"]),
            )
            _transition_workflow(
                conn, organization_id=record["organization_id"], load_id=record["load_id"],
                target=LoadState.DISPATCH_ACKNOWLEDGED, actor_user_id=None,
                idempotency_key=f"dispatch:{record['public_uuid']}:driver-ack",
                reason="Assigned driver acknowledged the dispatch link",
            )
        record_audit_event(
            "dispatch.driver_acknowledged", int(record["organization_id"]), None,
            {"load_id": record["load_id"], "approval_uuid": record["public_uuid"]},
        )
    return _render_driver_dispatch_page(request, token, _driver_dispatch_record(token))


def _delivery_documents_for_load(organization_id: int, load_id: int) -> list[dict[str, Any]]:
    rows = query_all(
        """SELECT l.*,d.public_uuid AS document_public_uuid,d.original_filename,
        d.content_type,d.size_bytes,d.page_count,d.malware_status,d.processing_status,
        d.storage_key,d.created_at AS document_created_at
        FROM delivery_document_links l
        JOIN operational_documents d ON d.id=l.document_id
        WHERE l.organization_id=? AND l.load_id=? AND d.deleted_at IS NULL
        ORDER BY l.created_at DESC,l.id DESC""",
        (organization_id, load_id),
    )
    return [dict(row) for row in rows]


def _driver_dispatch_payload(record: dict[str, Any]) -> dict[str, Any]:
    return {
        **record,
        "driver": {"name": record.get("driver_name"), "phone": record.get("driver_phone")},
    }


def _render_driver_dispatch_page(request: Request, token: str, record: dict[str, Any]) -> HTMLResponse:
    package = driver_dispatch_package(_driver_dispatch_payload(record))
    return render(request, "driver_dispatch_ack.html", {
        "dispatch_record": record,
        "dispatch_package": package,
        "dispatch_token": token,
        "delivery_documents": _delivery_documents_for_load(int(record["organization_id"]), int(record["load_id"])),
        "public_driver_page": True,
        "public_page": True,
    })


@app.post("/driver/dispatch/{token}/status", response_class=HTMLResponse)
async def update_driver_load_status(request: Request, token: str):
    record = _driver_dispatch_record(token)
    if record["status"] not in {"AWAITING_DRIVER_ACK", "ACKNOWLEDGED"}:
        raise HTTPException(409, "A driver status update requires an approved dispatch")
    if not driver_dispatch_package(_driver_dispatch_payload(record))["ready"]:
        raise HTTPException(409, "Complete the RateCon dispatch details before sending delivery status updates")
    form = await request.form()
    target_value = str(form.get("status") or "").strip()
    reason = str(form.get("reason") or "Driver status update").strip()[:500]
    idempotency_key = str(form.get("idempotency_key") or "").strip()[:200]
    if not idempotency_key:
        raise HTTPException(400, "A status update idempotency key is required")
    try:
        target = validate_driver_transition(str(record.get("status_code") or ""), target_value)
    except LoadStateError as exc:
        raise HTTPException(409, str(exc)) from exc
    try:
        with db_session() as conn:
            _transition_workflow(
                conn,
                organization_id=int(record["organization_id"]),
                load_id=int(record["load_id"]),
                target=target,
                actor_user_id=None,
                idempotency_key=f"driver:{record['public_uuid']}:{idempotency_key}",
                reason=reason,
            )
    except LoadStateError as exc:
        raise HTTPException(409, str(exc)) from exc
    record_audit_event(
        "driver.load_status_updated", int(record["organization_id"]), None,
        {"load_id": record["load_id"], "approval_uuid": record["public_uuid"], "status": target.value},
    )
    return _render_driver_dispatch_page(request, token, _driver_dispatch_record(token))


@app.post("/driver/dispatch/{token}/documents", response_class=HTMLResponse)
async def upload_driver_delivery_document(request: Request, token: str):
    record = _driver_dispatch_record(token)
    if record["status"] not in {"AWAITING_DRIVER_ACK", "ACKNOWLEDGED"}:
        raise HTTPException(409, "A delivery document requires an approved dispatch")
    if not driver_dispatch_package(_driver_dispatch_payload(record))["ready"]:
        raise HTTPException(409, "Complete the RateCon dispatch details before uploading delivery documents")
    form = await request.form()
    try:
        kind = parse_delivery_document_kind(str(form.get("document_kind") or ""))
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    upload = form.get("document")
    if upload is None or not hasattr(upload, "read"):
        raise HTTPException(400, "Choose a BOL, POD, receipt, or detention-evidence file")
    payload = await upload.read(MAX_RATECON_BYTES + 1)
    filename = Path(str(getattr(upload, "filename", "delivery-document"))).name.replace("\x00", "")[:255]
    try:
        validated = validate_ratecon_upload(
            payload,
            filename=filename,
            claimed_content_type=str(getattr(upload, "content_type", "")),
        )
    except RateConError as exc:
        raise HTTPException(400, str(exc)) from exc
    storage = configured_storage_provider()
    if not storage.secure_at_rest:
        raise HTTPException(503, "Private encrypted document storage is not configured")
    scan = configured_malware_scan_provider().scan(payload, filename=filename)
    if scan.status != "CLEAN":
        raise HTTPException(400, f"The delivery document cannot be accepted until malware screening is CLEAN ({scan.status})")
    document_uuid = str(uuid.uuid4())
    storage_key = f"organizations/{record['organization_id']}/loads/{record['load_id']}/delivery/{document_uuid}.{validated.extension}"
    storage.put(storage_key, payload, content_type=validated.media_type)
    assignment = query_one(
        """SELECT driver_id FROM load_assignments
        WHERE organization_id=? AND load_id=? AND ended_at IS NULL
        ORDER BY id DESC LIMIT 1""",
        (record["organization_id"], record["load_id"]),
    )
    try:
        with db_session() as conn:
            document_id = int(conn.execute(
                """INSERT INTO operational_documents
                (public_uuid,organization_id,load_id,document_type,storage_key,storage_provider,
                 original_filename,content_type,size_bytes,page_count,sha256,malware_status,
                 processing_status,retention_date,created_by)
                VALUES (?,?,?, ?,?,?,?,?,?,?,?,?,?,?,?)""",
                (
                    document_uuid, record["organization_id"], record["load_id"], kind.value,
                    storage_key, storage.name, filename, validated.media_type, validated.size_bytes,
                    validated.page_count, validated.sha256, scan.status, "RECEIVED",
                    default_retention_date(), None,
                ),
            ).lastrowid)
            conn.execute(
                """INSERT INTO delivery_document_links
                (public_uuid,organization_id,document_id,load_id,document_kind,source,captured_by_driver,notes)
                VALUES (?,?,?,?,?,'driver_portal',?,?)""",
                (
                    str(uuid.uuid4()), record["organization_id"], document_id, record["load_id"],
                    kind.value, assignment["driver_id"] if assignment else None,
                    str(form.get("notes") or "").strip()[:500],
                ),
            )
    except sqlite3.IntegrityError as exc:
        storage.delete(storage_key)
        raise HTTPException(409, "This delivery document was already uploaded for the load") from exc
    record_audit_event(
        "delivery_document.uploaded", int(record["organization_id"]), None,
        {"load_id": record["load_id"], "document_uuid": document_uuid, "document_kind": kind.value, "sha256": validated.sha256},
    )
    return _render_driver_dispatch_page(request, token, _driver_dispatch_record(token))


@app.post("/loads/{load_uuid}/status")
async def update_load_status(request: Request, load_uuid: str):
    user = require_permission(request, "loads.manage")
    form = await verified_form(request)
    load = _load_by_public_uuid(user["organization_id"], load_uuid)
    target_value = str(form.get("status") or "").strip()
    reason = str(form.get("reason") or "Office status update").strip()[:500]
    idempotency_key = str(form.get("idempotency_key") or "").strip()[:200]
    if not idempotency_key:
        raise HTTPException(400, "A status update idempotency key is required")
    try:
        target = LoadState(target_value.upper())
    except ValueError as exc:
        raise HTTPException(400, "Choose a valid load status") from exc
    try:
        with db_session() as conn:
            _transition_workflow(
                conn,
                organization_id=int(user["organization_id"]), load_id=int(load["id"]),
                target=target, actor_user_id=int(user["id"]),
                idempotency_key=f"office:{load_uuid}:{idempotency_key}", reason=reason,
            )
    except LoadStateError as exc:
        raise HTTPException(409, str(exc)) from exc
    record_audit_event(
        "load.status_updated", int(user["organization_id"]), int(user["id"]),
        {"load_id": load["id"], "status": target.value, "reason": reason},
    )
    return redirect(f"/loads/{load['id']}")


@app.get("/vehicles", response_class=HTMLResponse)
def vehicles_page(request: Request):
    user = require_user(request)
    rows = query_all("SELECT * FROM vehicles WHERE organization_id=? ORDER BY active DESC,name", (user["organization_id"],))
    trailers = query_all(
        "SELECT * FROM trailers WHERE organization_id=? ORDER BY active DESC,unit_number",
        (user["organization_id"],),
    )
    return render(request, "vehicles.html", {
        "vehicles": [dict(r) for r in rows],
        "trailers": [dict(r) for r in trailers],
        "unit_limit": int(user["active_unit_limit"] if user.get("active_unit_limit") is not None else 2),
    })


@app.post("/vehicles")
async def create_vehicle(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    active = 1 if yes(form.get("active", "on")) else 0
    if active:
        active_count = query_one(
            "SELECT COUNT(*) AS total FROM vehicles WHERE organization_id=? AND active=1",
            (user["organization_id"],),
        )
        unit_limit = int(user["active_unit_limit"] if user.get("active_unit_limit") is not None else 2)
        if int(active_count["total"] if active_count else 0) >= unit_limit:
            set_flash(request, "Your plan's active-unit limit has been reached. Upgrade from Billing to add another active unit.")
            return redirect("/vehicles")
    try:
        name = str(form.get("name", "New unit")).strip()
        equipment_type = str(form.get("equipment_type", "Truck")).strip()
        with db_session() as conn:
            public_uuid = str(uuid.uuid4())
            vehicle_id = int(conn.execute(
                "INSERT INTO vehicles (organization_id,public_uuid,name,equipment_type,active) VALUES (?,?,?,?,?)",
                (user["organization_id"], public_uuid, name, equipment_type, active),
            ).lastrowid)
            conn.execute(
                """INSERT INTO power_units
                (public_uuid,organization_id,legacy_vehicle_id,public_identifier,
                 company_unit_number,equipment_type,active,legacy_snapshot_json)
                VALUES (?,?,?,?,?,?,?,?)""",
                (
                    str(uuid.uuid4()), user["organization_id"], vehicle_id, name, name,
                    equipment_type, active, json.dumps({"source": "vehicle_create", "vehicle_public_uuid": public_uuid}),
                ),
            )
    except Exception as exc:
        raise HTTPException(400, "That unit name already exists") from exc
    set_flash(request, "Unit added. Assign it on Driver & Equipment Setup.")
    return redirect("/vehicles")


@app.post("/trailers")
async def create_trailer(request: Request):
    user = require_permission(request, "loads.manage")
    form = await verified_form(request)
    unit_number = str(form.get("unit_number") or "").strip()
    trailer_type = str(form.get("trailer_type") or "").strip()
    if not unit_number or not trailer_type:
        raise HTTPException(400, "Enter a trailer number and type")
    try:
        execute(
            """INSERT INTO trailers
            (public_uuid,organization_id,unit_number,trailer_type,gvwr_lbs,empty_weight_lbs,
             dimensions,active)
            VALUES (?,?,?,?,?,?,?,?)""",
            (
                str(uuid.uuid4()), user["organization_id"], unit_number, trailer_type,
                integer(form.get("gvwr_lbs")) or None, integer(form.get("empty_weight_lbs")) or None,
                str(form.get("dimensions") or "").strip(), 1 if yes(form.get("active", "on")) else 0,
            ),
        )
    except sqlite3.IntegrityError as exc:
        raise HTTPException(400, "That trailer number already exists") from exc
    record_audit_event(
        "trailer.created", int(user["organization_id"]), int(user["id"]),
        {"unit_number": unit_number, "trailer_type": trailer_type},
    )
    set_flash(request, "Trailer added and available for RateCon assignment checks.")
    return redirect("/vehicles")


@app.post("/vehicles/{vehicle_id}/update")
async def update_vehicle(request: Request, vehicle_id: int):
    user = require_user(request)
    form = await verified_form(request)
    existing = query_one(
        "SELECT * FROM vehicles WHERE id=? AND organization_id=?",
        (vehicle_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Unit not found")
    active = 1 if yes(form.get("active")) else 0
    if active and not int(existing["active"]):
        active_count = query_one(
            "SELECT COUNT(*) AS total FROM vehicles WHERE organization_id=? AND active=1 AND id<>?",
            (user["organization_id"], vehicle_id),
        )
        unit_limit = int(user["active_unit_limit"] if user.get("active_unit_limit") is not None else 2)
        if int(active_count["total"] if active_count else 0) >= unit_limit:
            set_flash(request, "Your plan's active-unit limit has been reached. Deactivate another unit or upgrade first.")
            return redirect("/vehicles")
    try:
        name = str(form.get("name", "")).strip() or str(existing["name"])
        equipment_type = str(form.get("equipment_type", "Truck")).strip()
        with db_session() as conn:
            conn.execute(
                "UPDATE vehicles SET name=?,equipment_type=?,active=? WHERE id=? AND organization_id=?",
                (name, equipment_type, active, vehicle_id, user["organization_id"]),
            )
            conn.execute(
                """UPDATE power_units SET public_identifier=?,company_unit_number=?,equipment_type=?,
                active=?,updated_at=? WHERE organization_id=? AND legacy_vehicle_id=?""",
                (name, name, equipment_type, active, utc_now_iso(), user["organization_id"], vehicle_id),
            )
    except Exception as exc:
        raise HTTPException(400, "That unit name already exists") from exc
    record_audit_event(
        "vehicle.updated",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"vehicle_id": vehicle_id, "active": bool(active)},
    )
    set_flash(request, "Unit updated. Existing assignments remain connected.")
    return redirect("/vehicles")


@app.get("/drivers", response_class=HTMLResponse)
def drivers_page(request: Request):
    user = require_user(request)
    bundle, state = get_state(user["organization_id"])
    vehicles = {int(v["id"]): v for v in bundle["vehicles"]}
    drivers = []
    balances = {int(b["driver_id"]): b for b in state["driver_balances"]}
    locations = latest_driver_locations(user["organization_id"])
    for row in bundle["drivers"]:
        item = dict(row)
        item["vehicle"] = vehicles.get(int(row.get("vehicle_id") or 0))
        item["monthly_fixed"] = driver_monthly_fixed(row)
        item["balance"] = balances.get(int(row["id"]))
        item["location"] = locations.get(int(row["id"]))
        drivers.append(item)
    return render(request, "drivers.html", {"drivers": drivers, "vehicles": bundle["vehicles"]})


@app.post("/drivers")
async def create_driver(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    driver_id = execute(
        """INSERT INTO drivers
        (organization_id,public_uuid,vehicle_id,name,email,phone,role,pay_model,equipment_type,
         fixed_cost_start,fixed_cost_end,truck_financing_monthly,auto_insurance_monthly,
         trailer_financing_monthly,trailer_insurance_monthly,other_fixed_monthly,mpg,
         maintenance_per_mile,driver_profit_split_pct,contractor_gross_split_pct,
         owner_operator_split_pct,flat_rate_per_load,pay_per_loaded_mile,
         pay_per_total_mile,day_rate,payroll_burden_applies,notes,active)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            user["organization_id"], str(uuid.uuid4()), integer(form.get("vehicle_id")) or None, str(form.get("name", "New driver")).strip(),
            str(form.get("email", "")).strip(), str(form.get("phone", "")).strip(), str(form.get("role", "Driver")),
            str(form.get("pay_model", "Profit Split")), str(form.get("equipment_type", "")).strip(),
            str(form.get("fixed_cost_start", "")) or None, str(form.get("fixed_cost_end", "")) or None,
            number(form.get("truck_financing_monthly")), number(form.get("auto_insurance_monthly")),
            number(form.get("trailer_financing_monthly")), number(form.get("trailer_insurance_monthly")),
            number(form.get("other_fixed_monthly")), number(form.get("mpg"), 10), number(form.get("maintenance_per_mile"), .20),
            number(form.get("driver_profit_split_pct")), number(form.get("contractor_gross_split_pct")),
            number(form.get("owner_operator_split_pct")), number(form.get("flat_rate_per_load")),
            number(form.get("pay_per_loaded_mile")), number(form.get("pay_per_total_mile")),
            number(form.get("day_rate")), 1 if yes(form.get("payroll_burden_applies")) else 0,
            str(form.get("notes", "")).strip(), 1 if yes(form.get("active", "on")) else 0,
        ),
    )
    set_flash(request, "Driver and equipment profile added.")
    return redirect(f"/drivers#driver-{driver_id}")


@app.post("/drivers/{driver_id}/update")
async def update_driver(request: Request, driver_id: int):
    user = require_user(request)
    form = await verified_form(request)
    with db_session() as conn:
        conn.execute(
            """UPDATE drivers SET vehicle_id=?,name=?,email=?,phone=?,role=?,pay_model=?,equipment_type=?,
            fixed_cost_start=?,fixed_cost_end=?,truck_financing_monthly=?,auto_insurance_monthly=?,
            trailer_financing_monthly=?,trailer_insurance_monthly=?,other_fixed_monthly=?,mpg=?,
            maintenance_per_mile=?,driver_profit_split_pct=?,contractor_gross_split_pct=?,
            owner_operator_split_pct=?,flat_rate_per_load=?,pay_per_loaded_mile=?,
            pay_per_total_mile=?,day_rate=?,payroll_burden_applies=?,notes=?,active=?
            WHERE id=? AND organization_id=?""",
            (
                integer(form.get("vehicle_id")) or None, str(form.get("name", "")).strip(), str(form.get("email", "")).strip(),
                str(form.get("phone", "")).strip(), str(form.get("role", "Driver")), str(form.get("pay_model", "Profit Split")),
                str(form.get("equipment_type", "")).strip(), str(form.get("fixed_cost_start", "")) or None,
                str(form.get("fixed_cost_end", "")) or None, number(form.get("truck_financing_monthly")),
                number(form.get("auto_insurance_monthly")), number(form.get("trailer_financing_monthly")),
                number(form.get("trailer_insurance_monthly")), number(form.get("other_fixed_monthly")),
                number(form.get("mpg"), 10), number(form.get("maintenance_per_mile"), .20),
                number(form.get("driver_profit_split_pct")), number(form.get("contractor_gross_split_pct")),
                number(form.get("owner_operator_split_pct")), number(form.get("flat_rate_per_load")),
                number(form.get("pay_per_loaded_mile")), number(form.get("pay_per_total_mile")),
                number(form.get("day_rate")), 1 if yes(form.get("payroll_burden_applies")) else 0,
                str(form.get("notes", "")).strip(), 1 if yes(form.get("active")) else 0,
                driver_id, user["organization_id"],
            ),
        )
    set_flash(request, "Driver setup updated. All load pay and fixed-cost results were recalculated.")
    return redirect(f"/drivers#driver-{driver_id}")


@app.get("/fuel", response_class=HTMLResponse)
def fuel_page(request: Request, edit_id: int | None = None):
    user = require_user(request)
    bundle = get_bundle(user["organization_id"])
    current_week = monday_for(date.today())
    current_effective = fuel_price_for(date.today(), bundle["weekly_fuel"], number(bundle["settings"]["fallback_diesel_price"]))
    editing = None
    if edit_id is not None:
        row = query_one(
            "SELECT * FROM weekly_fuel WHERE id=? AND organization_id=?",
            (edit_id, user["organization_id"]),
        )
        if not row:
            raise HTTPException(404, "Fuel record not found")
        editing = dict(row)
    return render(request, "fuel.html", {
        "rows": list(reversed(bundle["weekly_fuel"])),
        "current_week": current_week,
        "current_effective": current_effective,
        "fallback": bundle["settings"]["fallback_diesel_price"],
        "editing": editing,
    })


@app.post("/fuel")
async def save_fuel(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    week = parse_date(form.get("week_start"))
    if not week:
        raise HTTPException(400, "Enter a valid week-start date")
    week = monday_for(week)
    existing = query_one(
        "SELECT * FROM weekly_fuel WHERE organization_id=? AND week_start=?",
        (user["organization_id"], week.isoformat()),
    )
    with db_session() as conn:
        conn.execute(
            """INSERT INTO weekly_fuel (organization_id,week_start,average_price,source_notes,entered_by)
            VALUES (?,?,?,?,?) ON CONFLICT(organization_id,week_start) DO UPDATE SET
            average_price=excluded.average_price,source_notes=excluded.source_notes,entered_by=excluded.entered_by""",
            (user["organization_id"], week.isoformat(), number(form.get("average_price")), str(form.get("source_notes", "")).strip(), str(form.get("entered_by", user["full_name"])).strip()),
        )
    record_audit_event(
        "fuel.updated" if existing else "fuel.created",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"week_start": week.isoformat()},
    )
    set_flash(request, f"Diesel price saved for the week of {week:%b %d, %Y}. Loads were recalculated.")
    return redirect("/fuel")


@app.get("/payments", response_class=HTMLResponse)
def payments_page(request: Request):
    user = require_user(request)
    bundle, state = get_state(user["organization_id"])
    rows = query_all(
        """SELECT p.*,d.name AS driver_name,l.load_number FROM payments p
        LEFT JOIN drivers d ON d.id=p.driver_id LEFT JOIN loads l ON l.id=p.load_id
        WHERE p.organization_id=? ORDER BY p.paid_at DESC,p.id DESC""",
        (user["organization_id"],),
    )
    return render(request, "payments.html", {
        "payments": [dict(r) for r in rows],
        "drivers": bundle["drivers"],
        "loads": list(reversed(bundle["loads"])),
        "balances": state["driver_balances"],
        "owner": state["owner_pay"],
    })


def payment_form_values(user: dict[str, Any], form: Any) -> tuple[int, int | None, str, str, float]:
    driver_id = integer(form.get("driver_id"))
    if not query_one(
        "SELECT 1 FROM drivers WHERE id=? AND organization_id=?",
        (driver_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid payee")
    load_id = integer(form.get("load_id")) or None
    if load_id and not query_one(
        "SELECT 1 FROM loads WHERE id=? AND organization_id=?",
        (load_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid load")
    paid_at = str(form.get("paid_at", ""))
    if not parse_date(paid_at):
        raise HTTPException(400, "Enter a valid payment date")
    amount = number(form.get("amount"))
    if amount <= 0:
        raise HTTPException(400, "Payment amount must be greater than zero")
    payment_type = str(form.get("payment_type", "Regular payout")).strip()
    return driver_id, load_id, paid_at, payment_type, amount


@app.post("/payments")
async def add_payment(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    driver_id, load_id, paid_at, payment_type, amount = payment_form_values(user, form)
    payment_id = execute(
        """INSERT INTO payments
        (organization_id,driver_id,load_id,paid_at,payment_type,amount,method,reference,notes,
         counts_against_load_pay,include_in_reports)
        VALUES (?,?,?,?,?,?,?,?,?,?,?)""",
        (
            user["organization_id"], driver_id, load_id, paid_at, payment_type,
            amount, str(form.get("method", "")).strip(), str(form.get("reference", "")).strip(),
            str(form.get("notes", "")).strip(), 1 if yes(form.get("counts_against_load_pay")) else 0,
            1 if yes(form.get("include_in_reports")) else 0,
        ),
    )
    record_audit_event(
        "payment.created",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"payment_id": payment_id, "amount": amount},
    )
    set_flash(request, "Payment recorded. Driver balances were recalculated cumulatively.")
    return redirect("/payments")


@app.get("/payments/{payment_id}/edit", response_class=HTMLResponse)
def edit_payment_page(request: Request, payment_id: int):
    user = require_user(request)
    payment = query_one(
        "SELECT * FROM payments WHERE id=? AND organization_id=?",
        (payment_id, user["organization_id"]),
    )
    if not payment:
        raise HTTPException(404, "Payment not found")
    if payment["voided_at"]:
        set_flash(request, "Voided payments are locked to preserve the audit trail.")
        return redirect("/payments")
    bundle = get_bundle(user["organization_id"])
    return render(request, "payment_edit.html", {
        "payment": dict(payment),
        "drivers": bundle["drivers"],
        "loads": list(reversed(bundle["loads"])),
    })


@app.post("/payments/{payment_id}/edit")
async def update_payment(request: Request, payment_id: int):
    user = require_user(request)
    form = await verified_form(request)
    existing = query_one(
        "SELECT * FROM payments WHERE id=? AND organization_id=? AND voided_at IS NULL",
        (payment_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Active payment not found")
    driver_id, load_id, paid_at, payment_type, amount = payment_form_values(user, form)
    execute(
        """UPDATE payments SET driver_id=?,load_id=?,paid_at=?,payment_type=?,amount=?,method=?,
        reference=?,notes=?,counts_against_load_pay=?,include_in_reports=?
        WHERE id=? AND organization_id=? AND voided_at IS NULL""",
        (
            driver_id, load_id, paid_at, payment_type, amount,
            str(form.get("method", "")).strip(), str(form.get("reference", "")).strip(),
            str(form.get("notes", "")).strip(), 1 if yes(form.get("counts_against_load_pay")) else 0,
            1 if yes(form.get("include_in_reports")) else 0,
            payment_id, user["organization_id"],
        ),
    )
    record_audit_event(
        "payment.updated",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"payment_id": payment_id, "previous_amount": number(existing["amount"]), "amount": amount},
    )
    set_flash(request, "Payment updated and every affected balance was recalculated.")
    return redirect("/payments")


@app.post("/payments/{payment_id}/void")
async def void_payment(request: Request, payment_id: int):
    user = require_user(request)
    form = await verified_form(request)
    existing = query_one(
        "SELECT * FROM payments WHERE id=? AND organization_id=? AND voided_at IS NULL",
        (payment_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Active payment not found")
    reason = str(form.get("void_reason", "Correction requested from payment ledger")).strip()
    execute(
        "UPDATE payments SET voided_at=?,voided_by=?,void_reason=? WHERE id=? AND organization_id=? AND voided_at IS NULL",
        (utc_now_iso(), user["id"], reason, payment_id, user["organization_id"]),
    )
    record_audit_event(
        "payment.voided",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"payment_id": payment_id, "amount": number(existing["amount"]), "reason": reason[:120]},
    )
    set_flash(request, "Payment voided. It remains in the ledger but no longer affects balances or reports.")
    return redirect("/payments")


OPPORTUNITY_STATUSES = ("Draft", "Evaluated", "Negotiating", "Declined", "Booked")


def opportunity_row(organization_id: int, opportunity_id: int) -> dict[str, Any]:
    row = query_one(
        "SELECT * FROM load_opportunities WHERE id=? AND organization_id=?",
        (opportunity_id, organization_id),
    )
    if not row:
        raise HTTPException(404, "Rate quote not found")
    return dict(row)


def latest_driver_locations(organization_id: int) -> dict[int, dict[str, Any]]:
    rows = query_all(
        """SELECT location.* FROM driver_locations location
        WHERE location.organization_id=? AND location.id=(
          SELECT candidate.id FROM driver_locations candidate
          WHERE candidate.organization_id=location.organization_id
            AND candidate.driver_id=location.driver_id
          ORDER BY candidate.observed_at DESC,candidate.id DESC LIMIT 1
        )""",
        (organization_id,),
    )
    locations = {int(row["driver_id"]): dict(row) for row in rows}
    fallback_rows = query_all(
        """SELECT d.id AS driver_id,l.destination,l.delivery_date,l.status
        FROM drivers d LEFT JOIN loads l ON l.id=(
          SELECT candidate.id FROM loads candidate
          WHERE candidate.organization_id=d.organization_id AND candidate.driver_id=d.id
            AND lower(trim(candidate.status)) NOT IN ('cancelled','canceled','quote')
          ORDER BY candidate.delivery_date DESC,candidate.id DESC LIMIT 1
        ) WHERE d.organization_id=? AND d.active=1""",
        (organization_id,),
    )
    for row in fallback_rows:
        driver_id = int(row["driver_id"])
        if driver_id in locations or not row["destination"]:
            continue
        destination = str(row["destination"]).strip()
        city, separator, state = destination.partition(",")
        observed_at = f"{row['delivery_date']}T23:59:59+00:00" if row["delivery_date"] else utc_now_iso()
        locations[driver_id] = {
            "driver_id": driver_id,
            "city": city.strip(),
            "state": state.strip() if separator else "",
            "postal_code": "",
            "latitude": None,
            "longitude": None,
            "source": "projected load destination",
            "confidence": "projected",
            "observed_at": observed_at,
        }
    return locations


def opportunity_operational_warnings(
    organization_id: int,
    opportunity: dict[str, Any],
    drivers: list[dict[str, Any]],
) -> dict[int, list[str]]:
    pickup = str(opportunity.get("pickup_at") or "")[:10]
    delivery = str(opportunity.get("delivery_at") or "")[:10]
    warnings: dict[int, list[str]] = {}
    for driver in drivers:
        driver_id = int(driver["id"])
        items: list[str] = []
        if pickup and delivery:
            overlap = query_one(
                """SELECT load_number FROM loads WHERE organization_id=? AND driver_id=?
                AND lower(trim(status)) NOT IN ('cancelled','canceled','quote')
                AND COALESCE(pickup_date,'')<=? AND COALESCE(delivery_date,'')>=?
                ORDER BY pickup_date LIMIT 1""",
                (organization_id, driver_id, delivery, pickup),
            )
            if overlap:
                items.append(f"Schedule conflict with load {overlap['load_number']}.")
        expired = query_one(
            """SELECT document_type FROM compliance_items
            WHERE organization_id=? AND expiration_date<? AND (
              (lower(subject_type)='driver' AND subject_id=?) OR
              (lower(subject_type)='vehicle' AND subject_id=?))
            ORDER BY expiration_date LIMIT 1""",
            (organization_id, date.today().isoformat(), driver_id, driver.get("vehicle_id")),
        )
        if expired:
            items.append(f"Expired compliance item: {expired['document_type']}.")
        if not driver.get("vehicle_id"):
            items.append("No active power unit is assigned to this driver.")
        warnings[driver_id] = items
    return warnings


def opportunity_form_values(form: Any, *, original_rate: float | None = None) -> dict[str, Any]:
    offered = number(form.get("original_offered_rate")) if original_rate is None else original_rate
    return {
        "original_offered_rate": offered,
        "broker_customer": str(form.get("broker_customer", "")).strip(),
        "origin_city": str(form.get("origin_city", "")).strip(),
        "origin_state": str(form.get("origin_state", "")).strip().upper(),
        "origin_postal_code": str(form.get("origin_postal_code", "")).strip(),
        "destination_city": str(form.get("destination_city", "")).strip(),
        "destination_state": str(form.get("destination_state", "")).strip().upper(),
        "destination_postal_code": str(form.get("destination_postal_code", "")).strip(),
        "pickup_at": str(form.get("pickup_at", "")).strip(),
        "delivery_at": str(form.get("delivery_at", "")).strip(),
        "equipment_type": str(form.get("equipment_type", "")).strip(),
        "weight_lbs": optional_number(form.get("weight_lbs")),
        "stops": max(0, integer(form.get("stops"))),
        "commodity": str(form.get("commodity", "")).strip(),
        "pallets_pieces": str(form.get("pallets_pieces", "")).strip(),
        "selected_driver_id": integer(form.get("selected_driver_id")) or None,
        "selected_vehicle_id": integer(form.get("selected_vehicle_id")) or None,
        "loaded_miles": max(0.0, number(form.get("loaded_miles"))),
        "deadhead_miles": max(0.0, number(form.get("deadhead_miles"))),
        "mileage_source": str(form.get("mileage_source", "manual")).strip() or "manual",
        "linehaul_revenue": max(0.0, number(form.get("linehaul_revenue"))),
        "fuel_surcharge": max(0.0, number(form.get("fuel_surcharge"))),
        "additional_revenue": max(0.0, number(form.get("additional_revenue"))),
        "stop_pay_revenue": max(0.0, number(form.get("stop_pay_revenue"))),
        "tolls": max(0.0, number(form.get("tolls"))),
        "lumper": max(0.0, number(form.get("lumper"))),
        "factoring_pct": max(0.0, number(form.get("factoring_pct"))),
        "quick_pay_pct": max(0.0, number(form.get("quick_pay_pct"))),
        "misc_expenses": max(0.0, number(form.get("misc_expenses"))),
        "hazmat": 1 if yes(form.get("hazmat")) else 0,
        "team_required": 1 if yes(form.get("team_required")) else 0,
        "driver_assist": 1 if yes(form.get("driver_assist")) else 0,
        "liftgate": 1 if yes(form.get("liftgate")) else 0,
        "inside_delivery": 1 if yes(form.get("inside_delivery")) else 0,
        "special_equipment": str(form.get("special_equipment", "")).strip(),
        "notes": str(form.get("notes", "")).strip(),
    }


def validate_opportunity_values(organization_id: int, values: dict[str, Any]) -> None:
    if values["original_offered_rate"] <= 0:
        raise HTTPException(400, "Enter the broker's offered rate")
    if not all((values["origin_city"], values["destination_city"], values["pickup_at"], values["delivery_at"])):
        raise HTTPException(400, "Enter origin, destination, pickup, and delivery")
    pickup = parse_date(values["pickup_at"][:10])
    delivery = parse_date(values["delivery_at"][:10])
    if not pickup or not delivery or delivery < pickup:
        raise HTTPException(400, "Enter a valid pickup and delivery window")
    if values["selected_driver_id"]:
        driver = query_one(
            "SELECT vehicle_id FROM drivers WHERE id=? AND organization_id=? AND active=1",
            (values["selected_driver_id"], organization_id),
        )
        if not driver:
            raise HTTPException(400, "Select a valid active driver")
        values["selected_vehicle_id"] = values["selected_vehicle_id"] or driver["vehicle_id"]
    if values["selected_vehicle_id"] and not query_one(
        "SELECT 1 FROM vehicles WHERE id=? AND organization_id=? AND active=1",
        (values["selected_vehicle_id"], organization_id),
    ):
        raise HTTPException(400, "Select a valid active power unit")


def insert_opportunity_snapshot(
    conn: Any,
    opportunity: dict[str, Any],
    settings: dict[str, Any],
    driver: dict[str, Any] | None,
    result: Any,
    stage: str,
    user_id: int,
    revenue: float,
) -> int:
    revision_row = conn.execute(
        "SELECT COALESCE(MAX(revision),0)+1 AS revision FROM opportunity_snapshots WHERE opportunity_id=? AND stage=?",
        (opportunity["id"], stage),
    ).fetchone()
    revision = int(revision_row["revision"])
    return int(conn.execute(
        """INSERT INTO opportunity_snapshots
        (organization_id,opportunity_id,stage,revision,input_json,result_json,created_by)
        VALUES (?,?,?,?,?,?,?)""",
        (
            opportunity["organization_id"], opportunity["id"], stage, revision,
            opportunity_input_snapshot(opportunity, settings, driver, revenue=revenue),
            json.dumps(result.to_dict(), sort_keys=True, separators=(",", ":")), user_id,
        ),
    ).lastrowid)


def calculate_saved_opportunity(
    organization_id: int,
    opportunity: dict[str, Any],
    *,
    revenue_override: float | None = None,
    extra_warnings: list[str] | None = None,
) -> tuple[dict[str, Any], dict[str, Any] | None, Any, float]:
    bundle = get_bundle(organization_id)
    driver = next(
        (item for item in bundle["drivers"] if int(item["id"]) == int(opportunity.get("selected_driver_id") or 0)),
        None,
    )
    pickup = parse_date(str(opportunity.get("pickup_at") or "")[:10]) or date.today()
    fuel_price = fuel_price_for(pickup, bundle["weekly_fuel"], number(bundle["settings"]["fallback_diesel_price"]))
    result = calculate_opportunity(
        bundle["settings"], driver, opportunity,
        revenue_override=revenue_override, fuel_price=fuel_price,
        warnings=extra_warnings or (),
    )
    return bundle, driver, result, fuel_price


@app.get("/rate-quotes", response_class=HTMLResponse)
def rate_quotes_page(request: Request, status: str = ""):
    user = require_user(request)
    params: list[Any] = [user["organization_id"]]
    where = "WHERE organization_id=?"
    if status in OPPORTUNITY_STATUSES:
        where += " AND status=?"
        params.append(status)
    rows = [dict(row) for row in query_all(
        f"SELECT * FROM load_opportunities {where} ORDER BY created_at DESC,id DESC",
        params,
    )]
    return render(request, "rate_quotes.html", {
        "opportunities": rows,
        "selected_status": status if status in OPPORTUNITY_STATUSES else "",
        "statuses": OPPORTUNITY_STATUSES,
    })


def opportunity_form_context(organization_id: int, values: dict[str, Any] | None = None, editing: bool = False):
    bundle = get_bundle(organization_id)
    defaults = values or {
        "pickup_at": f"{date.today().isoformat()}T08:00",
        "delivery_at": f"{date.today().isoformat()}T17:00",
        "mileage_source": "manual",
        "factoring_pct": 0,
        "quick_pay_pct": 0,
        "stops": 0,
    }
    return {"drivers": bundle["drivers"], "vehicles": bundle["vehicles"], "values": defaults, "editing": editing}


@app.get("/rate-quotes/new", response_class=HTMLResponse)
def new_rate_quote_page(request: Request):
    user = require_user(request)
    return render(request, "rate_quote_form.html", opportunity_form_context(user["organization_id"]))


@app.post("/rate-quotes/new")
async def create_rate_quote(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    values = opportunity_form_values(form)
    validate_opportunity_values(user["organization_id"], values)
    columns = ",".join(values)
    placeholders = ",".join("?" for _ in values)
    with db_session() as conn:
        opportunity_id = int(conn.execute(
            f"""INSERT INTO load_opportunities
            (organization_id,status,{columns},created_by,updated_by)
            VALUES (?, 'Evaluated', {placeholders}, ?, ?)""",
            (user["organization_id"], *values.values(), user["id"], user["id"]),
        ).lastrowid)
        opportunity = dict(conn.execute("SELECT * FROM load_opportunities WHERE id=?", (opportunity_id,)).fetchone())
        settings = dict(conn.execute("SELECT * FROM organizations WHERE id=?", (user["organization_id"],)).fetchone())
        driver_row = conn.execute(
            "SELECT * FROM drivers WHERE id=? AND organization_id=?",
            (values["selected_driver_id"] or 0, user["organization_id"]),
        ).fetchone()
        driver = dict(driver_row) if driver_row else None
        bundle = get_bundle(user["organization_id"])
        pickup = parse_date(values["pickup_at"][:10]) or date.today()
        fuel_price = fuel_price_for(pickup, bundle["weekly_fuel"], number(settings["fallback_diesel_price"]))
        result = calculate_opportunity(settings, driver, opportunity, fuel_price=fuel_price)
        insert_opportunity_snapshot(conn, opportunity, settings, driver, result, "evaluation", user["id"], offered_revenue(opportunity))
        conn.execute(
            """INSERT INTO opportunity_negotiations
            (organization_id,opportunity_id,action,amount,notes,created_by)
            VALUES (?,?,?,?,?,?)""",
            (user["organization_id"], opportunity_id, "Offer received", values["original_offered_rate"], "Initial manual offer", user["id"]),
        )
    record_audit_event("opportunity.created", int(user["organization_id"]), int(user["id"]), {"opportunity_id": opportunity_id})
    return redirect(f"/rate-quotes/{opportunity_id}")


@app.get("/rate-quotes/{opportunity_id}/edit", response_class=HTMLResponse)
def edit_rate_quote_page(request: Request, opportunity_id: int):
    user = require_user(request)
    opportunity = opportunity_row(user["organization_id"], opportunity_id)
    if opportunity["status"] in {"Booked", "Declined"}:
        raise HTTPException(409, "Booked or declined quotes are locked")
    return render(request, "rate_quote_form.html", opportunity_form_context(user["organization_id"], opportunity, True))


@app.post("/rate-quotes/{opportunity_id}/edit")
async def update_rate_quote(request: Request, opportunity_id: int):
    user = require_user(request)
    existing = opportunity_row(user["organization_id"], opportunity_id)
    if existing["status"] in {"Booked", "Declined"}:
        raise HTTPException(409, "Booked or declined quotes are locked")
    form = await verified_form(request)
    values = opportunity_form_values(form, original_rate=number(existing["original_offered_rate"]))
    validate_opportunity_values(user["organization_id"], values)
    assignments = ",".join(f"{column}=?" for column in values if column != "original_offered_rate")
    editable_values = [value for column, value in values.items() if column != "original_offered_rate"]
    with db_session() as conn:
        conn.execute(
            f"UPDATE load_opportunities SET {assignments},status='Evaluated',updated_by=?,updated_at=? WHERE id=? AND organization_id=?",
            (*editable_values, user["id"], utc_now_iso(), opportunity_id, user["organization_id"]),
        )
        opportunity = dict(conn.execute("SELECT * FROM load_opportunities WHERE id=?", (opportunity_id,)).fetchone())
        settings = dict(conn.execute("SELECT * FROM organizations WHERE id=?", (user["organization_id"],)).fetchone())
        driver_row = conn.execute("SELECT * FROM drivers WHERE id=? AND organization_id=?", (opportunity.get("selected_driver_id") or 0, user["organization_id"])).fetchone()
        driver = dict(driver_row) if driver_row else None
        bundle = get_bundle(user["organization_id"])
        pickup = parse_date(str(opportunity["pickup_at"])[:10]) or date.today()
        fuel_price = fuel_price_for(pickup, bundle["weekly_fuel"], number(settings["fallback_diesel_price"]))
        result = calculate_opportunity(settings, driver, opportunity, fuel_price=fuel_price)
        insert_opportunity_snapshot(conn, opportunity, settings, driver, result, "evaluation", user["id"], offered_revenue(opportunity))
    record_audit_event("opportunity.updated", int(user["organization_id"]), int(user["id"]), {"opportunity_id": opportunity_id})
    return redirect(f"/rate-quotes/{opportunity_id}")


@app.post("/rate-quotes/{opportunity_id}/select-driver/{driver_id}")
async def select_quote_driver(request: Request, opportunity_id: int, driver_id: int):
    user = require_user(request)
    await verified_form(request)
    opportunity = opportunity_row(user["organization_id"], opportunity_id)
    if opportunity["status"] in {"Booked", "Declined"}:
        raise HTTPException(409, "This quote is locked")
    driver = query_one("SELECT * FROM drivers WHERE id=? AND organization_id=? AND active=1", (driver_id, user["organization_id"]))
    if not driver:
        raise HTTPException(404, "Driver not found")
    execute(
        "UPDATE load_opportunities SET selected_driver_id=?,selected_vehicle_id=?,status='Evaluated',updated_by=?,updated_at=? WHERE id=? AND organization_id=?",
        (driver_id, driver["vehicle_id"], user["id"], utc_now_iso(), opportunity_id, user["organization_id"]),
    )
    revised = opportunity_row(user["organization_id"], opportunity_id)
    bundle, selected_driver, result, _ = calculate_saved_opportunity(user["organization_id"], revised)
    with db_session() as conn:
        insert_opportunity_snapshot(
            conn, revised, bundle["settings"], selected_driver, result,
            "evaluation", user["id"], offered_revenue(revised),
        )
    set_flash(request, f"{driver['name']} selected. The quote was recalculated with that pay profile.")
    return redirect(f"/rate-quotes/{opportunity_id}")


@app.post("/rate-quotes/{opportunity_id}/negotiate")
async def negotiate_rate_quote(request: Request, opportunity_id: int):
    user = require_user(request)
    opportunity = opportunity_row(user["organization_id"], opportunity_id)
    if opportunity["status"] in {"Booked", "Declined"}:
        raise HTTPException(409, "This quote is locked")
    form = await verified_form(request)
    amount = optional_number(form.get("counteroffer_rate"))
    if amount is None or amount <= 0:
        raise HTTPException(400, "Enter a counteroffer")
    response = str(form.get("broker_response", "")).strip()
    notes = str(form.get("negotiation_notes", "")).strip()
    with db_session() as conn:
        conn.execute(
            """UPDATE load_opportunities SET status='Negotiating',counteroffer_rate=?,broker_response=?,
            negotiation_notes=?,updated_by=?,updated_at=? WHERE id=? AND organization_id=?""",
            (amount, response, notes, user["id"], utc_now_iso(), opportunity_id, user["organization_id"]),
        )
        conn.execute(
            """INSERT INTO opportunity_negotiations
            (organization_id,opportunity_id,action,amount,broker_response,notes,created_by)
            VALUES (?,?,?,?,?,?,?)""",
            (user["organization_id"], opportunity_id, "Counteroffer sent", amount, response, notes, user["id"]),
        )
    record_audit_event("opportunity.negotiated", int(user["organization_id"]), int(user["id"]), {"opportunity_id": opportunity_id, "counteroffer": amount})
    return redirect(f"/rate-quotes/{opportunity_id}")


@app.post("/rate-quotes/{opportunity_id}/decline")
async def decline_rate_quote(request: Request, opportunity_id: int):
    user = require_user(request)
    opportunity = opportunity_row(user["organization_id"], opportunity_id)
    if opportunity["status"] == "Booked":
        raise HTTPException(409, "A booked quote cannot be declined")
    form = await verified_form(request)
    notes = str(form.get("notes", "")).strip()
    with db_session() as conn:
        conn.execute(
            "UPDATE load_opportunities SET status='Declined',declined_at=?,updated_by=?,updated_at=? WHERE id=? AND organization_id=?",
            (utc_now_iso(), user["id"], utc_now_iso(), opportunity_id, user["organization_id"]),
        )
        conn.execute(
            """INSERT INTO opportunity_negotiations
            (organization_id,opportunity_id,action,notes,created_by) VALUES (?,?,?,?,?)""",
            (user["organization_id"], opportunity_id, "Declined", notes, user["id"]),
        )
    record_audit_event("opportunity.declined", int(user["organization_id"]), int(user["id"]), {"opportunity_id": opportunity_id})
    return redirect(f"/rate-quotes/{opportunity_id}")


@app.post("/rate-quotes/{opportunity_id}/book")
async def book_rate_quote(request: Request, opportunity_id: int):
    user = require_user(request)
    opportunity = opportunity_row(user["organization_id"], opportunity_id)
    if opportunity.get("booked_load_id") or opportunity["status"] == "Booked":
        raise HTTPException(409, "This opportunity has already been converted to a load")
    form = await verified_form(request)
    final_rate = optional_number(form.get("final_agreed_rate"))
    if final_rate is None or final_rate <= 0:
        final_rate = optional_number(opportunity.get("counteroffer_rate")) or number(opportunity["original_offered_rate"])
    driver_id = int(opportunity.get("selected_driver_id") or 0)
    if not driver_id:
        raise HTTPException(400, "Select a driver before booking")
    driver = query_one("SELECT * FROM drivers WHERE id=? AND organization_id=? AND active=1", (driver_id, user["organization_id"]))
    if not driver:
        raise HTTPException(400, "The selected driver is unavailable")
    vehicle_id = int(opportunity.get("selected_vehicle_id") or driver["vehicle_id"] or 0)
    if not vehicle_id or not query_one("SELECT 1 FROM vehicles WHERE id=? AND organization_id=? AND active=1", (vehicle_id, user["organization_id"])):
        raise HTTPException(400, "Assign an active power unit before booking")
    operational = opportunity_operational_warnings(user["organization_id"], opportunity, [dict(driver)])
    bundle, driver_data, _, fuel_price = calculate_saved_opportunity(
        user["organization_id"], opportunity, revenue_override=final_rate,
    )
    location_comparison = compare_drivers(
        bundle["settings"], [dict(driver)], opportunity,
        latest_driver_locations(user["organization_id"]), configured_route_provider(),
        fuel_price=fuel_price, operational_warnings=operational,
    )
    review_warnings = list(location_comparison[0]["result"].warnings) if location_comparison else operational.get(driver_id, [])
    bundle, driver_data, result, _ = calculate_saved_opportunity(
        user["organization_id"], opportunity, revenue_override=final_rate,
        extra_warnings=review_warnings,
    )
    if result.recommendation == "REVIEW REQUIRED" and not yes(form.get("confirm_review")):
        raise HTTPException(400, "Resolve the review warnings or confirm the reviewed override before booking")
    if result.loaded_miles <= 0:
        raise HTTPException(400, "Verified loaded miles are required before booking")
    due_at = (datetime.now(timezone.utc) + timedelta(hours=max(1, integer(bundle["settings"].get("ratecon_due_hours"), 4)))).replace(microsecond=0).isoformat()
    origin = ", ".join(part for part in (opportunity.get("origin_city"), opportunity.get("origin_state")) if part)
    destination = ", ".join(part for part in (opportunity.get("destination_city"), opportunity.get("destination_state")) if part)
    factoring_fee = final_rate * (number(opportunity.get("factoring_pct")) + number(opportunity.get("quick_pay_pct"))) / 100
    direct_costs = number(opportunity.get("lumper")) + number(opportunity.get("misc_expenses")) + factoring_fee
    with db_session() as conn:
        locked = conn.execute(
            "SELECT * FROM load_opportunities WHERE id=? AND organization_id=?",
            (opportunity_id, user["organization_id"]),
        ).fetchone()
        if not locked or locked["booked_load_id"] or locked["status"] == "Booked":
            raise HTTPException(409, "This opportunity has already been converted to a load")
        quote_snapshot = conn.execute(
            """SELECT id FROM opportunity_snapshots WHERE opportunity_id=? AND stage='evaluation'
            ORDER BY revision DESC,id DESC LIMIT 1""",
            (opportunity_id,),
        ).fetchone()
        booking_snapshot_id = insert_opportunity_snapshot(
            conn, opportunity, bundle["settings"], driver_data, result, "booking", user["id"], final_rate,
        )
        load_number = f"CO-{date.today():%Y}-{opportunity_id:05d}"
        load_public_uuid = str(uuid.uuid4())
        booked_at = utc_now_iso()
        load_id = int(conn.execute(
            """INSERT INTO loads
            (organization_id,public_uuid,status_code,updated_at,load_number,pickup_date,delivery_date,driver_id,vehicle_id,
             broker,origin,destination,status,revenue,loaded_miles,deadhead_miles,fuel_override,
             tolls_misc,other_direct_costs,notes,include_in_model,opportunity_id,
             original_offered_rate,final_agreed_rate,quote_snapshot_id,booking_snapshot_id,ratecon_due_at,
             pickup_window_start,delivery_window_start)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
            (
                user["organization_id"], load_public_uuid, LoadState.BOOKED_AWAITING_RATECON.value,
                booked_at, load_number, str(opportunity["pickup_at"])[:10],
                str(opportunity["delivery_at"])[:10], driver_id, vehicle_id,
                opportunity.get("broker_customer") or "", origin, destination,
                WORKFLOW_LABELS[LoadState.BOOKED_AWAITING_RATECON], final_rate, result.loaded_miles, result.deadhead_miles,
                None, number(opportunity.get("tolls")), direct_costs,
                opportunity.get("notes") or "", 1, opportunity_id,
                opportunity["original_offered_rate"], final_rate,
                int(quote_snapshot["id"]) if quote_snapshot else None, booking_snapshot_id, due_at,
                str(opportunity["pickup_at"])[:16], str(opportunity["delivery_at"])[:16],
            ),
        ).lastrowid)
        conn.execute(
            "UPDATE loads SET booking_snapshot_id=? WHERE id=?",
            (booking_snapshot_id, load_id),
        )
        conn.execute(
            """INSERT INTO load_status_history
            (organization_id,load_id,prior_status,new_status,changed_by,idempotency_key,reason)
            VALUES (?,?,NULL,?,?,?,?)""",
            (
                user["organization_id"], load_id, LoadState.BOOKED_AWAITING_RATECON.value,
                user["id"], f"booking:{opportunity_id}:initial", "Booked from approved rate quote",
            ),
        )
        for sequence, stop_type, city, state, postal, appointment in (
            (1, "PICKUP", opportunity.get("origin_city"), opportunity.get("origin_state"), opportunity.get("origin_postal_code"), opportunity.get("pickup_at")),
            (2, "DELIVERY", opportunity.get("destination_city"), opportunity.get("destination_state"), opportunity.get("destination_postal_code"), opportunity.get("delivery_at")),
        ):
            conn.execute(
                """INSERT INTO load_stops
                (public_uuid,organization_id,load_id,sequence_number,stop_type,city,state,postal_code,
                 appointment_local_start,detention_eligible)
                VALUES (?,?,?,?,?,?,?,?,?,1)""",
                (
                    str(uuid.uuid4()), user["organization_id"], load_id, sequence, stop_type,
                    city or "", state or "", postal or "", appointment or None,
                ),
            )
        conn.execute(
            """INSERT INTO load_revenue_items
            (public_uuid,organization_id,load_id,category,description,amount_cents,stage,source)
            VALUES (?,?,?,'LINEHAUL','Final agreed all-in rate',?,'BOOKED','quote_booking')""",
            (str(uuid.uuid4()), user["organization_id"], load_id, money_to_cents(final_rate, field="Final agreed rate")),
        )
        power_unit = conn.execute(
            "SELECT id FROM power_units WHERE organization_id=? AND legacy_vehicle_id=?",
            (user["organization_id"], vehicle_id),
        ).fetchone()
        if power_unit:
            conn.execute(
                """INSERT INTO load_assignments
                (public_uuid,organization_id,load_id,driver_id,power_unit_id,assignment_stage,
                 provisional,deadhead_miles_micros,route_source)
                VALUES (?,?,?, ?,?,'QUOTE_SELECTED',1,?,'booking_snapshot')""",
                (
                    str(uuid.uuid4()), user["organization_id"], load_id, driver_id, power_unit["id"],
                    round(result.deadhead_miles * 1_000_000),
                ),
            )
        _insert_load_snapshot(
            conn, organization_id=user["organization_id"], load_id=load_id, stage="BOOKED",
            inputs={"opportunity_id": opportunity_id, "opportunity_booking_snapshot_id": booking_snapshot_id},
            result=result.to_dict(), user_id=user["id"],
        )
        conn.execute(
            """UPDATE load_opportunities SET status='Booked',final_agreed_rate=?,booked_load_id=?,
            ratecon_due_at=?,booked_at=?,updated_by=?,updated_at=?
            WHERE id=? AND organization_id=? AND booked_load_id IS NULL""",
            (final_rate, load_id, due_at, utc_now_iso(), user["id"], utc_now_iso(), opportunity_id, user["organization_id"]),
        )
        conn.execute(
            """INSERT INTO opportunity_negotiations
            (organization_id,opportunity_id,action,amount,notes,created_by) VALUES (?,?,?,?,?,?)""",
            (user["organization_id"], opportunity_id, "Booked — Awaiting RateCon", final_rate, "Converted once to operational load", user["id"]),
        )
    record_audit_event("opportunity.booked", int(user["organization_id"]), int(user["id"]), {"opportunity_id": opportunity_id, "load_id": load_id, "final_rate": final_rate})
    set_flash(request, "Load booked. The original quote and final booking calculations are locked as audit snapshots; RateCon is now due.")
    return redirect(f"/loads/{load_id}")


@app.get("/rate-quotes/{opportunity_id}", response_class=HTMLResponse)
def rate_quote_detail(request: Request, opportunity_id: int):
    user = require_user(request)
    opportunity = opportunity_row(user["organization_id"], opportunity_id)
    bundle, driver, result, fuel_price = calculate_saved_opportunity(user["organization_id"], opportunity)
    locations = latest_driver_locations(user["organization_id"])
    comparison = compare_drivers(
        bundle["settings"], bundle["drivers"], opportunity, locations,
        configured_route_provider(), fuel_price=fuel_price,
        operational_warnings=opportunity_operational_warnings(user["organization_id"], opportunity, bundle["drivers"]),
    )
    selected_comparison = next(
        (row for row in comparison if int(row["driver"]["id"]) == int(opportunity.get("selected_driver_id") or 0)),
        None,
    )
    if selected_comparison:
        driver = selected_comparison["driver"]
        result = selected_comparison["result"]
    negotiations = [dict(row) for row in query_all(
        "SELECT * FROM opportunity_negotiations WHERE organization_id=? AND opportunity_id=? ORDER BY created_at DESC,id DESC",
        (user["organization_id"], opportunity_id),
    )]
    snapshots = [dict(row) for row in query_all(
        "SELECT id,stage,revision,created_at FROM opportunity_snapshots WHERE organization_id=? AND opportunity_id=? ORDER BY created_at DESC,id DESC",
        (user["organization_id"], opportunity_id),
    )]
    counter = result.opening_counteroffer or result.minimum_acceptable_rate or opportunity["original_offered_rate"]
    message = (
        f"We can cover {opportunity.get('origin_city')}, {opportunity.get('origin_state')} to "
        f"{opportunity.get('destination_city')}, {opportunity.get('destination_state')} for ${counter:,.0f}. "
        "Please confirm the final all-in rate and send the RateCon after we agree."
    )
    return render(request, "rate_quote_detail.html", {
        "opportunity": opportunity, "selected_driver": driver, "result": result,
        "comparison": comparison, "negotiations": negotiations, "snapshots": snapshots,
        "negotiation_message": message, "fuel_price": fuel_price,
    })


@app.post("/drivers/{driver_id}/location")
async def save_driver_location(request: Request, driver_id: int):
    user = require_user(request)
    driver = query_one("SELECT id FROM drivers WHERE id=? AND organization_id=?", (driver_id, user["organization_id"]))
    if not driver:
        raise HTTPException(404, "Driver not found")
    form = await verified_form(request)
    city = str(form.get("city", "")).strip()
    state = str(form.get("state", "")).strip().upper()
    if not city or not state:
        raise HTTPException(400, "Enter a city and state")
    observed_at = str(form.get("observed_at", "")).strip() or utc_now_iso()
    try:
        datetime.fromisoformat(observed_at.replace("Z", "+00:00"))
    except ValueError as exc:
        raise HTTPException(400, "Enter a valid location timestamp") from exc
    execute(
        """INSERT INTO driver_locations
        (organization_id,driver_id,city,state,postal_code,latitude,longitude,source,confidence,
         observed_at,override_reason,created_by) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            user["organization_id"], driver_id, city, state,
            str(form.get("postal_code", "")).strip(), optional_number(form.get("latitude")),
            optional_number(form.get("longitude")), "manual", "manual", observed_at,
            str(form.get("override_reason", "")).strip(), user["id"],
        ),
    )
    record_audit_event("driver.location_recorded", int(user["organization_id"]), int(user["id"]), {"driver_id": driver_id})
    set_flash(request, "Driver location recorded with its source and timestamp.")
    return_to = str(form.get("return_to") or "/drivers")
    if not return_to.startswith("/") or return_to.startswith("//"):
        return_to = "/drivers"
    return redirect(return_to)


@app.get("/quotes", response_class=HTMLResponse)
def quote_page(request: Request):
    user = require_user(request)
    bundle = get_bundle(user["organization_id"])
    current_fuel = fuel_price_for(date.today(), bundle["weekly_fuel"], number(bundle["settings"]["fallback_diesel_price"]))
    return render(request, "quotes.html", {
        "drivers": bundle["drivers"],
        "result": None,
        "defaults": {
            "pickup_date": date.today().isoformat(),
            "fuel_price": current_fuel,
            "trip_days": 1,
            "target_margin_pct": bundle["settings"]["target_margin_pct"],
        },
    })


@app.post("/quotes", response_class=HTMLResponse)
async def calculate_quote_page(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    bundle = get_bundle(user["organization_id"])
    driver = next((d for d in bundle["drivers"] if int(d["id"]) == integer(form.get("driver_id"))), None)
    pickup = parse_date(form.get("pickup_date"))
    if not driver or not pickup:
        raise HTTPException(400, "Select a driver and valid pickup date")
    result = calculate_quote(
        settings=bundle["settings"], driver=driver, pickup_date=pickup,
        loaded_miles=number(form.get("loaded_miles")), deadhead_miles=number(form.get("deadhead_miles")),
        trip_days=integer(form.get("trip_days"), 1), fuel_price=number(form.get("fuel_price")),
        tolls_misc=number(form.get("tolls_misc")), other_direct_costs=number(form.get("other_direct_costs")),
        quoted_revenue=optional_number(form.get("quoted_revenue")), target_margin_pct=number(form.get("target_margin_pct"), bundle["settings"]["target_margin_pct"]),
    )
    return render(request, "quotes.html", {
        "drivers": bundle["drivers"], "result": result, "submitted": dict(form), "selected_driver": driver,
        "defaults": {"pickup_date": pickup.isoformat(), "fuel_price": form.get("fuel_price"), "trip_days": form.get("trip_days"), "target_margin_pct": form.get("target_margin_pct")},
    })


@app.get("/financials", response_class=HTMLResponse)
def financials_page(request: Request):
    user = require_user(request)
    bundle, state = get_state(user["organization_id"])
    context = load_report_context(request, bundle, state)
    context.update({
        "rows": state["monthly_financials"],
        "owner": state["owner_pay"],
        "settings": bundle["settings"],
        "monthly_overhead": state["summary"]["monthly_company_overhead"],
        "filter_action": "/financials",
        "filter_page": "financials",
    })
    return render(request, "financials.html", context)


@app.get("/idle", response_class=HTMLResponse)
def idle_page(request: Request, month: str | None = None, edit_id: int | None = None):
    user = require_user(request)
    settings = query_one("SELECT default_report_month FROM organizations WHERE id=?", (user["organization_id"],))
    report_month = selected_month(month, settings["default_report_month"] if settings else None)
    bundle, state = get_state(user["organization_id"], report_month)
    editing = None
    if edit_id is not None:
        row = query_one(
            "SELECT * FROM idle_periods WHERE id=? AND organization_id=?",
            (edit_id, user["organization_id"]),
        )
        if not row:
            raise HTTPException(404, "Idle period not found")
        editing = dict(row)
    return render(request, "idle.html", {
        "month": report_month,
        "bridges": state["selected_month_bridges"],
        "periods": list(reversed(state["idle_state"]["periods"])),
        "drivers": bundle["drivers"],
        "vehicles": bundle["vehicles"],
        "editing": editing,
    })


@app.post("/idle")
async def add_idle_period(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    driver_id = integer(form.get("driver_id"))
    driver = query_one("SELECT * FROM drivers WHERE id=? AND organization_id=?", (driver_id, user["organization_id"]))
    if not driver:
        raise HTTPException(400, "Select a valid driver")
    vehicle_id = integer(form.get("vehicle_id")) or integer(driver["vehicle_id"])
    if vehicle_id and not query_one(
        "SELECT 1 FROM vehicles WHERE id=? AND organization_id=?",
        (vehicle_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid unit")
    idle_id = execute(
        """INSERT INTO idle_periods
        (organization_id,start_date,end_date,driver_id,vehicle_id,situation,include_in_model,notes)
        VALUES (?,?,?,?,?,?,?,?)""",
        (
            user["organization_id"], str(form.get("start_date", "")) or None, str(form.get("end_date", "")) or None,
            driver_id, vehicle_id or None, str(form.get("situation", "Company Responsibility")),
            1 if yes(form.get("include_in_model")) else 0, str(form.get("notes", "")).strip(),
        ),
    )
    record_audit_event(
        "idle_period.created",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"idle_period_id": idle_id},
    )
    month = str(form.get("start_date", ""))[:7]
    set_flash(request, "Idle/time-off period recorded. Load-covered days were excluded automatically.")
    return redirect(f"/idle?month={month}-01" if month else "/idle")


@app.post("/idle/{idle_id}/edit")
async def update_idle_period(request: Request, idle_id: int):
    user = require_user(request)
    form = await verified_form(request)
    existing = query_one(
        "SELECT * FROM idle_periods WHERE id=? AND organization_id=?",
        (idle_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Idle period not found")
    driver_id = integer(form.get("driver_id"))
    driver = query_one(
        "SELECT * FROM drivers WHERE id=? AND organization_id=?",
        (driver_id, user["organization_id"]),
    )
    if not driver:
        raise HTTPException(400, "Select a valid driver")
    vehicle_id = integer(form.get("vehicle_id")) or integer(driver["vehicle_id"])
    if vehicle_id and not query_one(
        "SELECT 1 FROM vehicles WHERE id=? AND organization_id=?",
        (vehicle_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid unit")
    start_date = str(form.get("start_date", ""))
    end_date = str(form.get("end_date", ""))
    if not parse_date(start_date) or not parse_date(end_date):
        raise HTTPException(400, "Enter valid start and end dates")
    execute(
        """UPDATE idle_periods SET start_date=?,end_date=?,driver_id=?,vehicle_id=?,situation=?,
        include_in_model=?,notes=? WHERE id=? AND organization_id=?""",
        (
            start_date, end_date, driver_id, vehicle_id or None,
            str(form.get("situation", "Company Responsibility")),
            1 if yes(form.get("include_in_model")) else 0,
            str(form.get("notes", "")).strip(), idle_id, user["organization_id"],
        ),
    )
    record_audit_event(
        "idle_period.updated",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"idle_period_id": idle_id},
    )
    set_flash(request, "Idle period updated and fixed-cost calculations were refreshed.")
    return redirect(f"/idle?month={start_date[:7]}-01")


@app.get("/settings", response_class=HTMLResponse)
def settings_page(request: Request):
    user = require_user(request)
    bundle = get_bundle(user["organization_id"])
    return render(request, "settings.html", {"settings": bundle["settings"], "overhead": bundle["overhead_items"]})


@app.post("/settings")
async def update_settings(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    with db_session() as conn:
        conn.execute(
            """UPDATE organizations SET name=?,owner_name=?,owner_email=?,fallback_diesel_price=?,
            processing_fee_pct=?,admin_fee_per_load=?,payroll_burden_pct=?,target_margin_pct=?,
            target_max_deadhead_pct=?,min_company_profit_per_mile=?,owner_distribution_pct=?,
            min_total_profit=?,min_profit_per_day=?,min_revenue_per_total_mile=?,
            quote_counteroffer_pct=?,ratecon_due_hours=?,location_stale_hours=?,default_payment_days=?,
            supported_start_date=?,supported_end_date=?,max_active_days=?,tax_reserve_pct=?,
            growth_reserve_pct=?,reporting_start_month=?,default_report_month=? WHERE id=?""",
            (
                str(form.get("name", "")).strip(), str(form.get("owner_name", "")).strip(), str(form.get("owner_email", "")).strip(),
                number(form.get("fallback_diesel_price")), number(form.get("processing_fee_pct")), number(form.get("admin_fee_per_load")),
                number(form.get("payroll_burden_pct")), number(form.get("target_margin_pct")), number(form.get("target_max_deadhead_pct")),
                number(form.get("min_company_profit_per_mile")), number(form.get("owner_distribution_pct")),
                number(form.get("min_total_profit"), 200), number(form.get("min_profit_per_day"), 100),
                number(form.get("min_revenue_per_total_mile"), 1.75), number(form.get("quote_counteroffer_pct"), 5),
                integer(form.get("ratecon_due_hours"), 4), integer(form.get("location_stale_hours"), 24),
                integer(form.get("default_payment_days"), 30),
                str(form.get("supported_start_date", "")), str(form.get("supported_end_date", "")), integer(form.get("max_active_days"), 31),
                number(form.get("tax_reserve_pct")), number(form.get("growth_reserve_pct")), str(form.get("reporting_start_month", "")),
                str(form.get("default_report_month", "")), user["organization_id"],
            ),
        )
        for row in query_all("SELECT id FROM overhead_items WHERE organization_id=?", (user["organization_id"],)):
            conn.execute("UPDATE overhead_items SET monthly_cost=? WHERE id=? AND organization_id=?", (number(form.get(f"overhead_{row['id']}")), row["id"], user["organization_id"]))
    set_flash(request, "Assumptions and company overhead updated. The entire model was recalculated.")
    return redirect("/settings")


@app.post("/settings/overhead")
async def add_overhead(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    order = query_one("SELECT COALESCE(MAX(sort_order),0)+1 AS n FROM overhead_items WHERE organization_id=?", (user["organization_id"],))
    execute("INSERT INTO overhead_items (organization_id,name,monthly_cost,sort_order) VALUES (?,?,?,?)", (user["organization_id"], str(form.get("name", "Other overhead")).strip(), number(form.get("monthly_cost")), integer(order["n"] if order else 1)))
    return redirect("/settings")


@app.post("/settings/export")
async def export_company_data(request: Request):
    user = require_user(request)
    if not user.get("is_admin"):
        raise HTTPException(status_code=403, detail="Administrator access required")
    await verified_form(request)
    payload = export_organization_data(int(user["organization_id"]))
    filename = f"carrieros-company-export-{date.today().isoformat()}.json"
    return Response(
        json.dumps(payload, indent=2, ensure_ascii=False),
        media_type="application/json",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/compliance", response_class=HTMLResponse)
def compliance_page(request: Request, edit_id: int | None = None):
    user = require_user(request)
    rows = query_all("SELECT * FROM compliance_items WHERE organization_id=? ORDER BY expiration_date", (user["organization_id"],))
    items = []
    for row in rows:
        item = dict(row)
        item["status"], item["days"] = compliance_status(row["expiration_date"])
        items.append(item)
    bundle = get_bundle(user["organization_id"])
    editing = None
    if edit_id is not None:
        row = query_one(
            "SELECT * FROM compliance_items WHERE id=? AND organization_id=?",
            (edit_id, user["organization_id"]),
        )
        if not row:
            raise HTTPException(404, "Compliance item not found")
        editing = dict(row)
    return render(request, "compliance.html", {"items": items, "vehicles": bundle["vehicles"], "drivers": bundle["drivers"], "editing": editing})


def validated_compliance_subject(user: dict[str, Any], form: Any) -> tuple[str, int | None]:
    subject_type = str(form.get("subject_type", "Company")).strip()
    if subject_type not in {"Company", "Vehicle", "Driver"}:
        raise HTTPException(400, "Select a valid subject type")
    subject_id = integer(form.get("subject_id")) or None
    if subject_type == "Company":
        return subject_type, None
    table = "drivers" if subject_type == "Driver" else "vehicles"
    if subject_id and not query_one(
        f"SELECT 1 FROM {table} WHERE id=? AND organization_id=?",
        (subject_id, user["organization_id"]),
    ):
        raise HTTPException(400, "Select a valid linked record")
    return subject_type, subject_id


@app.post("/compliance")
async def create_compliance_item(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    subject_type, subject_id = validated_compliance_subject(user, form)
    item_id = execute(
        """INSERT INTO compliance_items
        (organization_id,subject_type,subject_id,subject_name,document_type,expiration_date,notes)
        VALUES (?,?,?,?,?,?,?)""",
        (
            user["organization_id"], subject_type, subject_id,
            str(form.get("subject_name", user["organization_name"])).strip(), str(form.get("document_type", "Document")).strip(),
            str(form.get("expiration_date", "")) or None, str(form.get("notes", "")).strip(),
        ),
    )
    record_audit_event(
        "compliance_item.created",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"compliance_item_id": item_id},
    )
    return redirect("/compliance")


@app.post("/compliance/{item_id}/edit")
async def update_compliance_item(request: Request, item_id: int):
    user = require_user(request)
    form = await verified_form(request)
    if not query_one(
        "SELECT 1 FROM compliance_items WHERE id=? AND organization_id=?",
        (item_id, user["organization_id"]),
    ):
        raise HTTPException(404, "Compliance item not found")
    subject_type, subject_id = validated_compliance_subject(user, form)
    execute(
        """UPDATE compliance_items SET subject_type=?,subject_id=?,subject_name=?,document_type=?,
        expiration_date=?,notes=? WHERE id=? AND organization_id=?""",
        (
            subject_type, subject_id, str(form.get("subject_name", user["organization_name"])).strip(),
            str(form.get("document_type", "Document")).strip(), str(form.get("expiration_date", "")) or None,
            str(form.get("notes", "")).strip(), item_id, user["organization_id"],
        ),
    )
    record_audit_event(
        "compliance_item.updated",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"compliance_item_id": item_id},
    )
    set_flash(request, "Compliance item updated.")
    return redirect("/compliance")


@app.post("/compliance/{item_id}/delete")
async def delete_compliance_item(request: Request, item_id: int):
    user = require_user(request)
    await verified_form(request)
    existing = query_one(
        "SELECT * FROM compliance_items WHERE id=? AND organization_id=?",
        (item_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Compliance item not found")
    execute(
        "DELETE FROM compliance_items WHERE id=? AND organization_id=?",
        (item_id, user["organization_id"]),
    )
    record_audit_event(
        "compliance_item.deleted",
        organization_id=int(user["organization_id"]), user_id=int(user["id"]),
        details={"compliance_item_id": item_id, "document_type": str(existing["document_type"])},
    )
    set_flash(request, "Compliance item removed.")
    return redirect("/compliance")


@app.get("/onboarding", response_class=HTMLResponse)
def onboarding_page(request: Request):
    user = require_user(request)
    rows = query_all("SELECT * FROM onboarding_applications WHERE organization_id=? ORDER BY created_at DESC", (user["organization_id"],))
    return render(request, "onboarding.html", {"applications": [dict(r) for r in rows]})


@app.post("/onboarding/invite")
async def onboarding_invite(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    token = new_onboarding_token()
    expires_at = (
        datetime.now(timezone.utc) + timedelta(days=14)
    ).replace(microsecond=0).isoformat()
    execute(
        """INSERT INTO onboarding_applications
        (organization_id,token,full_name,email,status,expires_at)
        VALUES (?,?,?,?, 'Invited', ?)""",
        (
            user["organization_id"], token_digest(token),
            str(form.get("full_name", "")).strip(),
            str(form.get("email", "")).strip(), expires_at,
        ),
    )
    set_flash(request, f"Onboarding link created (expires in 14 days): /onboard/{token}")
    return redirect("/onboarding")


@app.get("/onboard/{token}", response_class=HTMLResponse)
def public_onboarding_page(request: Request, token: str):
    row = query_one(
        """SELECT a.*,o.name AS organization_name FROM onboarding_applications a
        JOIN organizations o ON o.id=a.organization_id WHERE a.token IN (?, ?)""",
        (token, token_digest(token)),
    )
    if not row:
        raise HTTPException(404, "Onboarding link not found")
    if row["expires_at"] and str(row["expires_at"]) < utc_now_iso():
        raise HTTPException(410, "Onboarding link has expired")
    return render(request, "public_onboarding.html", {"application": dict(row), "public_page": True})


@app.post("/onboard/{token}", response_class=HTMLResponse)
async def submit_onboarding(request: Request, token: str):
    row = query_one(
        "SELECT * FROM onboarding_applications WHERE token IN (?, ?)",
        (token, token_digest(token)),
    )
    if not row:
        raise HTTPException(404, "Onboarding link not found")
    if row["expires_at"] and str(row["expires_at"]) < utc_now_iso():
        raise HTTPException(410, "Onboarding link has expired")
    if str(row["status"]).lower() == "submitted":
        organization = query_one(
            "SELECT name FROM organizations WHERE id=?", (row["organization_id"],)
        )
        application = dict(row)
        application["organization_name"] = organization["name"] if organization else "the carrier"
        return render(
            request,
            "public_onboarding.html",
            {"application": application, "submitted": True, "public_page": True},
        )
    form = await verified_form(request)
    with db_session() as conn:
        conn.execute(
            """UPDATE onboarding_applications SET full_name=?,email=?,phone=?,license_number=?,
            license_state=?,medical_card_expiration=?,employment_history=?,emergency_contact=?,
            status='Submitted',submitted_at=? WHERE id=? AND organization_id=?""",
            (
                str(form.get("full_name", "")).strip(), str(form.get("email", "")).strip(), str(form.get("phone", "")).strip(),
                str(form.get("license_number", "")).strip(), str(form.get("license_state", "")).strip(),
                str(form.get("medical_card_expiration", "")) or None, str(form.get("employment_history", "")).strip(),
                str(form.get("emergency_contact", "")).strip(), utc_now_iso(),
                row["id"], row["organization_id"],
            ),
        )
    updated = query_one(
        """SELECT a.*,o.name AS organization_name FROM onboarding_applications a
        JOIN organizations o ON o.id=a.organization_id WHERE a.id=?""", (row["id"],),
    )
    return render(request, "public_onboarding.html", {"application": dict(updated), "submitted": True, "public_page": True})


AUDIT_DOCUMENT_LABELS = {
    "ratecon": "Rate confirmation",
    "bank_statement": "Business bank statement",
    "bill_statement": "Business bill statement",
}


def _audit_page_context(organization_id: int) -> dict[str, Any]:
    loads = query_all(
        """SELECT id,load_number,pickup_date,delivery_date,status,revenue
        FROM loads WHERE organization_id=? ORDER BY pickup_date DESC,id DESC""",
        (organization_id,),
    )
    overhead = query_all(
        "SELECT id,name,monthly_cost FROM overhead_items WHERE organization_id=? ORDER BY sort_order,id",
        (organization_id,),
    )
    audits = query_all(
        """SELECT a.*,l.load_number,o.name AS overhead_name
        FROM document_audits a
        LEFT JOIN loads l ON l.id=a.linked_load_id AND l.organization_id=a.organization_id
        LEFT JOIN overhead_items o ON o.id=a.overhead_item_id AND o.organization_id=a.organization_id
        WHERE a.organization_id=? ORDER BY a.created_at DESC,a.id DESC""",
        (organization_id,),
    )
    return {
        "loads": [dict(row) for row in loads],
        "overhead_items": [dict(row) for row in overhead],
        "audits": [dict(row) for row in audits],
        "document_types": AUDIT_DOCUMENT_LABELS,
        "max_upload_mb": MAX_AUDIT_FILE_BYTES // (1024 * 1024),
    }


@app.get("/audits", response_class=HTMLResponse)
def audits_page(request: Request):
    user = require_user(request)
    return render(request, "audits.html", _audit_page_context(int(user["organization_id"])))


@app.post("/audits/upload")
async def upload_audit_document(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    document_type = str(form.get("document_type") or "").strip()
    upload = form.get("document")
    filename = str(getattr(upload, "filename", "") or "").strip()
    content_type = str(getattr(upload, "content_type", "") or "")
    if not upload or not filename or not hasattr(upload, "read"):
        raise HTTPException(400, "Choose a PDF or CSV file to audit")
    payload = await upload.read(MAX_AUDIT_FILE_BYTES + 1)
    if hasattr(upload, "close"):
        await upload.close()

    organization_id = int(user["organization_id"])
    linked_load_id = integer(form.get("linked_load_id")) or None
    overhead_item_id = integer(form.get("overhead_item_id")) or None
    period_start = parse_date(form.get("period_start"))
    period_end = parse_date(form.get("period_end"))
    context: dict[str, Any] = {}
    expected_amount = 0.0

    if document_type == "bank_statement":
        if not period_start or not period_end:
            raise HTTPException(400, "Choose the beginning and ending date for the bank statement")
        if period_end < period_start:
            raise HTTPException(400, "The statement ending date must be on or after its beginning date")
        bundle, state = get_state(organization_id)
        period_rows = filter_and_sort_loads(
            loads_with_results(bundle, state),
            date_from=period_start,
            date_to=period_end,
            date_field="delivery_date",
        )
        summary = summarize_load_rows(period_rows)
        expected_amount = float(summary["revenue"])
        context = {
            "expected_revenue": expected_amount,
            "expected_expense": float(summary["operating_expense"]),
            "period_label": f"{period_start.isoformat()} through {period_end.isoformat()}",
        }
    elif document_type == "ratecon":
        if not linked_load_id:
            raise HTTPException(400, "Select the CarrierOS load that belongs to this RateCon")
        bundle, state = get_state(organization_id)
        load = next(
            (
                row for row in loads_with_results(bundle, state)
                if int(row["id"]) == linked_load_id
            ),
            None,
        )
        if not load:
            raise HTTPException(404, "Load not found")
        dispatch = driver_dispatch_package(load)
        expected_amount = float(load.get("revenue") or 0)
        context = {
            "expected_revenue": expected_amount,
            "load_number": load.get("load_number"),
            "dispatch_blockers": [] if dispatch.get("ready") else dispatch.get("blockers", []),
        }
    elif document_type == "bill_statement":
        overhead = None
        if overhead_item_id:
            overhead = query_one(
                "SELECT * FROM overhead_items WHERE id=? AND organization_id=?",
                (overhead_item_id, organization_id),
            )
            if not overhead:
                raise HTTPException(404, "Overhead estimate not found")
        expected_amount = float(overhead["monthly_cost"] if overhead else optional_number(form.get("expected_amount")) or 0)
        context = {
            "expected_amount": expected_amount,
            "estimate_name": str(overhead["name"] if overhead else "the entered monthly estimate"),
        }

    try:
        result = audit_uploaded_document(
            document_type=document_type,
            filename=filename,
            content_type=content_type,
            payload=payload,
            context=context,
        )
    except AuditFileError as exc:
        raise HTTPException(400, str(exc)) from exc

    duplicate = query_one(
        """SELECT id FROM document_audits
        WHERE organization_id=? AND document_type=? AND sha256=? ORDER BY id DESC LIMIT 1""",
        (organization_id, document_type, result["sha256"]),
    )
    if duplicate:
        set_flash(request, "That exact file was already audited. CarrierOS opened the existing result.")
        return redirect(f"/audits/{duplicate['id']}")

    extracted = result.get("extracted") or {}
    if document_type == "bank_statement":
        observed_amount = extracted.get("deposits")
    elif document_type == "ratecon":
        observed_amount = extracted.get("rate")
    else:
        observed_amount = extracted.get("amount_due")
    audit_id = execute(
        """INSERT INTO document_audits
        (organization_id,document_type,original_filename,content_type,size_bytes,sha256,status,
         linked_load_id,overhead_item_id,period_start,period_end,expected_amount,observed_amount,
         summary,extracted_json,findings_json,created_by)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (
            organization_id, document_type, result["filename"], result["content_type"],
            result["size_bytes"], result["sha256"], result["status"], linked_load_id,
            overhead_item_id, period_start.isoformat() if period_start else None,
            period_end.isoformat() if period_end else None, expected_amount,
            observed_amount, result["summary"],
            json.dumps({
                "extracted": result.get("extracted") or {},
                "metrics": result.get("metrics") or [],
                "raw_file_retained": False,
            }, separators=(",", ":")),
            json.dumps(result.get("findings") or [], separators=(",", ":")),
            int(user["id"]),
        ),
    )
    if document_type == "ratecon" and yes(form.get("mark_ratecon_verified")) and linked_load_id:
        load_status = query_one(
            "SELECT status,ratecon_received_at FROM loads WHERE id=? AND organization_id=?",
            (linked_load_id, organization_id),
        )
        if load_status:
            status = str(load_status["status"] or "Booked")
            if "Awaiting RateCon" in status:
                status = "RateCon Received"
            execute(
                """UPDATE loads SET ratecon_reference=?,ratecon_received_at=?,status=?
                WHERE id=? AND organization_id=?""",
                (
                    result["filename"], load_status["ratecon_received_at"] or utc_now_iso(),
                    status, linked_load_id, organization_id,
                ),
            )
    record_audit_event(
        "document.audit_created", organization_id, int(user["id"]),
        {"audit_id": audit_id, "document_type": document_type, "raw_file_retained": False},
    )
    set_flash(request, "Audit complete. Review every extracted figure before changing a company estimate.")
    return redirect(f"/audits/{audit_id}")


@app.get("/audits/{audit_id}", response_class=HTMLResponse)
def audit_detail(request: Request, audit_id: int):
    user = require_user(request)
    row = query_one(
        """SELECT a.*,l.load_number,o.name AS overhead_name
        FROM document_audits a
        LEFT JOIN loads l ON l.id=a.linked_load_id AND l.organization_id=a.organization_id
        LEFT JOIN overhead_items o ON o.id=a.overhead_item_id AND o.organization_id=a.organization_id
        WHERE a.id=? AND a.organization_id=?""",
        (audit_id, user["organization_id"]),
    )
    if not row:
        raise HTTPException(404, "Audit not found")
    item = dict(row)
    try:
        extracted = json.loads(item.get("extracted_json") or "{}")
        findings = json.loads(item.get("findings_json") or "[]")
    except json.JSONDecodeError:
        extracted, findings = {}, []
    return render(request, "audit_detail.html", {
        "audit": item,
        "extracted": extracted.get("extracted") or {},
        "metrics": extracted.get("metrics") or [],
        "findings": findings,
        "document_label": AUDIT_DOCUMENT_LABELS.get(item["document_type"], item["document_type"]),
    })


@app.post("/audits/{audit_id}/delete")
async def delete_audit_record(request: Request, audit_id: int):
    user = require_user(request)
    await verified_form(request)
    existing = query_one(
        "SELECT id,document_type FROM document_audits WHERE id=? AND organization_id=?",
        (audit_id, user["organization_id"]),
    )
    if not existing:
        raise HTTPException(404, "Audit not found")
    execute(
        "DELETE FROM document_audits WHERE id=? AND organization_id=?",
        (audit_id, user["organization_id"]),
    )
    record_audit_event(
        "document.audit_deleted", int(user["organization_id"]), int(user["id"]),
        {"audit_id": audit_id, "document_type": existing["document_type"]},
    )
    set_flash(request, "The audit result and checksum metadata were deleted. No raw file was retained.")
    return redirect("/audits")


def _growth_page_context(organization_id: int, values: dict[str, Any] | None = None) -> dict[str, Any]:
    bundle, state = get_state(organization_id)
    start = date.today() - timedelta(days=89)
    rows = filter_and_sort_loads(
        loads_with_results(bundle, state), date_from=start, date_to=date.today(), date_field="delivery_date"
    )
    summary = summarize_load_rows(rows)
    active_units = sum(1 for vehicle in bundle["vehicles"] if int(vehicle.get("active") or 0) == 1)
    driver_mpg = [float(driver.get("mpg") or 0) for driver in bundle["drivers"] if float(driver.get("mpg") or 0) > 0]
    defaults = {
        "purchase_price": 75000,
        "down_payment": 15000,
        "apr_pct": 9.0,
        "term_months": 60,
        "monthly_insurance": 1800,
        "other_monthly_costs": 500,
        "monthly_miles": 8500,
        "revenue_per_mile": round(float(summary.get("avg_revenue_per_mile") or 2.5), 2),
        "mpg": round(sum(driver_mpg) / len(driver_mpg), 1) if driver_mpg else 8.0,
        "diesel_price": float(bundle["settings"].get("fallback_diesel_price") or 4.0),
        "maintenance_per_mile": 0.20,
        "driver_pay_pct": 0,
        "cash_reserve": 30000,
    }
    return {
        "summary": summary,
        "active_units": active_units,
        "mentor_findings": growth_mentor_findings(summary, bundle["settings"], active_units),
        "values": values or defaults,
        "settings": bundle["settings"],
        "target_margin_decimal": float(bundle["settings"].get("target_margin_pct") or 10.0) / 100.0,
        "window_start": start,
    }


@app.get("/growth", response_class=HTMLResponse)
def growth_page(request: Request):
    user = require_user(request)
    return render(request, "growth.html", _growth_page_context(int(user["organization_id"])))


@app.post("/growth", response_class=HTMLResponse)
async def growth_audit(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    values = {
        "purchase_price": number(form.get("purchase_price")),
        "down_payment": number(form.get("down_payment")),
        "apr_pct": number(form.get("apr_pct")),
        "term_months": max(1, integer(form.get("term_months"), 60)),
        "monthly_insurance": number(form.get("monthly_insurance")),
        "other_monthly_costs": number(form.get("other_monthly_costs")),
        "monthly_miles": number(form.get("monthly_miles")),
        "revenue_per_mile": number(form.get("revenue_per_mile")),
        "mpg": number(form.get("mpg"), 8),
        "diesel_price": number(form.get("diesel_price")),
        "maintenance_per_mile": number(form.get("maintenance_per_mile")),
        "driver_pay_pct": number(form.get("driver_pay_pct")),
        "cash_reserve": number(form.get("cash_reserve")),
    }
    context = _growth_page_context(int(user["organization_id"]), values)
    context["scenario"] = equipment_finance_audit(values, context["settings"])
    record_audit_event(
        "growth.equipment_audited", int(user["organization_id"]), int(user["id"]),
        {"purchase_price": values["purchase_price"], "financed_amount": context["scenario"]["financed_amount"]},
    )
    return render(request, "growth.html", context)


@app.get("/startup", response_class=HTMLResponse)
def startup_page(request: Request):
    user = require_user(request)
    rows = query_all(
        "SELECT * FROM startup_checklist_progress WHERE organization_id=?",
        (user["organization_id"],),
    )
    progress = {str(row["item_key"]): dict(row) for row in rows}
    steps = [{**step, "progress": progress.get(step["key"])} for step in STARTUP_STEPS]
    completed = sum(1 for step in steps if step["progress"] and step["progress"].get("completed_at"))
    return render(request, "startup.html", {
        "steps": steps,
        "completed": completed,
        "total": len(steps),
        "progress_pct": completed / len(steps) if steps else 0,
    })


@app.post("/startup/{item_key}/toggle")
async def toggle_startup_step(request: Request, item_key: str):
    user = require_user(request)
    form = await verified_form(request)
    allowed = {step["key"] for step in STARTUP_STEPS}
    if item_key not in allowed:
        raise HTTPException(404, "Startup step not found")
    completed = yes(form.get("completed"))
    notes = str(form.get("notes") or "").strip()
    if len(notes) > 2000:
        raise HTTPException(400, "Startup checklist notes are limited to 2,000 characters")
    with db_session() as conn:
        conn.execute(
            """INSERT INTO startup_checklist_progress
            (organization_id,item_key,completed_at,notes,updated_by,updated_at)
            VALUES (?,?,?,?,?,?)
            ON CONFLICT(organization_id,item_key) DO UPDATE SET
              completed_at=excluded.completed_at,notes=excluded.notes,
              updated_by=excluded.updated_by,updated_at=excluded.updated_at""",
            (
                user["organization_id"], item_key, utc_now_iso() if completed else None,
                notes, user["id"], utc_now_iso(),
            ),
        )
    record_audit_event(
        "startup.step_updated", int(user["organization_id"]), int(user["id"]),
        {"item_key": item_key, "completed": completed},
    )
    return redirect(f"/startup#{item_key}")


DOCUMENT_TEMPLATES = [
    ("driver_handbook", "Driver handbook"),
    ("independent_contractor", "Independent-contractor template"),
    ("inspection_policy", "Vehicle inspection policy"),
    ("hours_of_service", "Hours-of-service policy"),
    ("load_securement", "Load-securement acknowledgement"),
    ("accident_response", "Accident-response procedure"),
    ("advance_authorization", "Advance and deduction authorization"),
    ("equipment_checkout", "Equipment checkout form"),
    ("detention_request", "Detention request"),
    ("payment_demand", "Past-due payment notice"),
]


def generate_document(document_type: str, company: str, effective_date: str, state: str, contact: str, notes: str) -> tuple[str, str]:
    labels = dict(DOCUMENT_TEMPLATES)
    title = labels.get(document_type, "Company document")
    intro = f"{company}\n{title.upper()}\nEffective date: {effective_date}\nPrimary state: {state}\nContact: {contact}\n"
    sections = {
        "driver_handbook": """PURPOSE\nThis handbook summarizes dispatch, safety, communication, equipment-care, documentation, payment, and incident-reporting expectations.\n\nCORE EXPECTATIONS\nDrivers must operate lawfully, protect cargo and equipment, communicate delays promptly, preserve shipping documents, and report accidents or damage immediately.\n\nPAY AND SETTLEMENTS\nPay follows the written compensation arrangement and completed-load documentation. Advances, reimbursements, non-load bonuses, and deductions must be identified separately.""",
        "independent_contractor": """SERVICES\nThe contractor may accept transportation work under separately confirmed terms and remains responsible for duties allocated by the signed master agreement.\n\nCOMPENSATION\nCompensation, expense responsibility, insurance, equipment, chargebacks, and termination terms must be completed in the signed agreement before work begins.""",
        "inspection_policy": """POLICY\nA pre-trip and post-trip inspection is required. Safety defects must be reported before continued operation. Inspection, repair, and out-of-service records must be retained as required.""",
        "hours_of_service": """POLICY\nNo dispatch instruction authorizes a driver to violate hours-of-service requirements. Drivers must keep records current and report when available hours are insufficient.""",
        "load_securement": """ACKNOWLEDGEMENT\nThe driver is responsible for inspecting cargo securement before movement and during transit at required intervals. Equipment must be appropriate, serviceable, and used according to applicable rules and shipper instructions.""",
        "accident_response": """IMMEDIATE STEPS\nProtect life, contact emergency services when needed, secure the scene, notify the company, avoid admitting fault, photograph conditions, collect witness information, and preserve all documents and electronic records.""",
        "advance_authorization": """AUTHORIZATION\nThe payee requests or acknowledges the listed advance, reimbursement, or deduction. Whether it reduces load pay must be marked explicitly on the settlement record and supported by the governing agreement.""",
        "equipment_checkout": """EQUIPMENT RECORD\nRecord the unit, trailer, keys, cards, permits, straps, chains, tarps, tools, and condition at issue and return. Damage or missing items must be documented with photographs.""",
        "detention_request": """SUBJECT: Detention request\nPlease process detention for the referenced shipment. Attach the rate confirmation, arrival/departure evidence, bill of lading, and calculation.\n\nLoad: [ENTER]\nArrival: [ENTER]\nDeparture: [ENTER]\nFree time: [ENTER]\nRequested amount: [ENTER]""",
        "payment_demand": """SUBJECT: Past-due transportation invoice\nOur records show the referenced invoice remains unpaid. Please confirm payment status, remittance date, or the specific documentation issue preventing payment.\n\nInvoice: [ENTER]\nLoad: [ENTER]\nBalance: [ENTER]\nDue date: [ENTER]""",
    }
    disclaimer = "\n\nREVIEW NOTICE\nThis operational template must be customized and reviewed by qualified legal, tax, insurance, safety, or compliance professionals before reliance."
    body = intro + "\n" + sections.get(document_type, notes or "[ENTER CONTENT]")
    if notes:
        body += f"\n\nCUSTOM NOTES\n{notes}"
    return title, (body + disclaimer).strip()


@app.get("/documents", response_class=HTMLResponse)
def documents_page(request: Request):
    user = require_user(request)
    rows = query_all("SELECT * FROM generated_documents WHERE organization_id=? ORDER BY created_at DESC", (user["organization_id"],))
    return render(request, "documents.html", {"documents": [dict(r) for r in rows], "document_types": DOCUMENT_TEMPLATES})


@app.post("/documents/generate")
async def create_document(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    doc_type = str(form.get("document_type", "driver_handbook"))
    title, body = generate_document(
        doc_type, str(form.get("company", user["organization_name"])), str(form.get("effective_date", date.today().isoformat())),
        str(form.get("state", "Tennessee")), str(form.get("contact", user["email"])), str(form.get("notes", "")).strip(),
    )
    doc_id = execute("INSERT INTO generated_documents (organization_id,title,document_type,body) VALUES (?,?,?,?)", (user["organization_id"], title, doc_type, body))
    return redirect(f"/documents/{doc_id}")


@app.get("/documents/{doc_id}", response_class=HTMLResponse)
def document_detail(request: Request, doc_id: int):
    user = require_user(request)
    row = query_one("SELECT * FROM generated_documents WHERE id=? AND organization_id=?", (doc_id, user["organization_id"]))
    if not row:
        raise HTTPException(404, "Document not found")
    return render(request, "document_detail.html", {"document": dict(row)})


@app.get("/documents/{doc_id}/docx")
def download_document_docx(request: Request, doc_id: int):
    user = require_user(request)
    row = query_one("SELECT * FROM generated_documents WHERE id=? AND organization_id=?", (doc_id, user["organization_id"]))
    if not row:
        raise HTTPException(404, "Document not found")
    doc = Document()
    doc.add_heading(row["title"], level=0)
    for block in row["body"].split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.isupper() and len(block) < 80:
            doc.add_heading(block.title(), level=1)
        else:
            doc.add_paragraph(block)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = "_".join(row["title"].lower().split()) + ".docx"
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@app.get("/receivables", response_class=HTMLResponse)
def receivables_page(request: Request):
    user = require_user(request)
    invoices = [invoice_aging(r) for r in query_all("SELECT * FROM invoices WHERE organization_id=? ORDER BY due_date", (user["organization_id"],))]
    claims = query_all(
        """SELECT c.*,l.load_number FROM detention_claims c LEFT JOIN loads l ON l.id=c.load_id
        WHERE c.organization_id=? ORDER BY c.created_at DESC""", (user["organization_id"],),
    )
    loads = query_all("SELECT id,load_number,broker FROM loads WHERE organization_id=? ORDER BY pickup_date DESC,id DESC", (user["organization_id"],))
    return render(request, "receivables.html", {"invoices": invoices, "claims": [dict(c) for c in claims], "loads": [dict(load_row) for load_row in loads]})


@app.post("/invoices")
async def create_invoice(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    invoice_date = str(form.get("invoice_date", date.today().isoformat()))
    due_date = str(form.get("due_date", (date.today() + timedelta(days=30)).isoformat()))
    execute(
        """INSERT INTO invoices
        (organization_id,invoice_number,customer,load_id,amount,invoice_date,due_date,status,notes)
        VALUES (?,?,?,?,?,?,?,'Unpaid',?)""",
        (user["organization_id"], str(form.get("invoice_number", "")).strip(), str(form.get("customer", "")).strip(), integer(form.get("load_id")) or None, number(form.get("amount")), invoice_date, due_date, str(form.get("notes", "")).strip()),
    )
    return redirect("/receivables")


@app.post("/invoices/{invoice_id}/paid")
async def mark_invoice_paid(request: Request, invoice_id: int):
    user = require_user(request)
    form = await verified_form(request)
    with db_session() as conn:
        conn.execute("UPDATE invoices SET status='Paid',paid_date=? WHERE id=? AND organization_id=?", (str(form.get("paid_date", date.today().isoformat())), invoice_id, user["organization_id"]))
    return redirect("/receivables")


@app.post("/detention")
async def create_detention_claim(request: Request):
    user = require_user(request)
    form = await verified_form(request)
    arrival_raw = str(form.get("arrival_at", ""))
    departure_raw = str(form.get("departure_at", ""))
    try:
        arrival = datetime.fromisoformat(arrival_raw)
        departure = datetime.fromisoformat(departure_raw)
    except ValueError as exc:
        raise HTTPException(400, "Enter valid arrival and departure times") from exc
    total_minutes = max(0.0, (departure - arrival).total_seconds() / 60)
    free_minutes = integer(form.get("free_minutes"), 120)
    rate = number(form.get("hourly_rate"), 75)
    amount = round(max(0.0, total_minutes - free_minutes) / 60 * rate, 2)
    execute(
        """INSERT INTO detention_claims
        (organization_id,load_id,broker,arrival_at,departure_at,free_minutes,hourly_rate,amount_claimed,status,notes)
        VALUES (?,?,?,?,?,?,?,?,'Draft',?)""",
        (user["organization_id"], integer(form.get("load_id")) or None, str(form.get("broker", "")).strip(), arrival_raw, departure_raw, free_minutes, rate, amount, str(form.get("notes", "")).strip()),
    )
    return redirect("/receivables")


@app.get("/detention/{claim_id}/draft", response_class=HTMLResponse)
def detention_draft(request: Request, claim_id: int):
    user = require_user(request)
    row = query_one(
        """SELECT c.*,l.load_number FROM detention_claims c LEFT JOIN loads l ON l.id=c.load_id
        WHERE c.id=? AND c.organization_id=?""", (claim_id, user["organization_id"]),
    )
    if not row:
        raise HTTPException(404, "Claim not found")
    total_minutes = (datetime.fromisoformat(row["departure_at"]) - datetime.fromisoformat(row["arrival_at"])).total_seconds() / 60
    billable = max(0.0, total_minutes - row["free_minutes"])
    draft = (
        f"Subject: Detention request - Load {row['load_number'] or 'N/A'}\n\n"
        f"Please process detention for load {row['load_number'] or 'N/A'}. The driver arrived at {row['arrival_at']} and departed at {row['departure_at']}. "
        f"After {row['free_minutes']} minutes of free time, {billable / 60:.2f} hours were billable at {money(row['hourly_rate'])} per hour. "
        f"The requested detention amount is {money(row['amount_claimed'])}.\n\n"
        "Please confirm receipt and advise the expected payment date. Attach the rate confirmation and supporting timestamps before sending."
    )
    return render(request, "detention_draft.html", {"claim": dict(row), "draft": draft})


@app.post("/api/quote")
async def api_quote(request: Request):
    user = require_user(request)
    payload = await request.json()
    bundle = get_bundle(user["organization_id"])
    driver = next((d for d in bundle["drivers"] if int(d["id"]) == integer(payload.get("driver_id"))), None)
    pickup = parse_date(payload.get("pickup_date"))
    if not driver or not pickup:
        return JSONResponse({"error": "Invalid driver or pickup date"}, status_code=400)
    result = calculate_quote(
        settings=bundle["settings"], driver=driver, pickup_date=pickup,
        loaded_miles=number(payload.get("loaded_miles")), deadhead_miles=number(payload.get("deadhead_miles")),
        trip_days=integer(payload.get("trip_days"), 1), fuel_price=number(payload.get("fuel_price")),
        tolls_misc=number(payload.get("tolls_misc")), other_direct_costs=number(payload.get("other_direct_costs")),
        quoted_revenue=optional_number(payload.get("quoted_revenue")), target_margin_pct=number(payload.get("target_margin_pct"), bundle["settings"]["target_margin_pct"]),
    )
    return result.to_dict()
